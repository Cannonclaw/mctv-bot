"""Word document builder with MCTV Elite Advertising branding."""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
import subprocess
import shutil


# ── Color Schemes ────────────────────────────────────────────────────────────
# Each scheme defines the colors used throughout the proposal.
# Keys: primary (headers/bg), accent (highlights/bullets), text (body),
#        light (alt-row bg), bg_hex (cover bg hex), accent_hex (border hex),
#        light_hex (callout bg hex), cover_logo (filename in assets/branding/)

COLOR_SCHEMES = {
    "original": {
        "label": "Original Primary",
        "primary":    RGBColor(0x1B, 0x1F, 0x3B),  # Navy
        "accent":     RGBColor(0xC5, 0xA5, 0x5A),  # Gold
        "white":      RGBColor(0xFF, 0xFF, 0xFF),
        "gray":       RGBColor(0x80, 0x80, 0x80),
        "text":       RGBColor(0x33, 0x33, 0x33),
        "light":      RGBColor(0xF5, 0xF5, 0xF5),
        "bg_hex":     "1B1F3B",
        "accent_hex": "C5A55A",
        "light_hex":  "F0EDE4",
        "cover_logo": "mctv_logo_on_navy.png",
    },
    "light": {
        "label": "Light, Bright & Airy",
        "primary":    RGBColor(0x2E, 0x5E, 0x86),  # Sky blue
        "accent":     RGBColor(0xE8, 0x9E, 0x3C),  # Warm amber
        "white":      RGBColor(0xFF, 0xFF, 0xFF),
        "gray":       RGBColor(0x99, 0x99, 0x99),
        "text":       RGBColor(0x3A, 0x3A, 0x3A),
        "light":      RGBColor(0xF8, 0xF9, 0xFA),
        "bg_hex":     "2E5E86",
        "accent_hex": "E89E3C",
        "light_hex":  "F0F6FB",
        "cover_logo": "mctv_logo_on_light.png",
    },
    "dark": {
        "label": "Dark & Sophisticated",
        "primary":    RGBColor(0x1A, 0x1A, 0x2E),  # Deep charcoal-navy
        "accent":     RGBColor(0xD4, 0xAF, 0x37),  # Rich gold
        "white":      RGBColor(0xFF, 0xFF, 0xFF),
        "gray":       RGBColor(0x88, 0x88, 0x88),
        "text":       RGBColor(0x2A, 0x2A, 0x2A),
        "light":      RGBColor(0xF2, 0xF2, 0xF2),
        "bg_hex":     "1A1A2E",
        "accent_hex": "D4AF37",
        "light_hex":  "F5F0E6",
        "cover_logo": "mctv_logo_on_dark.png",
    },
    "pastel": {
        "label": "Peaceful Pastels",
        "primary":    RGBColor(0x5B, 0x7B, 0x7A),  # Sage green
        "accent":     RGBColor(0xC4, 0x8D, 0x78),  # Dusty rose
        "white":      RGBColor(0xFF, 0xFF, 0xFF),
        "gray":       RGBColor(0x99, 0x99, 0x99),
        "text":       RGBColor(0x3D, 0x3D, 0x3D),
        "light":      RGBColor(0xF7, 0xF5, 0xF3),
        "bg_hex":     "5B7B7A",
        "accent_hex": "C48D78",
        "light_hex":  "F3EEED",
        "cover_logo": "mctv_logo_on_pastel.png",
    },
}

# Legacy module-level aliases (for any code still referencing these directly)
NAVY = COLOR_SCHEMES["original"]["primary"]
GOLD = COLOR_SCHEMES["original"]["accent"]
WHITE = COLOR_SCHEMES["original"]["white"]
GRAY = COLOR_SCHEMES["original"]["gray"]
LIGHT_GRAY = COLOR_SCHEMES["original"]["light"]
DARK_TEXT = COLOR_SCHEMES["original"]["text"]

OUTPUT_DIR = Path(__file__).parent.parent / "output"
PROJECT_ROOT = Path(__file__).parent.parent


class DocxService:
    """Creates branded MCTV Word documents."""

    def __init__(self, config: dict, color_scheme: str = "original"):
        self.config = config
        self.company = config["company"]
        self.scheme_name = color_scheme
        self.c = COLOR_SCHEMES.get(color_scheme, COLOR_SCHEMES["original"])

    def create_document(self) -> Document:
        """Create a new document with MCTV default styles."""
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = self.c["text"]

        # Set narrow margins for maximum content density
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

        return doc

    def add_cover_page(self, doc: Document, title: str, subtitle: str,
                       prepared_for: str, prepared_by: dict, date: str = None,
                       client_logo_path: str = None):
        """Add a branded cover page with navy background, gold/white text.

        Uses a full-page table cell with navy fill. The cell height is forced
        to fill the entire printable area so there is no white gap.
        The cover page uses its own section with minimal margins so the navy
        fills edge-to-edge; a section break restores normal margins for
        subsequent pages.
        """
        if date is None:
            date = datetime.now().strftime("%B %Y")

        # Shrink cover-page margins so navy fills nearly edge-to-edge
        cover_section = doc.sections[0]
        cover_section.top_margin = Cm(0.8)
        cover_section.bottom_margin = Cm(0.5)
        cover_section.left_margin = Cm(1.3)
        cover_section.right_margin = Cm(1.3)
        # Hide footer on cover page
        cover_section.footer.is_linked_to_previous = False
        cover_section.different_first_page_header_footer = False

        # Full-page background table (single cell with navy fill)
        bg_table = doc.add_table(rows=1, cols=1)
        bg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = bg_table.rows[0].cells[0]

        # Set navy background
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:fill"): self.c["bg_hex"],
            qn("w:val"): "clear",
        })
        tc_pr.append(shd)

        # Force cell to fill full page height (Letter: 11in = 15840 twips)
        # With tiny cover margins (0.8+0.5 cm = ~740 twips), almost all page is usable
        cell_height_twips = 14800  # ~10.3 inches
        tr_pr = cell._element.getparent().get_or_add_trPr()
        tr_height = tr_pr.makeelement(qn("w:trHeight"), {
            qn("w:val"): str(cell_height_twips),
            qn("w:hRule"): "atLeast",
        })
        tr_pr.append(tr_height)

        # Remove default cell margins for edge-to-edge feel
        cell_margin = tc_pr.makeelement(qn("w:tcMar"), {})
        for side in ("top", "left", "bottom", "right"):
            margin = cell_margin.makeelement(qn(f"w:{side}"), {
                qn("w:w"): "0",
                qn("w:type"): "dxa",
            })
            cell_margin.append(margin)
        tc_pr.append(cell_margin)

        # Vertically center content in cell
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self._remove_table_borders(bg_table)

        # ── All cover content goes inside this cell ──

        # MCTV logo — use the scheme-specific pre-composited version
        cover_logo_name = self.c.get("cover_logo", "mctv_logo_on_navy.png")
        mctv_logo_navy = PROJECT_ROOT / "assets" / "branding" / cover_logo_name
        if mctv_logo_navy.exists():
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(0)
            p.space_after = Pt(16)
            run = p.add_run()
            run.add_picture(str(mctv_logo_navy), width=Inches(3.0))
        else:
            # Fallback text logo
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_after = Pt(16)
            run = p.add_run("MCTV ELITE ADVERTISING")
            run.font.size = Pt(16)
            run.font.color.rgb = self.c["accent"]
            run.font.bold = True
            run.font.name = "Calibri"

        # "Prepared for" label
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(2)
        run = p.add_run("Prepared for")
        run.font.size = Pt(11)
        run.font.color.rgb = self.c["white"]
        run.font.italic = True
        run.font.name = "Calibri"

        # Client name (big, white, bold)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(2)
        run = p.add_run(prepared_for.upper())
        run.font.size = Pt(24)
        run.font.color.rgb = self.c["white"]
        run.font.bold = True
        run.font.name = "Calibri"

        # Subtitle (business name)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(4)
        run = p.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = self.c["accent"]
        run.font.name = "Calibri"

        # Client logo (if provided)
        if client_logo_path:
            logo_path = Path(client_logo_path)
            if logo_path.exists():
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(8)
                p.space_after = Pt(8)
                run = p.add_run()
                run.add_picture(str(logo_path), width=Inches(1.8))

        # Gold accent line
        accent = cell.add_paragraph()
        accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
        accent.space_before = Pt(16)
        accent.space_after = Pt(16)
        run = accent.add_run("\u2500" * 30)
        run.font.size = Pt(10)
        run.font.color.rgb = self.c["accent"]

        # Title (ADVERTISING PARTNERSHIP PROPOSAL)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(0)
        p.space_after = Pt(4)
        run = p.add_run(title)
        run.font.size = Pt(28)
        run.font.color.rgb = self.c["accent"]
        run.font.bold = True
        run.font.name = "Calibri"

        # Gold accent line
        accent = cell.add_paragraph()
        accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
        accent.space_before = Pt(16)
        accent.space_after = Pt(16)
        run = accent.add_run("\u2500" * 30)
        run.font.size = Pt(10)
        run.font.color.rgb = self.c["accent"]

        # Date
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(16)
        run = p.add_run(date)
        run.font.size = Pt(12)
        run.font.color.rgb = self.c["white"]
        run.font.name = "Calibri"

        # Prepared by (rep info)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(2)
        run = p.add_run(f"{prepared_by['name']}  |  MCTV Elite Advertising")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = self.c["accent"]
        run.font.name = "Calibri"

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(2)
        run = p.add_run(f"{prepared_by['email']}  |  {prepared_by['phone']}")
        run.font.size = Pt(9)
        run.font.color.rgb = self.c["white"]
        run.font.name = "Calibri"

        # Website
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(0)
        run = p.add_run("www.mctvofms.com")
        run.font.size = Pt(9)
        run.font.color.rgb = self.c["accent"]
        run.font.name = "Calibri"

        # Add a section break (new page) to restore normal margins for
        # subsequent content pages.  This avoids the blank-page problem
        # that doc.add_page_break() caused, while also letting us use
        # different margins on the cover vs. content pages.
        new_section = doc.add_section()      # defaults to NEW_PAGE
        new_section.top_margin = Cm(1.5)
        new_section.bottom_margin = Cm(1.5)
        new_section.left_margin = Cm(2.0)
        new_section.right_margin = Cm(2.0)

    def add_section_header(self, doc: Document, text: str):
        """Add a full-width section header bar with navy background, white text,
        and a thin gold accent line underneath — matching the MUC gold standard."""

        # ── Navy background bar (full-width single-cell table) ──
        header_table = doc.add_table(rows=1, cols=1)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = header_table.rows[0].cells[0]

        # Navy background fill
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:fill"): self.c["bg_hex"],
            qn("w:val"): "clear",
        })
        tc_pr.append(shd)

        # White bold text, centered
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(8)
        run = p.add_run(text.upper())
        run.font.size = Pt(16)
        run.font.color.rgb = self.c["white"]
        run.font.bold = True
        run.font.name = "Calibri"

        self._remove_table_borders(header_table)

        # ── Gold accent underline bar ──
        accent_bar = doc.add_table(rows=1, cols=1)
        accent_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
        accent_cell = accent_bar.rows[0].cells[0]
        accent_cell.height = Cm(0.08)

        # Gold background fill
        tc_pr2 = accent_cell._element.get_or_add_tcPr()
        shd2 = tc_pr2.makeelement(qn("w:shd"), {
            qn("w:fill"): self.c["accent_hex"],
            qn("w:val"): "clear",
        })
        tc_pr2.append(shd2)

        # Minimal content to hold the bar
        p2 = accent_cell.paragraphs[0]
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        run2 = p2.add_run()
        run2.font.size = Pt(1)

        self._remove_table_borders(accent_bar)

    def add_sub_header(self, doc: Document, text: str):
        """Add a bold sub-header with gold left accent."""
        p = doc.add_paragraph()
        p.space_before = Pt(6)
        p.space_after = Pt(1)
        run = p.add_run(f"\u275A  {text}")
        run.font.size = Pt(12)
        run.font.color.rgb = self.c["primary"]
        run.font.bold = True

    def add_section_divider(self, doc: Document):
        """Add a thin gold horizontal rule to visually separate sections."""
        p = doc.add_paragraph()
        p.space_before = Pt(6)
        p.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\u2500" * 50)
        run.font.size = Pt(6)
        run.font.color.rgb = self.c["accent"]

    def add_callout_box(self, doc: Document, text: str, bg_color: str = None):
        """Add a colored callout/highlight box with accent left border."""
        if bg_color is None:
            bg_color = self.c["light_hex"]
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Clear default table-level borders so cell borders render properly
        self._clear_table_borders(table)
        cell = table.rows[0].cells[0]

        # Set background color
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:fill"): bg_color,
            qn("w:val"): "clear",
        })
        tc_pr.append(shd)

        # Generous cell padding
        cell_margin = tc_pr.makeelement(qn("w:tcMar"), {})
        for side, val in (("top", "160"), ("bottom", "160"),
                          ("left", "200"), ("right", "140")):
            m = cell_margin.makeelement(qn(f"w:{side}"), {
                qn("w:w"): val, qn("w:type"): "dxa",
            })
            cell_margin.append(m)
        tc_pr.append(cell_margin)

        # Gold left border accent (~3pt) + thin gray on other sides
        self._set_cell_borders(cell, left_color=self.c["accent_hex"], left_sz=24,
                               other_color="D0D0D0", other_sz=4)

        # Add padding via paragraph formatting
        p = cell.paragraphs[0]
        p.space_before = Pt(4)
        p.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = self.c["text"]
        run.font.name = "Calibri"

    def add_selling_point(self, doc: Document, title: str, body: str):
        """Add a selling point with colored square bullet, bold title, and body text.

        Used for The Opportunity section to make selling points scannable
        rather than dumping them into a single callout box.
        """
        # Title line with colored square bullet
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.4)
        # Square bullet in accent color
        bullet_run = p.add_run("\u25A0  ")
        bullet_run.font.size = Pt(9)
        bullet_run.font.color.rgb = self.c["accent"]
        # Bold title in primary color
        title_run = p.add_run(title)
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.color.rgb = self.c["primary"]
        title_run.font.name = "Calibri"

        # Body text (if provided)
        if body:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(4)
            p2.paragraph_format.left_indent = Cm(0.9)
            run2 = p2.add_run(body)
            run2.font.size = Pt(10)
            run2.font.color.rgb = self.c["text"]
            run2.font.name = "Calibri"

    def add_accent_card(self, doc: Document, title: str, body: str):
        """Add a styled card with thick accent left border, light background,
        bold title, and body text — used for What's Included and Why MCTV sections."""

        # Spacer paragraph between cards for breathing room
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        spacer_run = spacer.add_run()
        spacer_run.font.size = Pt(4)

        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Clear default table-level borders so cell borders render properly
        self._clear_table_borders(table)
        cell = table.rows[0].cells[0]

        # Light background fill
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:fill"): self.c["light_hex"],
            qn("w:val"): "clear",
        })
        tc_pr.append(shd)

        # Cell padding — generous internal margins so content doesn't feel cramped
        # Left padding is larger to give breathing room from the thick accent border
        cell_margin = tc_pr.makeelement(qn("w:tcMar"), {})
        for side, val in (("top", "180"), ("bottom", "180"),
                          ("left", "240"), ("right", "140")):
            m = cell_margin.makeelement(qn(f"w:{side}"), {
                qn("w:w"): val, qn("w:type"): "dxa",
            })
            cell_margin.append(m)
        tc_pr.append(cell_margin)

        # Thick accent left border (~4.5pt = sz 36) + thin gray on other sides
        self._set_cell_borders(cell, left_color=self.c["accent_hex"], left_sz=36,
                               other_color="D0D0D0", other_sz=2)

        # Bold title in primary color
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = self.c["primary"]
        run.font.name = "Calibri"

        # Body text in text color
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        run2 = p2.add_run(body)
        run2.font.size = Pt(10)
        run2.font.color.rgb = self.c["text"]
        run2.font.name = "Calibri"

    def add_body_text(self, doc: Document, text: str):
        """Add body paragraphs with proper spacing.

        Automatically detects numbered items (e.g., '1. Bold Title\\nBody text')
        and formats them with a bold navy title and normal body text.
        """
        import re
        # Pre-process: ensure each numbered line (e.g. "2. Step") gets its own
        # paragraph block.  Claude sometimes returns steps separated by single
        # newlines, which groups them all into one block and only the first
        # match succeeds.  This inserts a double-newline before each numbered
        # line that isn't already preceded by one.
        text = re.sub(r'(?<!\n)\n(\d+\.\s)', r'\n\n\1', text.strip())
        paragraphs = text.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Detect numbered item: "1. Title\nBody..." or "1. Title: Body..."
            numbered = re.match(r"^(\d+)\.\s+(.+?)(?:\n(.+))?$", para_text, re.DOTALL)
            if numbered:
                num = numbered.group(1)
                first_line = numbered.group(2).strip()
                rest = (numbered.group(3) or "").strip()

                # Check if the first line has a title and inline body (split by newline)
                p = doc.add_paragraph()
                p.space_before = Pt(4)
                p.space_after = Pt(2)

                # Bold numbered title
                run = p.add_run(f"{num}. {first_line}")
                run.font.size = Pt(10.5)
                run.font.bold = True
                run.font.color.rgb = self.c["primary"]
                run.font.name = "Calibri"

                # Body text below the title
                if rest:
                    p = doc.add_paragraph()
                    p.space_after = Pt(4)
                    run = p.add_run(rest)
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = self.c["text"]
                    run.font.name = "Calibri"
            else:
                p = doc.add_paragraph()
                p.space_after = Pt(4)
                run = p.add_run(para_text)
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.c["text"]
                run.font.name = "Calibri"

    def add_bullet_point(self, doc: Document, title: str, description: str):
        """Add a bullet point with bold title and description."""
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        # Left indent with hanging indent for bullet alignment
        pf = p.paragraph_format
        pf.left_indent = Cm(0.6)
        pf.first_line_indent = Cm(-0.6)

        # Gold bullet character
        run = p.add_run("\u25CF  ")
        run.font.size = Pt(7)
        run.font.color.rgb = self.c["accent"]

        run = p.add_run(f"{title}: ")
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = self.c["primary"]

        run = p.add_run(description)
        run.font.size = Pt(10.5)
        run.font.color.rgb = self.c["text"]

    def add_bullet_list(self, doc: Document, text: str):
        """Parse dash-prefixed bullet items from Claude and render with bold navy titles.

        Expected format: '- Bold Title: Description sentence.'
        Falls back to add_body_text for non-matching lines.
        """
        import re
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Match "- Title: Description" or "- Title -- Description"
            bullet_match = re.match(r'^-\s+(.+?)(?::|--)\s+(.+)$', line)
            if bullet_match:
                title = bullet_match.group(1).strip()
                desc = bullet_match.group(2).strip()
                self.add_bullet_point(doc, title, desc)
            else:
                # Fallback: treat as regular body text
                self.add_body_text(doc, line)

    def add_inline_photos(self, doc: Document, photo_paths: list,
                          max_width: float = 2.0, cols: int = 2):
        """Add 1-4 photos inline within a section (no title, compact).

        Responsive layouts based on photo count:
          1 photo:  Single centered image, ~70% page width (3.0in)
          2 photos: Side-by-side, each ~48% page width (2.8in)
          3 photos: Top row 2 side-by-side + bottom row 1 centered
          4 photos: 2×2 grid

        Photos should complement the text, not dominate the page.
        """
        photo_paths = [p for p in (photo_paths or []) if Path(p).exists()]
        if not photo_paths:
            return

        count = len(photo_paths)

        # Single photo: centered, generous size
        if count == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            try:
                run = p.add_run()
                run.add_picture(photo_paths[0], width=Inches(3.0))
            except Exception:
                pass
            return

        # 2 photos: side-by-side
        if count == 2:
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i in range(2):
                cell = table.rows[0].cells[i]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(2)
                p.space_after = Pt(2)
                try:
                    run = p.add_run()
                    run.add_picture(photo_paths[i], width=Inches(2.8))
                except Exception:
                    run = p.add_run("[Image]")
                    run.font.size = Pt(9)
                    run.font.color.rgb = self.c["gray"]
            self._remove_table_borders(table)
            return

        # 3 photos: top row 2 side-by-side + bottom row 1 centered
        if count == 3:
            # Top row: 2 side-by-side
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i in range(2):
                cell = table.rows[0].cells[i]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(2)
                p.space_after = Pt(2)
                try:
                    run = p.add_run()
                    run.add_picture(photo_paths[i], width=Inches(2.8))
                except Exception:
                    run = p.add_run("[Image]")
                    run.font.size = Pt(9)
                    run.font.color.rgb = self.c["gray"]
            # Bottom row: merge cells and center one photo
            bottom_cell = table.rows[1].cells[0]
            bottom_cell.merge(table.rows[1].cells[1])
            p = bottom_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(2)
            p.space_after = Pt(2)
            try:
                run = p.add_run()
                run.add_picture(photo_paths[2], width=Inches(3.0))
            except Exception:
                run = p.add_run("[Image]")
                run.font.size = Pt(9)
                run.font.color.rgb = self.c["gray"]
            self._remove_table_borders(table)
            return

        # 4+ photos: 2×2 grid (cap at 4)
        photos = photo_paths[:4]
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, photo_path in enumerate(photos):
            row_idx = i // 2
            col_idx = i % 2
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(2)
            p.space_after = Pt(2)
            try:
                run = p.add_run()
                run.add_picture(photo_path, width=Inches(2.8))
            except Exception:
                run = p.add_run("[Image]")
                run.font.size = Pt(9)
                run.font.color.rgb = self.c["gray"]
        self._remove_table_borders(table)

    def add_photos_grid(self, doc: Document, photo_paths: list, title: str = None,
                        max_width: float = 2.5, cols: int = 2,
                        captions: list = None):
        """Add a grid of photos to the document.

        Args:
            photo_paths: List of file paths to images.
            title: Optional sub-header above the photos.
            max_width: Max width per image in inches.
            cols: Number of columns in the grid.
            captions: Optional list of caption strings (same length as photo_paths).
        """
        # Filter to only paths that actually exist on disk
        valid = []
        valid_captions = []
        for i, p in enumerate(photo_paths or []):
            if Path(p).exists():
                valid.append(p)
                if captions and i < len(captions):
                    valid_captions.append(captions[i])
                else:
                    valid_captions.append(None)
        if not valid:
            return

        if title:
            self.add_sub_header(doc, title)
            # Keep the title paragraph together with the grid below
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.keep_with_next = True

        # Build table grid for photos
        rows_needed = (len(valid) + cols - 1) // cols
        table = doc.add_table(rows=rows_needed, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, photo_path in enumerate(valid):
            row_idx = i // cols
            col_idx = i % cols
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Compact spacing to prevent blank page overflow
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            try:
                run = p.add_run()
                run.add_picture(photo_path, width=Inches(max_width))
            except Exception:
                run = p.add_run("[Image could not be loaded]")
                run.font.size = Pt(9)
                run.font.color.rgb = self.c["gray"]

            # Optional caption below the photo
            caption_text = valid_captions[i] if i < len(valid_captions) else None
            if caption_text:
                cap_p = cell.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_before = Pt(1)
                cap_p.paragraph_format.space_after = Pt(2)
                cap_run = cap_p.add_run(caption_text)
                cap_run.font.size = Pt(8)
                cap_run.font.italic = True
                cap_run.font.color.rgb = self.c["gray"]
                cap_run.font.name = "Calibri"

        self._remove_table_borders(table)

    def add_metrics_banner(self, doc: Document, metrics: dict):
        """Add a row of big metrics with navy background (e.g., 125+ Screens)."""
        table = doc.add_table(rows=2, cols=len(metrics))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (value, label) in enumerate(metrics.items()):
            # Value row — navy background, gold text
            cell = table.rows[0].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._element.get_or_add_tcPr()
            shd = tc_pr.makeelement(qn("w:shd"), {
                qn("w:fill"): self.c["bg_hex"],
                qn("w:val"): "clear",
            })
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(6)
            p.space_after = Pt(2)
            # Auto-scale font for long values (e.g., venue names in KPI row 2)
            value_str = str(value)
            if len(value_str) > 15:
                font_size = Pt(14)
            elif len(value_str) > 10:
                font_size = Pt(16)
            else:
                font_size = Pt(20)
            run = p.add_run(value_str)
            run.font.size = font_size
            run.font.bold = True
            run.font.color.rgb = self.c["accent"]

            # Label row — navy background, white text
            cell = table.rows[1].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._element.get_or_add_tcPr()
            shd = tc_pr.makeelement(qn("w:shd"), {
                qn("w:fill"): self.c["bg_hex"],
                qn("w:val"): "clear",
            })
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(0)
            p.space_after = Pt(6)
            run = p.add_run(label)
            run.font.size = Pt(8)
            run.font.color.rgb = self.c["white"]

        # Gold top border on the banner, remove other borders
        self._remove_table_borders(table)
        # Add gold top border to the first row of cells
        for cell in table.rows[0].cells:
            tc_pr = cell._element.get_or_add_tcPr()
            borders = tc_pr.makeelement(qn("w:tcBorders"), {})
            border = borders.makeelement(qn("w:top"), {
                qn("w:val"): "single", qn("w:sz"): "12",
                qn("w:space"): "0", qn("w:color"): self.c["accent_hex"],
            })
            borders.append(border)
            tc_pr.append(borders)

    def add_pricing_table(self, doc: Document, tiers: list,
                          recommended_idx: int = None):
        """Add a formatted pricing comparison table.

        Args:
            tiers: List of tier dicts from config.
            recommended_idx: 0-based index of the recommended tier row
                             (gets a gold highlight).
        """
        table = doc.add_table(rows=1 + len(tiers), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ["Monthly Rate", "Screens", "Ad Plays/Mo", "Cost/Screen"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(header)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = self.c["white"]
            # Navy background
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:fill"): self.c["bg_hex"],
                qn("w:val"): "clear",
            })
            shading.append(shading_elm)

        # Data rows
        for row_idx, tier in enumerate(tiers):
            row = table.rows[row_idx + 1]
            rate = tier['monthly_rate']
            rate_str = f"${rate:,.0f}" if rate == int(rate) else f"${rate:,.2f}"
            cps = tier.get('cost_per_screen', 0)
            cps_str = f"${cps:,.0f}" if cps == int(cps) else f"${cps:,.2f}"
            values = [
                f"{rate_str}/mo",
                f"{tier['screens']} Screens",
                tier.get("plays_per_month", ""),
                cps_str,
            ]
            is_recommended = (recommended_idx is not None and row_idx == recommended_idx)
            for col_idx, value in enumerate(values):
                cell = row.cells[col_idx]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(str(value))
                run.font.size = Pt(11)
                if col_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = self.c["accent"]
                    run.font.size = Pt(14)

                # Recommended row gets light gold highlight; others alternate
                if is_recommended:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn("w:shd"), {
                        qn("w:fill"): "FFF8E7",  # light gold tint
                        qn("w:val"): "clear",
                    })
                    shading.append(shading_elm)
                elif row_idx % 2 == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn("w:shd"), {
                        qn("w:fill"): "F5F5F5",  # alt-row, universal
                        qn("w:val"): "clear",
                    })
                    shading.append(shading_elm)

            # Add "RECOMMENDED" label on the highlighted tier
            if is_recommended:
                # Add a small label in the first cell
                rec_p = row.cells[0].add_paragraph()
                rec_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rec_p.paragraph_format.space_before = Pt(0)
                rec_p.paragraph_format.space_after = Pt(2)
                rec_run = rec_p.add_run("\u2605 RECOMMENDED")
                rec_run.font.size = Pt(7)
                rec_run.font.bold = True
                rec_run.font.color.rgb = self.c["accent"]
                rec_run.font.name = "Calibri"

        # Add thin gray borders for structure
        self._set_table_borders(table, color="D0D0D0", sz=4)

    def add_contract_terms(self, doc: Document, config: dict):
        """Add partnership terms section with 6-month and 12-month boxes."""
        terms = config["pricing"]["contract_terms"]
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, months in enumerate(terms["standard_months"]):
            cell = table.rows[0].cells[i]
            bonus = terms[f"prepay_{months}mo_bonus"]

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{months}-MONTH AGREEMENT")
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = self.c["accent"]

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label = "Minimum commitment" if months == 6 else "Best value \u2014 lock in your rate"
            run = p.add_run(label)
            run.font.size = Pt(9)
            run.font.color.rgb = self.c["gray"]

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"PREPAY BONUS: {bonus}")
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = self.c["primary"]

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("1 extra month free when you pay upfront")
            run.font.size = Pt(9)
            run.font.color.rgb = self.c["gray"]

        # Gold top border, thin gray sides for a polished box look
        for cell in table.rows[0].cells:
            self._set_cell_borders(cell, left_color=self.c["accent_hex"], left_sz=12,
                                   other_color="D0D0D0", other_sz=4)

    def add_data_table(self, doc: Document, headers: list, rows: list,
                       bold_rows: int = 0, totals_row: list = None):
        """Add a data table for traction reports.

        Args:
            headers: Column header strings.
            rows: List of row data (each row is a list of strings).
            bold_rows: Bold the first N data rows (top performers).
            totals_row: Optional summary row appended at the bottom.
        """
        total_rows = 1 + len(rows) + (1 if totals_row else 0)
        table = doc.add_table(rows=total_rows, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = self.c["white"]
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:fill"): self.c["bg_hex"],
                qn("w:val"): "clear",
            })
            shading.append(shading_elm)

        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(row_data):
                cell = row.cells[col_idx]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(value))
                run.font.size = Pt(9)
                run.font.color.rgb = self.c["text"]
                # Bold top N performers
                if row_idx < bold_rows:
                    run.font.bold = True

                if row_idx % 2 == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn("w:shd"), {
                        qn("w:fill"): "F5F5F5",  # alt-row, universal
                        qn("w:val"): "clear",
                    })
                    shading.append(shading_elm)

        # Totals row (styled like header but slightly lighter)
        if totals_row:
            totals_idx = 1 + len(rows)
            row = table.rows[totals_idx]
            for col_idx, value in enumerate(totals_row):
                cell = row.cells[col_idx]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(value))
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = self.c["white"]
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): self.c["bg_hex"],
                    qn("w:val"): "clear",
                })
                shading.append(shading_elm)

    def add_team_section(self, doc: Document,
                         closing_text: str = "We look forward to partnering with you.",
                         dark_mode: bool = False):
        """Add the Meet Your Team section with photos, closing statement, and logo.

        Args:
            dark_mode: If True, use dark navy background with white/gold text
                       (used for traction reports for a premium closing page).
        """
        self.add_section_header(doc, "Meet Your Team")

        team = list(self.config["team"])

        # Reorder so the preparer (sales rep) appears first
        preparer = getattr(self, 'preparer_name', None)
        if preparer:
            team.sort(key=lambda m: (0 if m["name"] == preparer else 1))

        table = doc.add_table(rows=1, cols=len(team))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Color selections based on mode
        name_color = self.c["white"] if dark_mode else self.c["primary"]
        title_color = self.c["accent"] if True else self.c["accent"]  # Gold in both modes
        contact_color = RGBColor(0xCC, 0xCC, 0xCC) if dark_mode else self.c["gray"]

        for i, member in enumerate(team):
            cell = table.rows[0].cells[i]

            # Dark navy cell background
            if dark_mode:
                tc_pr = cell._element.get_or_add_tcPr()
                shd = tc_pr.makeelement(qn("w:shd"), {
                    qn("w:fill"): self.c["bg_hex"],
                    qn("w:val"): "clear",
                })
                tc_pr.append(shd)

            # Team photo
            photo_path = member.get("photo", "")
            if photo_path:
                full_path = PROJECT_ROOT / photo_path
                if full_path.exists():
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(full_path), width=Inches(1.5))
                    # Name goes in next paragraph
                    p = cell.add_paragraph()
                else:
                    p = cell.paragraphs[0]
            else:
                p = cell.paragraphs[0]

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(member["name"])
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = name_color

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(member["title"])
            run.font.size = Pt(10)
            run.font.color.rgb = title_color
            run.font.italic = True

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(member["phone"])
            run.font.size = Pt(10)
            run.font.color.rgb = contact_color

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(member["email"])
            run.font.size = Pt(10)
            run.font.color.rgb = contact_color

        self._remove_table_borders(table)

        # ── Closing statement + Logo + URL ──
        # In dark mode: wrap in a navy-filled table to create a seamless dark page
        if dark_mode:
            # Build closing + logo + URL inside a single-cell navy table
            # so the dark background extends through the dead space below
            closing_table = doc.add_table(rows=1, cols=1)
            closing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._clear_table_borders(closing_table)
            closing_cell = closing_table.rows[0].cells[0]

            # Navy fill on the cell
            tc_pr = closing_cell._element.get_or_add_tcPr()
            shd = tc_pr.makeelement(qn("w:shd"), {
                qn("w:fill"): self.c["bg_hex"],
                qn("w:val"): "clear",
            })
            tc_pr.append(shd)

            # Generous padding to push content down and fill the page
            cell_margin = tc_pr.makeelement(qn("w:tcMar"), {})
            for side, val in (("top", "400"), ("bottom", "2000"),
                              ("left", "200"), ("right", "200")):
                m = cell_margin.makeelement(qn(f"w:{side}"), {
                    qn("w:w"): val, qn("w:type"): "dxa",
                })
                cell_margin.append(m)
            tc_pr.append(cell_margin)

            # Closing text
            if closing_text:
                p = closing_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(8)
                p.space_after = Pt(16)
                run = p.add_run(closing_text)
                run.font.size = Pt(12)
                run.font.italic = True
                run.font.color.rgb = self.c["accent"]
                run.font.name = "Calibri"

            # Logo — use white-on-transparent for dark mode (seamless on navy bg),
            # fall back to scheme cover logo, then generic logo as last resort.
            mctv_logo = PROJECT_ROOT / "assets" / "branding" / "mctv_logo_white.png"
            if not mctv_logo.exists():
                logo_name = self.c.get("cover_logo", "mctv_logo_on_navy.png")
                mctv_logo = PROJECT_ROOT / "assets" / "branding" / logo_name
            if not mctv_logo.exists():
                mctv_logo = PROJECT_ROOT / "assets" / "branding" / "mctv_logo.png"
            if mctv_logo.exists():
                p = closing_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(4)
                p.space_after = Pt(4)
                run = p.add_run()
                run.add_picture(str(mctv_logo), width=Inches(1.8))

            # Website URL
            p = closing_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(2)
            p.space_after = Pt(0)
            run = p.add_run("www.mctvofms.com")
            run.font.size = Pt(11)
            run.font.color.rgb = self.c["accent"]
            run.font.name = "Calibri"

        else:
            # Light mode: regular paragraphs (proposal closing)
            if closing_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(12)
                p.space_after = Pt(4)
                run = p.add_run(closing_text)
                run.font.size = Pt(12)
                run.font.italic = True
                run.font.color.rgb = self.c["accent"]
                run.font.name = "Calibri"

            # MCTV logo + website
            logo_name = "mctv_logo.png"
            mctv_logo = PROJECT_ROOT / "assets" / "branding" / logo_name
            if mctv_logo.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(2)
                p.space_after = Pt(0)
                run = p.add_run()
                run.add_picture(str(mctv_logo), width=Inches(1.6))

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(0)
            p.space_after = Pt(0)
            run = p.add_run("www.mctvofms.com")
            run.font.size = Pt(10)
            run.font.color.rgb = self.c["primary"]
        run.font.name = "Calibri"

    def add_venue_categories(self, doc: Document):
        """Add the 'Where Your Ads Play' venue category grid — ultra-compact 2-row format."""
        categories = [
            "Restaurants & Bars", "Barbershops & Salons",
            "Medical & Dental", "Gyms & Fitness",
            "Auto & Service Shops", "Retail & Boutiques",
            "Professional Offices", "Community Venues",
        ]

        # 2 rows x 4 columns for maximum compactness
        table = doc.add_table(rows=2, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, name in enumerate(categories):
            row_idx = i // 4
            col_idx = i % 4
            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            run = p.add_run(name)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = self.c["primary"]

        self._remove_table_borders(table)

    def add_footer(self, doc: Document, footer_text: str = "Confidential Partnership Proposal"):
        """Add branded footer: 'MCTV Elite Advertising  |  [text]  |  Page X'

        Skips the first section (cover page) so the navy background isn't
        interrupted by a page-number footer.

        Args:
            footer_text: Middle text (default: "Confidential Partnership Proposal").
                         Reports can pass different text.
        """
        for idx, section in enumerate(doc.sections):
            if idx == 0:
                # Cover page — no footer
                continue
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # "MCTV Elite Advertising" in accent color
            run_brand = p.add_run("MCTV Elite Advertising")
            run_brand.font.size = Pt(8)
            run_brand.font.color.rgb = self.c["accent"]
            run_brand.font.name = "Calibri"

            # Separator
            sep1 = p.add_run("   |   ")
            sep1.font.size = Pt(8)
            sep1.font.color.rgb = self.c["gray"]
            sep1.font.name = "Calibri"

            # Footer text in gray
            run_text = p.add_run(footer_text)
            run_text.font.size = Pt(8)
            run_text.font.color.rgb = self.c["gray"]
            run_text.font.name = "Calibri"

            # Separator
            sep2 = p.add_run("   |   ")
            sep2.font.size = Pt(8)
            sep2.font.color.rgb = self.c["gray"]
            sep2.font.name = "Calibri"

            # "Page " label in gray
            run_label = p.add_run("Page ")
            run_label.font.size = Pt(8)
            run_label.font.color.rgb = self.c["gray"]
            run_label.font.name = "Calibri"

            # PAGE field code
            run_fld1 = p.add_run()
            fld_begin = run_fld1._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
            run_fld1._element.append(fld_begin)
            run_fld2 = p.add_run()
            instr = run_fld2._element.makeelement(qn("w:instrText"), {})
            instr.text = " PAGE "
            run_fld2._element.append(instr)
            run_fld2.font.size = Pt(8)
            run_fld2.font.color.rgb = self.c["gray"]
            run_fld2.font.name = "Calibri"
            run_fld3 = p.add_run()
            fld_end = run_fld3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
            run_fld3._element.append(fld_end)

    def save_proposal(self, doc: Document, filename: str, also_pdf: bool = True) -> Path:
        """Save a proposal document and optionally convert to PDF. Returns docx path."""
        output_path = OUTPUT_DIR / "proposals" / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        if also_pdf:
            self._convert_to_pdf(output_path)
        return output_path

    def save_report(self, doc: Document, filename: str, also_pdf: bool = True) -> Path:
        """Save a report document and optionally convert to PDF. Returns docx path."""
        output_path = OUTPUT_DIR / "reports" / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        if also_pdf:
            self._convert_to_pdf(output_path)
        return output_path

    def save_email(self, content: str, filename: str) -> Path:
        """Save a cover email and return the path."""
        output_path = OUTPUT_DIR / "emails" / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")
        return output_path

    def _clear_table_borders(self, table):
        """Clear table-level borders AND style so cell-level borders render.

        python-docx applies a "Table Grid" style by default which defines
        borders that can override cell-level tcBorders.  We remove the style
        reference entirely and set explicit tblBorders to 'none' so the
        cell-level accent borders are the only borders that render.
        """
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = tbl.makeelement(qn("w:tblPr"), {})
            tbl.insert(0, tbl_pr)
        # Remove the "Table Grid" (or any) style — it defines its own borders
        for style_el in tbl_pr.findall(qn("w:tblStyle")):
            tbl_pr.remove(style_el)
        # Remove any existing tblBorders
        for existing in tbl_pr.findall(qn("w:tblBorders")):
            tbl_pr.remove(existing)
        tbl_borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = tbl_borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none", qn("w:sz"): "0",
                qn("w:space"): "0", qn("w:color"): "auto",
            })
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)

    def _remove_table_borders(self, table):
        """Remove all borders from a table (for layout tables)."""
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                borders = tc_pr.makeelement(qn("w:tcBorders"), {})
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = borders.makeelement(qn(f"w:{edge}"), {
                        qn("w:val"): "none",
                        qn("w:sz"): "0",
                        qn("w:space"): "0",
                        qn("w:color"): "auto",
                    })
                    borders.append(border)
                tc_pr.append(borders)

    def _set_cell_borders(self, cell, left_color=None, left_sz=12,
                          other_color=None, other_sz=4):
        """Set borders on a single cell. Colors are hex strings (no #)."""
        tc_pr = cell._element.get_or_add_tcPr()
        # Remove any existing tcBorders to prevent duplicates / conflicts
        for existing in tc_pr.findall(qn("w:tcBorders")):
            tc_pr.remove(existing)
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "bottom", "right"):
            if other_color:
                border = borders.makeelement(qn(f"w:{edge}"), {
                    qn("w:val"): "single",
                    qn("w:sz"): str(other_sz),
                    qn("w:space"): "0",
                    qn("w:color"): other_color,
                })
            else:
                border = borders.makeelement(qn(f"w:{edge}"), {
                    qn("w:val"): "none", qn("w:sz"): "0",
                    qn("w:space"): "0", qn("w:color"): "auto",
                })
            borders.append(border)
        # Left border — accent
        if left_color:
            border = borders.makeelement(qn("w:left"), {
                qn("w:val"): "single",
                qn("w:sz"): str(left_sz),
                qn("w:space"): "0",
                qn("w:color"): left_color,
            })
        else:
            border = borders.makeelement(qn("w:left"), {
                qn("w:val"): "none", qn("w:sz"): "0",
                qn("w:space"): "0", qn("w:color"): "auto",
            })
        borders.append(border)
        tc_pr.append(borders)

    def _set_table_borders(self, table, color="D0D0D0", sz=4):
        """Set uniform thin borders on all cells in a table."""
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                borders = tc_pr.makeelement(qn("w:tcBorders"), {})
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = borders.makeelement(qn(f"w:{edge}"), {
                        qn("w:val"): "single",
                        qn("w:sz"): str(sz),
                        qn("w:space"): "0",
                        qn("w:color"): color,
                    })
                    borders.append(border)
                tc_pr.append(borders)

    @staticmethod
    def _convert_to_pdf(docx_path: Path) -> Path | None:
        """Convert .docx to PDF. Tries multiple methods in order.
        Returns the PDF path on success, or None if conversion is unavailable."""
        pdf_path = docx_path.with_suffix(".pdf")

        # Method 1: LibreOffice headless (works on Linux cloud servers)
        try:
            libreoffice_cmd = shutil.which("libreoffice") or shutil.which("soffice")
            if libreoffice_cmd:
                subprocess.run(
                    [libreoffice_cmd, "--headless", "--convert-to", "pdf",
                     "--outdir", str(docx_path.parent), str(docx_path)],
                    capture_output=True, timeout=60,
                )
                if pdf_path.exists():
                    return pdf_path
        except Exception:
            pass

        # Method 2: Direct win32com (Windows + MS Word) via temp directory.
        # Using a temp path avoids COM errors from OneDrive/spaces in paths.
        try:
            import win32com.client
            import pythoncom
            import tempfile

            pythoncom.CoInitialize()
            tmp_dir = Path(tempfile.mkdtemp(prefix="mctv_pdf_"))
            tmp_docx = tmp_dir / "contract.docx"
            tmp_pdf = tmp_dir / "contract.pdf"

            shutil.copy2(docx_path, tmp_docx)

            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            try:
                doc = word.Documents.Open(str(tmp_docx), ReadOnly=True)
                doc.SaveAs2(str(tmp_pdf), FileFormat=17)  # 17 = wdFormatPDF
                doc.Close(False)
            finally:
                word.Quit()
                pythoncom.CoUninitialize()

            if tmp_pdf.exists():
                shutil.copy2(tmp_pdf, pdf_path)
                # Clean up temp files
                shutil.rmtree(tmp_dir, ignore_errors=True)
                if pdf_path.exists():
                    print(f"[docx_service] PDF converted via win32com: {pdf_path}")
                    return pdf_path
        except ImportError:
            pass  # win32com not available (Linux)
        except Exception as e:
            print(f"[docx_service] win32com PDF conversion failed: {e}")

        # Method 3: docx2pdf fallback (also uses Word COM but different approach)
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                return pdf_path
        except Exception:
            pass

        return None

    @staticmethod
    def get_pdf_path(docx_path: Path) -> Path | None:
        """Check if a PDF version exists for a given docx file."""
        pdf_path = docx_path.with_suffix(".pdf")
        return pdf_path if pdf_path.exists() else None
