# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Field audit brief generator.

Produces the working document a contracted technician carries into the field:
scope, instructions, the capture legend, a routed stop list, one reference page
per stop with every screen's license number, and an appendix of everything we
already know is missing.

The companion .xlsx (``field_audit_service.build_workbook``) is where data is
actually recorded — this document is what tells someone what to do and gives
them the license numbers to match labels against. The two render from the same
``CAPTURE_COLUMNS`` list so their fields cannot drift apart.

Usage:
    generator = FieldAuditGenerator(config, docx_service)
    path = generator.generate(rows, gaps, market="tupelo",
                              prepared_for="Xceed Technologies")
"""

import logging
from datetime import datetime
from pathlib import Path

from docx.shared import Pt

from services.config_service import get_team_member
from services.docx_service import DocxService
from services.field_audit_service import (
    CAPTURE_COLUMNS,
    CAPTURE_LEGEND,
    TARGET_SECONDS,
    roster_summary,
)
from services.screen_inventory_service import market_label, mmss

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).parent.parent / "output" / "field_audits"

# Fields reproduced as fill-in lines on the paper stop pages. The full set
# lives in the workbook; this is what is worth writing by hand at the screen.
PAPER_CAPTURE_FIELDS = [
    "Screen on?", "Content playing?", "Display brand", "Size (in)",
    "Orientation", "Mount type", "Pi model", "Pi serial",
    "Power source", "Network", "Label applied?", "Condition",
]


class FieldAuditGenerator:
    """Builds the field audit brief as a branded Word document."""

    def __init__(self, config: dict, docx: DocxService):
        self.config = config
        self.docx = docx

    # ── Entry point ──────────────────────────────────────────────────────────

    def generate(self, rows: list[dict], gaps: list[dict], market: str,
                 prepared_for: str = "Xceed Technologies",
                 prepared_by_name: str = "", sweep_date: str | None = None,
                 progress_callback=None) -> Path:
        """Build the brief and save it. Returns the .docx path.

        Args:
            rows: Roster from ``audit_roster`` after ``route_order``.
            gaps: Output of ``roster_gaps``.
            market: Market key, e.g. "tupelo".
            prepared_for: The contractor performing the audit.
            prepared_by_name: MCTV team member; defaults to the first on file.
            sweep_date: Whitelist sweep the license numbers came from.
            progress_callback: ``fn(section_name, step, total_steps)``.
        """
        stops = self._by_stop(rows)
        total = 6 + len(stops)
        step = 0

        def tick(name: str) -> None:
            nonlocal step
            step += 1
            if progress_callback:
                progress_callback(name, step, total)

        summary = roster_summary(rows)
        rep = get_team_member(self.config, prepared_by_name) if prepared_by_name \
            else self.config["team"][0]

        doc = self.docx.create_document()

        tick("Cover page")
        self._add_cover(doc, market, prepared_for, rep)

        tick("Scope")
        self._add_scope(doc, market, summary, sweep_date, prepared_for)

        tick("Instructions")
        self._add_instructions(doc, prepared_for, summary)

        tick("What to record")
        self._add_legend(doc)

        tick("Route")
        self._add_route(doc, stops)

        for stop_no, screens in stops:
            tick(f"Stop {stop_no}: {screens[0]['venue_name']}")
            self._add_stop_page(doc, stop_no, screens)

        tick("Exceptions and escalation")
        self._add_exceptions(doc, gaps)
        self._add_escalation(doc)

        self.docx.add_footer(doc, f"{market_label(market)} Screen Audit — Confidential")

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        filename = (f"MCTV_FieldAudit_{market_label(market).replace(' ', '')}_"
                    f"{datetime.now().strftime('%Y-%m-%d')}.docx")
        path = OUTPUT_DIR / filename
        doc.save(path)
        logger.info("Field audit brief saved to %s", path)
        return path

    # ── Sections ─────────────────────────────────────────────────────────────

    def _add_cover(self, doc, market: str, prepared_for: str, rep: dict) -> None:
        self.docx.add_cover_page(
            doc,
            title=f"{market_label(market)} Screen Audit",
            subtitle="Field working document, screen inventory, and labelling instructions",
            prepared_for=prepared_for,
            prepared_by=rep,
        )

    def _add_scope(self, doc, market: str, summary: dict, sweep_date: str | None,
                   prepared_for: str) -> None:
        self.docx.add_section_header(doc, "Scope")

        self.docx.add_body_text(doc, (
            f"{prepared_for} is auditing every MCTV screen in the "
            f"{market_label(market)} market. This document lists each screen, "
            f"where it is, who to ask for, and the license number of the "
            f"Raspberry Pi driving it. A companion spreadsheet is where findings "
            f"get recorded, and a companion label sheet carries one printed label "
            f"per license."
        ))

        self.docx.add_metrics_banner(doc, {
            str(summary["screens"]): "Screens\nto audit",
            str(summary["venues"]): "Host\nvenues",
            str(summary["stops"]): "Stops on\nthe route",
            str(summary["missing_license"]): "License numbers\nto capture on site",
        })

        source = (f"License numbers come from the n-compass whitelist sweep dated "
                  f"{sweep_date}." if sweep_date else
                  "License numbers come from the most recent n-compass whitelist sweep.")
        self.docx.add_body_text(doc, (
            f"{source} Each license corresponds to one Raspberry Pi player. Where a "
            f"venue runs more than one screen, each screen has its own license and "
            f"its own label — they are not interchangeable."
        ))

        self.docx.add_callout_box(doc, (
            "A stop is complete when every screen at that address has been checked, "
            "its label applied, and its row filled in on the spreadsheet. A venue "
            "with two screens is not finished until both rows are filled in."
        ))

    def _add_instructions(self, doc, prepared_for: str, summary: dict) -> None:
        self.docx.add_section_header(doc, "How to run the audit", new_page=True)

        self.docx.add_selling_point(doc, "1. Print the label sheet first", (
            "The label sheet is a separate file, sized for Avery 5163 stock "
            "(4\" x 2\", ten labels per sheet). Print it before leaving. Each "
            "label carries a venue name, a short code, the full license number, "
            "and a QR code of that license number. Run one test page on plain "
            "paper and hold it against a label sheet to confirm alignment before "
            "printing on stock."
        ))

        self.docx.add_selling_point(doc, "2. Match the label to the screen, not the venue", (
            "At a venue with a single screen there is one label and no ambiguity. "
            "At a venue with two, the labels are not interchangeable — confirm which "
            "license is on which player before affixing. If you cannot tell them "
            "apart, note both license numbers and where each screen physically is "
            "(for example \"front counter\" and \"back room\"), and we will "
            "reconcile from the office."
        ))

        self.docx.add_selling_point(doc, "3. Affix the label where it can be read later", (
            "Put it on a flat face of the Raspberry Pi or its case, positioned so "
            "the license number can be read without unmounting the unit or "
            "unplugging anything. Do not cover vents, ports, or the SD card slot. "
            "If the Pi is sealed behind the display and cannot be reached, say so "
            "in the spreadsheet rather than putting the label on the television."
        ))

        self.docx.add_selling_point(doc, "4. Watch the screen before you write it down", (
            f"Stand with the display long enough to see at least two spot changes "
            f"before judging whether content is playing. A loop in this market runs "
            f"roughly {mmss(TARGET_SECONDS)}, so a frozen screen can look fine for a "
            f"few seconds. Each stop page lists the loop length we expect for that "
            f"screen."
        ))

        self.docx.add_selling_point(doc, "5. Photograph every screen", (
            "Two photographs per screen: one wide enough to show the display in its "
            "setting and how it is mounted, one close enough to read the new label "
            "on the Pi. Name the files with the short code from the label."
        ))

        self.docx.add_selling_point(doc, "6. Record as you go", (
            "Fill the spreadsheet row at the screen rather than from memory at the "
            "end of the day. The dropdown columns exist so the results import "
            "cleanly; free text in a dropdown column has to be corrected by hand."
        ))

        if summary["missing_license"]:
            self.docx.add_callout_box(doc, (
                f"{summary['missing_license']} screens in this packet have no license "
                f"number on file yet, so no label was printed for them. For those, "
                f"record the license number shown in the player's own settings or on "
                f"any existing sticker, and note where the screen physically sits. "
                f"We will produce those labels afterwards."
            ))

    def _add_legend(self, doc) -> None:
        self.docx.add_section_header(doc, "What to record at every screen", new_page=True)
        self.docx.add_body_text(doc, (
            "These are the spreadsheet columns, in order. Columns with a fixed list "
            "of answers show a dropdown in the file — please use it."
        ))
        self.docx.add_data_table(
            doc,
            headers=["Column", "What to record"],
            rows=[[header, CAPTURE_LEGEND.get(header, "")]
                  for header, _, _ in CAPTURE_COLUMNS],
        )

    def _add_route(self, doc, stops: list[tuple[int, list[dict]]]) -> None:
        self.docx.add_section_header(doc, "Route", new_page=True)
        self.docx.add_body_text(doc, (
            "Stops are grouped into clusters that can be driven together and "
            "sequenced to keep the day's mileage down. The order is a "
            "recommendation, not a requirement — what matters is that every stop "
            "is covered."
        ))

        rows = []
        for stop_no, screens in stops:
            first = screens[0]
            shared = " (shared building)" if first["same_building"] else ""
            rows.append([
                str(stop_no),
                first["venue_name"],
                (first["address"] or "Address needed") + shared,
                str(len(screens)),
                first["cluster"],
            ])
        self.docx.add_data_table(
            doc,
            headers=["Stop", "Venue", "Address", "Screens", "Cluster"],
            rows=rows,
        )

        shared_venues = sorted({s["venue_name"] for _, screens in stops
                                for s in screens if s["same_building"]})
        if shared_venues:
            self.docx.add_callout_box(doc, (
                "Some of these venues share a building — one trip covers several "
                "screens. Doing them together saves a return visit: "
                + ", ".join(shared_venues) + "."
            ))

    def _add_stop_page(self, doc, stop_no: int, screens: list[dict]) -> None:
        """One reference page per stop: where it is, who to ask for, and every
        screen's license number with room to write."""
        first = screens[0]
        label = f"Stop {stop_no} — {first['venue_name']}"
        self.docx.add_section_header(doc, label, new_page=True)

        detail_rows = [
            ["Address", first["address"] or "No address on file — ask MCTV before driving"],
            ["Cluster", first["cluster"]],
            ["Contact", self._contact_line(first)],
            ["Venue type", first["category"] or "Not recorded"],
            ["Screens here", str(len(screens))],
        ]
        if first["same_building"]:
            detail_rows.append(
                ["Note", "Another MCTV venue shares this building — check the route "
                         "list and cover them in one visit."])
        self.docx.add_data_table(doc, headers=["", ""], rows=detail_rows)

        for screen in screens:
            self._add_screen_block(doc, screen, len(screens) > 1)

    def _add_screen_block(self, doc, screen: dict, multi: bool) -> None:
        title = screen["short_code"] or "License number to be recorded on site"
        if multi and screen["screen_label"]:
            title = f"{title} — {screen['screen_label']}"

        expected = ""
        if screen["expected_loop_seconds"]:
            expected = f"Expected loop {screen['expected_loop_mmss']}"
            if screen["expected_items"]:
                expected += f" across {screen['expected_items']} spots"
            if screen["over_target"]:
                expected += f" (over the {mmss(TARGET_SECONDS)} target)"

        body = f"License number: {screen['license_id'] or '________________________________'}"
        if expected:
            body += f"\n{expected}."
        if screen["note"]:
            body += f"\n{screen['note']}"
        self.docx.add_accent_card(doc, title, body)

        # Fill-in lines, two per row, for whoever works off paper.
        pairs = [PAPER_CAPTURE_FIELDS[i:i + 2]
                 for i in range(0, len(PAPER_CAPTURE_FIELDS), 2)]
        table = self.docx.add_data_table(
            doc,
            headers=["Field", "Recorded", "Field", "Recorded"],
            rows=[[pair[0], "", pair[1] if len(pair) > 1 else "", ""]
                  for pair in pairs],
        )
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

        notes = doc.add_paragraph()
        notes.paragraph_format.space_before = Pt(4)
        run = notes.add_run("Issues / notes: " + "_" * 78)
        run.font.size = Pt(8)
        run.font.name = "Arial"

    def _add_exceptions(self, doc, gaps: list[dict]) -> None:
        self.docx.add_section_header(doc, "Known gaps and exceptions", new_page=True)
        if not gaps:
            self.docx.add_body_text(doc, "Nothing outstanding — the roster is complete.")
            return

        self.docx.add_body_text(doc, (
            "These are things MCTV already knows are incomplete or unusual. They "
            "are not findings against the auditor, and none of them should stop a "
            "stop from being completed."
        ))

        by_kind: dict[str, list[dict]] = {}
        for gap in gaps:
            by_kind.setdefault(gap["kind"], []).append(gap)

        for kind, items in by_kind.items():
            self.docx.add_sub_header(doc, f"{kind} ({len(items)})")
            self.docx.add_data_table(
                doc,
                headers=["Venue", "Detail", "What to do"],
                rows=[[g["venue"], g["detail"], g["action"]] for g in items],
            )

    def _add_escalation(self, doc) -> None:
        self.docx.add_section_header(doc, "Who to call", new_page=True)
        self.docx.add_body_text(doc, (
            "Call from the site rather than driving away with an open question — "
            "a second trip costs more than a phone call."
        ))
        self.docx.add_bullet_list(doc, "\n".join([
            "- Screen is dark or shows an error: Record it, photograph it, and move on. Do not attempt a repair or a reboot unless MCTV asks for one.",
            "- Business has closed or moved: Record it and move on. Do not remove equipment.",
            "- Staff will not allow access: Record who you spoke to and when, and call MCTV before leaving.",
            "- Screen exists that is not in this packet: Record the venue, the location, and any license number visible on the player.",
            "- Screen in this packet does not exist at the venue: Say so explicitly. A missing screen is one of the most valuable findings in this audit.",
        ]))
        self.docx.add_team_section(
            doc,
            closing_text="Thank you for the careful work — this is the first "
                         "complete physical record of the network.",
            dark_mode=True,
        )

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _by_stop(rows: list[dict]) -> list[tuple[int, list[dict]]]:
        """Group the roster by stop, preserving route order."""
        stops: dict[int, list[dict]] = {}
        for row in rows:
            stops.setdefault(row["stop_no"] or 0, []).append(row)
        return sorted(stops.items())

    @staticmethod
    def _contact_line(screen: dict) -> str:
        name = screen["contact_name"]
        if not name or name.lower() == "on file":
            return "No named contact — ask for the manager on duty and record the name"
        parts = [name]
        if screen["contact_role"]:
            parts.append(f"({screen['contact_role']})")
        if screen["contact_phone"]:
            parts.append(f"— {screen['contact_phone']}")
        else:
            parts.append("— no phone on file")
        return " ".join(parts)
