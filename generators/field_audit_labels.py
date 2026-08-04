# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Raspberry Pi license labels, laid out for Avery 5163 stock.

One label per license: venue name, the short code, the full license UUID in
readable type, and a QR code of that UUID so a phone can read it back off a
unit mounted behind a television.

Avery 5163 is 4" x 2", two across and five down on US Letter:

    side margins   0.15625"    top / bottom margins   0.5"
    column gap     0.1875"     row height             2" exactly

Laid out as a three-column table (label, gap, label) with exact row heights so
the cells cannot drift as text wraps — the whole point of a label sheet is that
row eleven still lines up with the stock.

Screens with no license number on file get no label; they are reported back to
the caller so the brief can say how many are outstanding.

Usage:
    path, printed, skipped = generate_labels(rows, market="tupelo")
"""

import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from services.screen_inventory_service import market_label

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).parent.parent / "output" / "field_audits"

# Avery 5163 geometry, in inches.
LABEL_W = 4.0
LABEL_H = 2.0
COLS, ROWS = 2, 5
SIDE_MARGIN = 0.15625
TOP_MARGIN = 0.5
COL_GAP = 0.1875
PER_SHEET = COLS * ROWS

NAVY = RGBColor(0x1B, 0x1F, 0x3B)
GOLD = RGBColor(0xC5, 0xA5, 0x5A)
GREY = RGBColor(0x55, 0x55, 0x55)

SUPPORT_LINE = "Property of MCTV Digital  |  portal@mctvofms.com"

# Longest venue name that still fits the label's text column on one line.
MAX_VENUE_CHARS = 30


def _qr_png(payload: str, box_size: int = 8) -> BytesIO | None:
    """QR code for a license number, as an in-memory PNG.

    Returns None when the optional ``qrcode`` package is unavailable — the
    sheet must still print, just without the scan shortcut.
    """
    try:
        import qrcode
    except ImportError:
        logger.warning("qrcode not installed — printing text-only labels")
        return None
    try:
        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=box_size,
            border=1,
        )
        qr.add_data(payload)
        qr.make(fit=True)
        buf = BytesIO()
        qr.make_image(fill_color="black", back_color="white").save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception as exc:  # pragma: no cover - depends on optional backend
        logger.warning("QR generation failed (%s) — printing text-only labels", exc)
        return None


def _set_cell_margins(cell, top=0.04, left=0.08, bottom=0.02, right=0.06) -> None:
    """Padding inside a label cell, in inches."""
    tc_pr = cell._element.get_or_add_tcPr()
    mar = tc_pr.makeelement(qn("w:tcMar"), {})
    for tag, value in (("top", top), ("left", left),
                       ("bottom", bottom), ("right", right)):
        node = mar.makeelement(qn(f"w:{tag}"), {
            qn("w:w"): str(int(value * 1440)),
            qn("w:type"): "dxa",
        })
        mar.append(node)
    tc_pr.append(mar)


def _strip_borders(table) -> None:
    """Remove every border. Label stock is pre-cut; printed lines miss the die
    cut by a hair and look like a mistake."""
    tbl_pr = table._element.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none",
                                                     qn("w:sz"): "0"})
        borders.append(node)
    tbl_pr.append(borders)


def _fit(text: str, limit: int = MAX_VENUE_CHARS) -> str:
    text = (text or "").strip()
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def _draw_label(cell, screen: dict) -> None:
    """Render one label into a table cell: text on the left, QR on the right."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_margins(cell)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    inner = cell.add_table(rows=1, cols=2)
    inner.alignment = WD_TABLE_ALIGNMENT.LEFT
    _strip_borders(inner)
    text_cell, qr_cell = inner.rows[0].cells
    text_cell.width = Inches(2.85)
    qr_cell.width = Inches(0.95)
    for sub in (text_cell, qr_cell):
        _set_cell_margins(sub, top=0, left=0, bottom=0, right=0)

    def line(container, text, size, *, bold=False, color=NAVY, space_after=0):
        para = container.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
        return para

    # The first paragraph of a fresh cell is empty; reuse it rather than
    # leaving a blank line at the top of every label.
    head = text_cell.paragraphs[0]
    head.paragraph_format.space_before = Pt(0)
    head.paragraph_format.space_after = Pt(1)
    head_run = head.add_run("MCTV DIGITAL  ·  DO NOT REMOVE")
    head_run.font.name = "Arial"
    head_run.font.size = Pt(6.5)
    head_run.font.bold = True
    head_run.font.color.rgb = GOLD

    line(text_cell, _fit(screen["venue_name"]), 10.5, bold=True, space_after=1)
    if screen.get("screen_label"):
        line(text_cell, _fit(screen["screen_label"], 34), 7.5, color=GREY, space_after=1)
    line(text_cell, screen["short_code"], 17, bold=True, space_after=2)
    line(text_cell, screen["license_id"], 7.5, color=GREY, space_after=1)
    line(text_cell, SUPPORT_LINE, 6.5, color=GREY)

    qr_para = qr_cell.paragraphs[0]
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_para.paragraph_format.space_before = Pt(2)
    qr_para.paragraph_format.space_after = Pt(0)
    png = _qr_png(screen["license_id"])
    if png is not None:
        qr_para.add_run().add_picture(png, width=Inches(0.9))


def _new_sheet_table(doc) -> "object":
    """One Avery sheet: 5 rows of (label, gap, label) at exact heights."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    table = doc.add_table(rows=ROWS, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _strip_borders(table)

    for row in table.rows:
        row.height = Inches(LABEL_H)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for width, cell in zip((LABEL_W, COL_GAP, LABEL_W), row.cells):
            cell.width = Inches(width)
    return table


def generate_labels(rows: list[dict], market: str, copies: int = 1,
                    output_dir: Path | None = None) -> tuple[Path, int, list[dict]]:
    """Build the label sheet.

    Args:
        rows: Roster from ``audit_roster`` (route order is respected).
        market: Market key, used in the filename.
        copies: Labels printed per license. Two is worth it when a unit is
            awkward to reach and the first label goes on crooked.
        output_dir: Override the default ``output/field_audits``.

    Returns:
        ``(path, printed, skipped)`` — ``skipped`` is the roster rows that had
        no license number, so the caller can say how many are outstanding.
    """
    printable = [r for r in rows if r.get("license_id") and r.get("short_code")]
    skipped = [r for r in rows if not r.get("license_id")]

    labels: list[dict] = []
    for row in printable:
        labels.extend([row] * max(1, copies))

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(SIDE_MARGIN)
    section.right_margin = Inches(SIDE_MARGIN)
    section.top_margin = Inches(TOP_MARGIN)
    section.bottom_margin = Inches(TOP_MARGIN)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(0)

    sheets = [labels[i:i + PER_SHEET] for i in range(0, len(labels), PER_SHEET)] or [[]]

    for sheet_idx, sheet in enumerate(sheets):
        if sheet_idx:
            # A new section per sheet, so the page break cannot land mid-grid
            # and shift every label below it.
            new = doc.add_section(WD_SECTION.NEW_PAGE)
            new.left_margin = Inches(SIDE_MARGIN)
            new.right_margin = Inches(SIDE_MARGIN)
            new.top_margin = Inches(TOP_MARGIN)
            new.bottom_margin = Inches(TOP_MARGIN)
        table = _new_sheet_table(doc)
        for pos, screen in enumerate(sheet):
            row, col = divmod(pos, COLS)
            _draw_label(table.rows[row].cells[col * 2], screen)

    # A leading empty paragraph pushes the whole grid down by a line and every
    # label misses its die cut. Only the very first child qualifies, and only
    # when it carries no sectPr — the section breaks between sheets live in
    # empty paragraphs too, and removing one of those merges two sheets.
    body = doc.element.body
    children = list(body)
    if children and children[0].tag == qn("w:p"):
        first = children[0]
        has_sect = first.find(qn("w:pPr")) is not None and \
            first.find(qn("w:pPr")).find(qn("w:sectPr")) is not None
        if not has_sect and not "".join(first.itertext()).strip():
            body.remove(first)

    target = output_dir or OUTPUT_DIR
    target.mkdir(parents=True, exist_ok=True)
    filename = (f"MCTV_FieldAuditLabels_{market_label(market).replace(' ', '')}_"
                f"{datetime.now().strftime('%Y-%m-%d')}.docx")
    path = target / filename
    doc.save(path)

    logger.info("Label sheet saved to %s (%d labels, %d sheets, %d skipped)",
                path, len(labels), len(sheets), len(skipped))
    return path, len(labels), skipped
