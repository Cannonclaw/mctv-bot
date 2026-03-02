# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Branded contract document generator.

Produces professional MCTV advertising contracts as Word documents
with optional PDF conversion. Reuses DocxService for consistent branding.

Supports 5 contract types:
  - advertiser: Standard advertising partnership
  - host: Venue hosting agreement ($0 cost)
  - host_advertising: Host venue buying additional paid screens
  - category_exclusivity: Category-exclusive advertising agreement
  - bundle: Multi-brand bundle advertising agreement
"""

import re
import json
from pathlib import Path
from datetime import datetime, timedelta
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from services.docx_service import DocxService
from services.config_service import get_tier_impressions, calculate_cpm


# ── Config ──────────────────────────────────────────────────────────────────

ROOT = Path(__file__).parent.parent
CONFIG_PATH = ROOT / "config" / "config.json"
OUTPUT_DIR = ROOT / "output" / "contracts"


def _load_config() -> dict:
    """Load config.json."""
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


# ── Standard contract clauses ───────────────────────────────────────────────

STANDARD_CLAUSES = [
    {
        "title": "1. Services Provided",
        "body": (
            "MCTV Elite Advertising ('MCTV') agrees to display the Advertiser's "
            "digital content on the specified number of indoor digital billboard screens "
            "within the selected market(s). Content will play in rotation at a rate of "
            "4 plays per hour per screen during standard venue operating hours."
        ),
    },
    {
        "title": "2. Content Ownership",
        "body": (
            "All creative content produced by MCTV for this partnership remains the "
            "property of the Advertiser. The Advertiser may use any content created by "
            "MCTV on social media, websites, print, or any other medium at no additional "
            "cost. MCTV retains a non-exclusive right to display the content on the "
            "MCTV network for the duration of this agreement."
        ),
    },
    {
        "title": "3. Term and Renewal",
        "body": (
            "This agreement begins on the Start Date specified above and continues "
            "for the agreed term. Unless either party provides written notice of "
            "non-renewal at least 30 days before the end of the current term, this "
            "agreement will automatically renew for successive periods of equal length "
            "at the then-current rate."
        ),
    },
    {
        "title": "4. Payment Terms",
        "body": (
            "Invoices are issued monthly on the first business day of each month. "
            "Payment is due within 30 days of the invoice date. MCTV reserves the "
            "right to suspend ad display for accounts more than 45 days past due. "
            "A late fee of 1.5% per month may be applied to overdue balances."
        ),
    },
    {
        "title": "5. Content Guidelines",
        "body": (
            "All content must comply with applicable local, state, and federal laws. "
            "MCTV reserves the right to refuse or remove content that is offensive, "
            "misleading, or inappropriate for the venue environments in which screens "
            "are placed. The Advertiser will be notified promptly if any content is "
            "declined and given the opportunity to submit revised creative."
        ),
    },
    {
        "title": "6. Cancellation",
        "body": (
            "Either party may cancel this agreement with 60 days written notice. "
            "If the Advertiser cancels before the end of the agreed term, any prepay "
            "bonuses or discounts already applied will be prorated and billed. "
            "MCTV will make reasonable efforts to wind down ad display gracefully."
        ),
    },
    {
        "title": "7. Liability Limitation",
        "body": (
            "MCTV's total liability under this agreement shall not exceed the total "
            "fees paid by the Advertiser in the 12 months preceding the claim. "
            "MCTV shall not be liable for incidental, consequential, or indirect "
            "damages. MCTV makes no guarantee of specific business results from "
            "advertising display."
        ),
    },
    {
        "title": "8. Governing Law",
        "body": (
            "This agreement shall be governed by and construed in accordance with "
            "the laws of the State of Mississippi. Any disputes arising under this "
            "agreement shall be resolved in the courts of Hinds County, Mississippi."
        ),
    },
]


HOST_CLAUSES = [
    {
        "title": "1. Screen Hosting Agreement",
        "body": (
            "The Venue Host agrees to provide a suitable location for the installation "
            "and operation of MCTV indoor digital billboard screen(s). MCTV will provide, "
            "install, and maintain all hardware at no cost to the Venue Host. The screen(s) "
            "will display a mix of advertiser content and venue-specific content."
        ),
    },
    {
        "title": "2. Venue Benefits",
        "body": (
            "In exchange for hosting MCTV screen(s), the Venue Host receives complimentary "
            "advertising plays on all MCTV screens across the network. This includes free "
            "ad content creation by the MCTV creative team. Play frequency and screen count "
            "are specified in the Partnership Details section above."
        ),
    },
    {
        "title": "3. Screen Placement and Maintenance",
        "body": (
            "MCTV will work with the Venue Host to identify optimal screen placement for "
            "maximum visibility. MCTV is responsible for all hardware maintenance, repairs, "
            "and content updates. The Venue Host agrees to provide a standard electrical "
            "outlet and reliable WiFi connectivity for each screen."
        ),
    },
    {
        "title": "4. Term and Removal",
        "body": (
            "This agreement begins on the Start Date specified above and continues "
            "for the agreed term. Either party may terminate with 30 days written notice. "
            "Upon termination, MCTV will remove all hardware within 14 business days at "
            "no cost to the Venue Host."
        ),
    },
    {
        "title": "5. Content Control",
        "body": (
            "MCTV manages all content displayed on hosted screens. The Venue Host may "
            "request that specific advertiser categories be excluded from display in "
            "their venue (e.g., competitor ads). MCTV will make reasonable efforts to "
            "honor such requests. Venue Host content is subject to the same content "
            "guidelines as advertiser content."
        ),
    },
    {
        "title": "6. Liability and Insurance",
        "body": (
            "MCTV maintains liability insurance covering its equipment and operations. "
            "The Venue Host is not responsible for damage to MCTV equipment except in "
            "cases of gross negligence or intentional damage. MCTV shall not be liable "
            "for any loss of business or incidental damages arising from screen downtime."
        ),
    },
    {
        "title": "7. Governing Law",
        "body": (
            "This agreement shall be governed by and construed in accordance with "
            "the laws of the State of Mississippi. Any disputes arising under this "
            "agreement shall be resolved in the courts of Hinds County, Mississippi."
        ),
    },
]


HOST_ADVERTISING_CLAUSES = [
    {
        "title": "1. Host Advertising Partnership",
        "body": (
            "This agreement covers additional advertising screens for an existing "
            "MCTV Venue Host ('Host Advertiser'). As a current venue host, the Host "
            "Advertiser receives complimentary advertising on a specified number of "
            "screens across the MCTV network. This agreement provides advertising on "
            "additional screens beyond the complimentary allotment at a preferred "
            "discounted rate, as detailed in the Partnership Details section above."
        ),
    },
    {
        "title": "2. Complimentary Screens & Additional Screens",
        "body": (
            "The Host Advertiser's complimentary screen allotment (earned through "
            "the venue hosting agreement) remains unchanged by this agreement. The "
            "additional screens covered by this agreement are billed at the discounted "
            "monthly rate specified in the Partnership Details. The discount is applied "
            "as a courtesy to recognize the Host Advertiser's ongoing partnership with "
            "MCTV as a venue host."
        ),
    },
    {
        "title": "3. Content Ownership",
        "body": (
            "All creative content produced by MCTV for this partnership remains the "
            "property of the Host Advertiser. The Host Advertiser may use any content "
            "created by MCTV on social media, websites, print, or any other medium at "
            "no additional cost. MCTV retains a non-exclusive right to display the "
            "content on the MCTV network for the duration of this agreement."
        ),
    },
    {
        "title": "4. Term and Renewal",
        "body": (
            "This agreement begins on the Start Date specified above and continues "
            "for the agreed term. Unless either party provides written notice of "
            "non-renewal at least 30 days before the end of the current term, this "
            "agreement will automatically renew for successive periods of equal length "
            "at the then-current discounted rate."
        ),
    },
    {
        "title": "5. Payment Terms",
        "body": (
            "Invoices are issued monthly on the first business day of each month. "
            "Payment is due within 30 days of the invoice date. MCTV reserves the "
            "right to suspend additional ad display (beyond the complimentary allotment) "
            "for accounts more than 45 days past due. A late fee of 1.5% per month "
            "may be applied to overdue balances."
        ),
    },
    {
        "title": "6. Content Guidelines",
        "body": (
            "All content must comply with applicable local, state, and federal laws. "
            "MCTV reserves the right to refuse or remove content that is offensive, "
            "misleading, or inappropriate for the venue environments in which screens "
            "are placed. The Host Advertiser will be notified promptly if any content "
            "is declined and given the opportunity to submit revised creative."
        ),
    },
    {
        "title": "7. Cancellation",
        "body": (
            "Either party may cancel this agreement with 60 days written notice. "
            "Cancellation of this advertising add-on does not affect the underlying "
            "venue hosting agreement or the complimentary screen allotment. If the "
            "Host Advertiser cancels before the end of the agreed term, any prepay "
            "bonuses or discounts already applied will be prorated and billed."
        ),
    },
    {
        "title": "8. Liability Limitation",
        "body": (
            "MCTV's total liability under this agreement shall not exceed the total "
            "fees paid by the Host Advertiser in the 12 months preceding the claim. "
            "MCTV shall not be liable for incidental, consequential, or indirect "
            "damages. MCTV makes no guarantee of specific business results from "
            "advertising display."
        ),
    },
    {
        "title": "9. Governing Law",
        "body": (
            "This agreement shall be governed by and construed in accordance with "
            "the laws of the State of Mississippi. Any disputes arising under this "
            "agreement shall be resolved in the courts of Hinds County, Mississippi."
        ),
    },
]


CATEGORY_EXCLUSIVITY_CLAUSES = [
    {
        "title": "1. Exclusive Advertising Rights",
        "body": (
            "MCTV Elite Advertising ('MCTV') grants the Advertiser exclusive "
            "advertising rights within their designated business category across "
            "the specified market(s) and screen count. For the duration of this "
            "agreement, MCTV will not display advertising from any direct competitor "
            "within the same business category on the screens covered by this contract."
        ),
    },
    {
        "title": "2. Category Definition and Scope",
        "body": (
            "The exclusive business category is defined in the Partnership Details "
            "section above. MCTV and the Advertiser will agree on a clear category "
            "definition at the time of signing. If a borderline case arises (e.g., a "
            "business that partially overlaps the exclusive category), MCTV will "
            "consult the Advertiser before accepting that content for display. "
            "Exclusivity applies only to the screens and markets specified in this "
            "agreement."
        ),
    },
    {
        "title": "3. Premium Pricing and Value",
        "body": (
            "The monthly rate for this agreement reflects a premium for category "
            "exclusivity. This premium ensures the Advertiser is the only business "
            "in their category reaching audiences on the covered screens. The premium "
            "is calculated as a percentage above the standard advertising rate and is "
            "detailed in the Partnership Details section."
        ),
    },
    {
        "title": "4. Content Ownership",
        "body": (
            "All creative content produced by MCTV for this partnership remains the "
            "property of the Advertiser. The Advertiser may use any content created by "
            "MCTV on social media, websites, print, or any other medium at no additional "
            "cost. MCTV retains a non-exclusive right to display the content on the "
            "MCTV network for the duration of this agreement."
        ),
    },
    {
        "title": "5. Term and Renewal",
        "body": (
            "This agreement begins on the Start Date specified above and continues "
            "for the agreed term. Exclusivity rights are tied to the contract term. "
            "Unless either party provides written notice of non-renewal at least 30 "
            "days before the end of the current term, this agreement will automatically "
            "renew for successive periods of equal length at the then-current rate. "
            "Upon renewal, exclusivity rights continue uninterrupted."
        ),
    },
    {
        "title": "6. Payment Terms",
        "body": (
            "Invoices are issued monthly on the first business day of each month. "
            "Payment is due within 30 days of the invoice date. MCTV reserves the "
            "right to suspend ad display for accounts more than 45 days past due. "
            "If ad display is suspended due to non-payment, exclusivity protection "
            "will also be suspended until the account is current. A late fee of 1.5% "
            "per month may be applied to overdue balances."
        ),
    },
    {
        "title": "7. Content Guidelines",
        "body": (
            "All content must comply with applicable local, state, and federal laws. "
            "MCTV reserves the right to refuse or remove content that is offensive, "
            "misleading, or inappropriate for the venue environments in which screens "
            "are placed. The Advertiser will be notified promptly if any content is "
            "declined and given the opportunity to submit revised creative."
        ),
    },
    {
        "title": "8. Exclusivity Breach Remedy",
        "body": (
            "If MCTV inadvertently displays competitor content on the covered screens, "
            "MCTV will remove the competing content within 24 hours of notification and "
            "credit the Advertiser a pro-rated amount for the days the competing content "
            "was displayed. The Advertiser agrees to notify MCTV promptly upon discovery "
            "of any potential breach."
        ),
    },
    {
        "title": "9. Cancellation",
        "body": (
            "Either party may cancel this agreement with 60 days written notice. "
            "If the Advertiser cancels before the end of the agreed term, any prepay "
            "bonuses or discounts already applied will be prorated and billed. "
            "Upon cancellation, exclusivity rights end on the effective cancellation "
            "date and MCTV may accept new advertisers in the previously exclusive "
            "category."
        ),
    },
    {
        "title": "10. Liability Limitation",
        "body": (
            "MCTV's total liability under this agreement shall not exceed the total "
            "fees paid by the Advertiser in the 12 months preceding the claim. "
            "MCTV shall not be liable for incidental, consequential, or indirect "
            "damages. MCTV makes no guarantee of specific business results from "
            "advertising display."
        ),
    },
    {
        "title": "11. Governing Law",
        "body": (
            "This agreement shall be governed by and construed in accordance with "
            "the laws of the State of Mississippi. Any disputes arising under this "
            "agreement shall be resolved in the courts of Hinds County, Mississippi."
        ),
    },
]


BUNDLE_CLAUSES = [
    {
        "title": "1. Multi-Brand Bundle Agreement",
        "body": (
            "This agreement covers advertising for multiple brands or business "
            "locations ('Bundle') owned or operated by the same entity. Each brand "
            "included in this Bundle is listed in the Partnership Details section "
            "above. All brands share a single contract and billing relationship "
            "with MCTV Elite Advertising ('MCTV')."
        ),
    },
    {
        "title": "2. Bundle Pricing and Savings",
        "body": (
            "The monthly rate for this agreement reflects the multi-brand bundle "
            "discount. By consolidating multiple brands under a single contract, "
            "the Advertiser receives a reduced per-brand rate compared to individual "
            "contracts. The specific brands, screen allocations, and combined rate "
            "are detailed in the Partnership Details section."
        ),
    },
    {
        "title": "3. Adding and Removing Brands",
        "body": (
            "The Advertiser may request to add additional brands to this Bundle at "
            "any time. Additional brands will be added at the current bundle rate and "
            "the monthly total will be adjusted accordingly. Removing a brand from the "
            "Bundle requires 30 days written notice. If removing a brand reduces the "
            "Bundle below the minimum qualifying size, the remaining brands will revert "
            "to standard individual pricing."
        ),
    },
    {
        "title": "4. Content Ownership",
        "body": (
            "All creative content produced by MCTV for each brand in this Bundle "
            "remains the property of the respective brand owner. Each brand may use "
            "its MCTV-created content on social media, websites, print, or any other "
            "medium at no additional cost. MCTV retains a non-exclusive right to "
            "display the content on the MCTV network for the duration of this agreement."
        ),
    },
    {
        "title": "5. Term and Renewal",
        "body": (
            "This agreement begins on the Start Date specified above and continues "
            "for the agreed term. Unless either party provides written notice of "
            "non-renewal at least 30 days before the end of the current term, this "
            "agreement will automatically renew for successive periods of equal length "
            "at the then-current bundle rate."
        ),
    },
    {
        "title": "6. Payment Terms",
        "body": (
            "A single consolidated invoice is issued monthly on the first business "
            "day of each month covering all brands in the Bundle. Payment is due "
            "within 30 days of the invoice date. MCTV reserves the right to suspend "
            "ad display for all brands in the Bundle for accounts more than 45 days "
            "past due. A late fee of 1.5% per month may be applied to overdue balances."
        ),
    },
    {
        "title": "7. Content Guidelines",
        "body": (
            "All content for each brand must comply with applicable local, state, and "
            "federal laws. MCTV reserves the right to refuse or remove content that is "
            "offensive, misleading, or inappropriate for the venue environments in which "
            "screens are placed. The Advertiser will be notified promptly if any content "
            "is declined and given the opportunity to submit revised creative."
        ),
    },
    {
        "title": "8. Cancellation",
        "body": (
            "Either party may cancel this agreement with 60 days written notice. "
            "Cancellation applies to all brands in the Bundle. If the Advertiser "
            "wishes to continue advertising for individual brands after cancellation, "
            "new individual contracts must be established at standard rates. Any prepay "
            "bonuses or discounts already applied will be prorated and billed."
        ),
    },
    {
        "title": "9. Liability Limitation",
        "body": (
            "MCTV's total liability under this agreement shall not exceed the total "
            "fees paid by the Advertiser in the 12 months preceding the claim. "
            "MCTV shall not be liable for incidental, consequential, or indirect "
            "damages. MCTV makes no guarantee of specific business results from "
            "advertising display."
        ),
    },
    {
        "title": "10. Governing Law",
        "body": (
            "This agreement shall be governed by and construed in accordance with "
            "the laws of the State of Mississippi. Any disputes arising under this "
            "agreement shall be resolved in the courts of Hinds County, Mississippi."
        ),
    },
]


RENEWAL_CLAUSES = [
    {
        "title": "1. Continuation of Services",
        "body": (
            "MCTV Elite Advertising ('MCTV') and the Renewing Advertiser agree to "
            "continue the advertising partnership established under the original "
            "agreement. All services, screen placements, and content rotations will "
            "continue without interruption during the transition to this renewed term. "
            "This renewal supersedes and replaces the prior agreement in its entirety."
        ),
    },
    {
        "title": "2. Renewed Terms",
        "body": (
            "The terms of this renewed agreement are specified in the Partnership "
            "Details section above. Any changes to screen count, monthly rate, or "
            "market coverage from the original agreement are reflected in the updated "
            "Partnership Details. Where this renewal is silent on a matter, the terms "
            "of the original agreement shall apply."
        ),
    },
    {
        "title": "3. Rate & Loyalty Pricing",
        "body": (
            "The monthly rate for this renewal reflects MCTV's commitment to long-term "
            "partnerships. Renewing advertisers may receive preferential pricing or rate "
            "locks as a loyalty benefit. Any loyalty adjustments are reflected in the "
            "Partnership Details section. MCTV reserves the right to adjust rates at "
            "the next renewal period with 30 days advance written notice."
        ),
    },
    {
        "title": "4. Term & Future Renewal",
        "body": (
            "This renewed agreement begins on the Start Date specified above and "
            "continues for the agreed term. Unless either party provides written notice "
            "of non-renewal at least 30 days before the end of the current term, this "
            "agreement will automatically renew for successive periods of equal length "
            "at the then-current rate."
        ),
    },
    {
        "title": "5. Payment Continuation",
        "body": (
            "Invoices continue to be issued monthly on the first business day of each "
            "month. Payment is due within 30 days of the invoice date. MCTV reserves "
            "the right to suspend ad display for accounts more than 45 days past due. "
            "A late fee of 1.5% per month may be applied to overdue balances. Any "
            "outstanding balance from the prior agreement carries forward."
        ),
    },
    {
        "title": "6. Content Transition",
        "body": (
            "All existing creative content from the prior agreement carries over to "
            "this renewed term at no additional cost. The Renewing Advertiser retains "
            "ownership of all content created by MCTV during both the original and "
            "renewed terms. Fresh creative updates are included as part of this "
            "renewed partnership."
        ),
    },
    {
        "title": "7. Cancellation",
        "body": (
            "Either party may cancel this agreement with 60 days written notice. "
            "If the Renewing Advertiser cancels before the end of the renewed term, "
            "any prepay bonuses or discounts already applied will be prorated and "
            "billed. MCTV will make reasonable efforts to wind down ad display "
            "gracefully."
        ),
    },
    {
        "title": "8. Governing Law",
        "body": (
            "This agreement shall be governed by and construed in accordance with "
            "the laws of the State of Mississippi. Any disputes arising under this "
            "agreement shall be resolved in the courts of Hinds County, Mississippi."
        ),
    },
]


# Map contract types to their clause sets
_CLAUSE_MAP = {
    "advertiser": STANDARD_CLAUSES,
    "host": HOST_CLAUSES,
    "host_advertising": HOST_ADVERTISING_CLAUSES,
    "category_exclusivity": CATEGORY_EXCLUSIVITY_CLAUSES,
    "bundle": BUNDLE_CLAUSES,
    "renewal": RENEWAL_CLAUSES,
}


# ── Contract Generator ──────────────────────────────────────────────────────

def _format_date(iso_date: str) -> str:
    """Convert ISO date (YYYY-MM-DD) to human-readable format (e.g., February 24, 2026)."""
    if not iso_date:
        return "TBD"
    try:
        dt = datetime.strptime(iso_date[:10], "%Y-%m-%d")
        return dt.strftime("%B %d, %Y").replace(" 0", " ")  # "February 04" → "February 4"
    except (ValueError, TypeError):
        return iso_date  # Return as-is if parsing fails


class ContractGenerator:
    """Generate branded MCTV contract documents."""

    # Human-readable labels for each contract type
    CONTRACT_LABELS = {
        "advertiser": "Advertising Partnership Agreement",
        "host": "Venue Hosting Agreement",
        "host_advertising": "Host Advertising Agreement",
        "category_exclusivity": "Category Exclusivity Agreement",
        "bundle": "Multi-Brand Bundle Agreement",
        "renewal": "Advertising Partnership Renewal",
    }

    # Party labels for signature blocks
    PARTY_LABELS = {
        "advertiser": "Advertiser",
        "host": "Venue Host",
        "host_advertising": "Host Advertiser",
        "category_exclusivity": "Exclusive Advertiser",
        "bundle": "Bundle Advertiser",
        "renewal": "Renewing Advertiser",
    }

    def __init__(self, config: dict | None = None):
        self.config = config or _load_config()
        self.docx_service = DocxService(self.config, color_scheme="original")

    def generate(
        self,
        client_name: str,
        business_name: str,
        contract_type: str = "advertiser",
        tier_name: str = "",
        screen_count: int = 10,
        monthly_rate: float = 350.0,
        term_months: int = 6,
        markets: list[str] | None = None,
        start_date: str = "",
        auto_renew: bool = True,
        prepared_by: str = "",
        notes: str = "",
        # Category exclusivity fields
        exclusive_category: str = "",
        # Bundle fields
        bundle_brands: list[str] | None = None,
        # Multi-tier options (Good/Better/Best comparison)
        tier_options: list[dict] | None = None,
    ) -> tuple[Path, Path | None, bytes]:
        """Generate a contract document.

        Args:
            client_name: Contact person name
            business_name: Business/venue name
            contract_type: "advertiser", "host", "host_advertising",
                           "category_exclusivity", or "bundle"
            tier_name: Tier label (e.g., "20 Screens")
            screen_count: Number of screens
            monthly_rate: Monthly fee
            term_months: Contract length in months
            markets: List of market names
            start_date: Contract start date (YYYY-MM-DD), defaults to today
            auto_renew: Whether contract auto-renews
            prepared_by: Sales rep name
            notes: Additional notes/terms
            exclusive_category: Business category for exclusivity contracts
            bundle_brands: List of brand names for bundle contracts

        Returns:
            Tuple of (docx_path, pdf_path or None, docx_bytes).
            docx_bytes captured before PDF conversion to avoid file locking.
        """
        if not markets:
            markets = ["Oxford"]
        if not start_date:
            start_date = datetime.now().strftime("%Y-%m-%d")
        if not tier_name:
            tier_name = f"{screen_count} Screens"

        # Calculate end date
        start_dt = datetime.strptime(start_date, "%Y-%m-%d")
        end_dt = start_dt + timedelta(days=term_months * 30)
        end_date = end_dt.strftime("%Y-%m-%d")

        # Find prepared_by rep info
        rep_info = self._get_rep_info(prepared_by)

        # ── Build the document ──────────────────────────────────────
        doc = self.docx_service.create_document()

        # Cover page
        contract_label = self.CONTRACT_LABELS.get(contract_type, "Partnership Agreement")
        self.docx_service.add_cover_page(
            doc,
            title=contract_label,
            subtitle=business_name,
            prepared_for=f"{client_name}\n{business_name}",
            prepared_by=rep_info,
            date=datetime.now().strftime("%B %d, %Y"),
        )

        # ── Partnership Details section ─────────────────────────────
        self.docx_service.add_section_header(doc, "Partnership Details")

        # Multi-tier comparison table (if tier_options provided)
        if tier_options and len(tier_options) >= 2 and contract_type in (
            "advertiser", "renewal",
        ):
            self._add_tier_comparison_table(doc, tier_options, term_months,
                                            markets, start_date, end_date,
                                            auto_renew, business_name,
                                            client_name)
        elif contract_type == "host_advertising":
            self._add_host_advertising_details(
                doc, business_name, client_name, tier_name, screen_count,
                monthly_rate, term_months, markets, start_date, end_date,
                auto_renew,
            )
        elif contract_type == "host":
            self._add_host_details(
                doc, business_name, client_name, screen_count, term_months,
                markets, start_date, end_date,
            )
        elif contract_type == "category_exclusivity":
            self._add_exclusivity_details(
                doc, business_name, client_name, tier_name, screen_count,
                monthly_rate, term_months, markets, start_date, end_date,
                auto_renew, exclusive_category,
            )
        elif contract_type == "bundle":
            self._add_bundle_details(
                doc, business_name, client_name, tier_name, screen_count,
                monthly_rate, term_months, markets, start_date, end_date,
                auto_renew, bundle_brands or [],
            )
        elif contract_type == "renewal":
            self._add_renewal_details(
                doc, business_name, client_name, tier_name, screen_count,
                monthly_rate, term_months, markets, start_date, end_date,
                auto_renew,
            )
        else:
            self._add_advertiser_details(
                doc, business_name, client_name, tier_name, screen_count,
                monthly_rate, term_months, markets, start_date, end_date,
                auto_renew,
            )

        if notes:
            self.docx_service.add_body_text(doc, "")
            self.docx_service.add_callout_box(doc, f"Additional Notes: {notes}")

        # ── Value Recap section (paid contracts only) ────────────────
        if contract_type != "host":
            self._add_value_recap(
                doc, contract_type, monthly_rate, screen_count, term_months,
                business_name,
            )

        # ── Terms & Conditions section ──────────────────────────────
        self.docx_service.add_section_header(doc, "Terms & Conditions", new_page=True)

        clauses = _CLAUSE_MAP.get(contract_type, STANDARD_CLAUSES)
        for clause in clauses:
            self.docx_service.add_accent_card(doc, clause["title"], clause["body"])

        # ── Signature section ───────────────────────────────────────
        self.docx_service.add_section_header(doc, "Agreement & Signature", new_page=True)
        self._add_signature_section(doc, business_name, client_name, contract_type)

        # ── Footer ──────────────────────────────────────────────────
        self.docx_service.add_footer(doc, "Partnership Agreement")

        # ── Save ────────────────────────────────────────────────────
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        safe_name = "".join(c if c.isalnum() or c in " -_" else "" for c in business_name)
        safe_name = safe_name.strip().replace(" ", "_")[:40]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"MCTV_Contract_{safe_name}_{timestamp}.docx"

        docx_path = OUTPUT_DIR / filename
        doc.save(docx_path)

        # Read docx bytes immediately (before PDF conversion may lock the file)
        docx_bytes = docx_path.read_bytes()

        # Try PDF conversion
        pdf_path = self.docx_service._convert_to_pdf(docx_path)

        return docx_path, pdf_path, docx_bytes

    # ------------------------------------------------------------------
    # Value Recap — metrics banner + prepay callout between details & T&C
    # ------------------------------------------------------------------

    def _add_value_recap(self, doc, contract_type, monthly_rate, screen_count,
                         term_months, business_name):
        """Add a 'What You're Getting' value recap with metrics and prepay."""
        self.docx_service.add_section_header(
            doc, "What You're Getting", new_page=True,
        )

        # Calculate key metrics
        network = self.config.get("network", {})
        plays_per_hour = network.get("plays_per_hour", 4)
        hours_per_day = network.get("hours_per_day", 12)
        days_per_month = network.get("days_per_month", 30)
        daily_plays = screen_count * plays_per_hour * hours_per_day
        monthly_plays = daily_plays * days_per_month
        impressions = get_tier_impressions(self.config, screen_count)
        cpm = calculate_cpm(monthly_rate, impressions)
        daily_cost = monthly_rate / 30 if monthly_rate > 0 else 0

        # Metrics banner
        metrics = {}
        metrics[f"{screen_count}"] = "Screens\nWorking for You"
        metrics[f"{monthly_plays:,}"] = "Ad Plays\nPer Month"
        if cpm > 0:
            metrics[f"${cpm:.2f}"] = "Your CPM\n(Cost per 1K)"
        metrics[f"${daily_cost:.2f}"] = "Daily\nInvestment"

        self.docx_service.add_metrics_banner(doc, metrics)

        # ROI context line
        self.docx_service.add_body_text(
            doc,
            f"For ${daily_cost:.2f} a day, {business_name} gets {daily_plays:,} "
            f"ad plays across {screen_count} screens in the venues where your "
            f"customers already spend their time \u2014 restaurants, gyms, offices, "
            f"barbershops, and more. Every play is unskippable, full-screen, and "
            f"right at eye level."
        )

        # Prepay bonuses callout (prominent, before legal terms)
        if contract_type in ("advertiser", "host_advertising",
                             "category_exclusivity", "bundle"):
            pricing = self.config.get("pricing", {})
            contract_terms = pricing.get("contract_terms", {})
            bonus_6 = contract_terms.get("prepay_6mo_bonus", "")
            bonus_12 = contract_terms.get("prepay_12mo_bonus", "")

            if bonus_6 or bonus_12:
                lines = ["PREPAYMENT BONUSES"]
                if bonus_6:
                    lines.append(f"\u25A0  6-Month Prepay: {bonus_6}")
                if bonus_12:
                    lines.append(f"\u25A0  12-Month Prepay: {bonus_12}")
                lines.append("")
                lines.append(
                    "Pay upfront and get free months added to your contract. "
                    "Ask your rep about prepay options."
                )
                self.docx_service.add_callout_box(doc, "\n".join(lines))

    # ------------------------------------------------------------------
    # Partnership Details builders
    # ------------------------------------------------------------------

    def _get_rep_info(self, rep_name: str) -> dict:
        """Get rep info from config, or return default."""
        team = self.config.get("team", [])
        for member in team:
            if rep_name and rep_name.lower() in member.get("name", "").lower():
                return member
        # Default to first team member
        if team:
            return team[0]
        return {
            "name": "MCTV Elite Advertising",
            "title": "",
            "email": "info@mctvofms.com",
            "phone": "",
        }

    def _add_advertiser_details(self, doc, business_name, client_name,
                                 tier_name, screen_count, monthly_rate,
                                 term_months, markets, start_date, end_date,
                                 auto_renew):
        """Add the advertiser partnership details table."""
        total_value = monthly_rate * term_months

        details = [
            ("Advertiser", business_name),
            ("Contact", client_name),
            ("Package", tier_name),
            ("Screen Count", str(screen_count)),
            ("Monthly Rate", f"${monthly_rate:,.2f}"),
            ("Term", f"{term_months} months"),
            ("Total Contract Value", f"${total_value:,.2f}"),
            ("Market(s)", ", ".join(markets)),
            ("Start Date", _format_date(start_date)),
            ("End Date", _format_date(end_date)),
            ("Auto-Renew", "Yes" if auto_renew else "No"),
            ("Content Creation", "Included at no additional cost"),
            ("Plays per Hour", "4 per screen"),
        ]

        self.docx_service.add_data_table(
            doc,
            headers=["Detail", "Value"],
            rows=details,
        )

    def _add_host_details(self, doc, business_name, client_name,
                           screen_count, term_months, markets,
                           start_date, end_date):
        """Add the host venue partnership details table."""
        pricing = self.config.get("pricing", {})
        free_inside = pricing.get("host_free_inside_plays_per_hour", 8)
        free_outside_plays = pricing.get("host_outside_plays_per_hour", 4)
        free_outside_screens = pricing.get("host_free_outside_screens", 10)

        details = [
            ("Venue", business_name),
            ("Contact", client_name),
            ("Screens Hosted", str(screen_count)),
            ("Term", f"{term_months} months"),
            ("Market(s)", ", ".join(markets)),
            ("Start Date", _format_date(start_date)),
            ("End Date", _format_date(end_date)),
            ("Venue Ad Plays (In-Store)", f"{free_inside} plays/hour across network"),
            ("Venue Ad Plays (Outside)", f"{free_outside_plays} plays/hour on {free_outside_screens} screens"),
            ("Content Creation", "Included at no additional cost"),
            ("Hardware & Installation", "Provided and maintained by MCTV"),
            ("Venue Cost", "$0 / Free"),
        ]

        self.docx_service.add_data_table(
            doc,
            headers=["Detail", "Value"],
            rows=details,
        )

    def _add_host_advertising_details(self, doc, business_name, client_name,
                                       tier_name, screen_count, monthly_rate,
                                       term_months, markets, start_date,
                                       end_date, auto_renew):
        """Add the host advertising partnership details table."""
        pricing = self.config.get("pricing", {})
        free_screens = pricing.get("host_free_outside_screens", 10)
        total_value = monthly_rate * term_months

        # Parse discount from tier_name (e.g., "Host Discount 10% - 10 Screens")
        discount_str = "N/A"
        standard_rate_str = ""
        if "Host Discount" in tier_name:
            match = re.search(r"(\d+)%", tier_name)
            if match:
                pct = int(match.group(1))
                discount_str = f"{pct}%"
                if pct > 0:
                    standard_rate = monthly_rate / (1 - pct / 100)
                    standard_rate_str = f"${standard_rate:,.2f}"

        details = [
            ("Host Advertiser", business_name),
            ("Contact", client_name),
            ("Complimentary Screens", f"{free_screens} (included with hosting agreement)"),
            ("Additional Paid Screens", str(screen_count)),
            ("Total Screens", str(free_screens + screen_count)),
        ]

        if standard_rate_str:
            details.append(("Standard Rate", f"{standard_rate_str}/mo"))
            details.append(("Host Discount", discount_str))
        details += [
            ("Discounted Monthly Rate", f"${monthly_rate:,.2f}"),
            ("Term", f"{term_months} months"),
            ("Total Contract Value", f"${total_value:,.2f}"),
            ("Market(s)", ", ".join(markets)),
            ("Start Date", _format_date(start_date)),
            ("End Date", _format_date(end_date)),
            ("Auto-Renew", "Yes" if auto_renew else "No"),
            ("Content Creation", "Included at no additional cost"),
            ("Plays per Hour", "4 per screen"),
        ]

        self.docx_service.add_data_table(
            doc,
            headers=["Detail", "Value"],
            rows=details,
        )

    def _add_exclusivity_details(self, doc, business_name, client_name,
                                  tier_name, screen_count, monthly_rate,
                                  term_months, markets, start_date, end_date,
                                  auto_renew, exclusive_category):
        """Add the category exclusivity partnership details table."""
        total_value = monthly_rate * term_months

        details = [
            ("Exclusive Advertiser", business_name),
            ("Contact", client_name),
            ("Exclusive Category", exclusive_category or "To be defined at signing"),
            ("Package", tier_name),
            ("Screen Count", str(screen_count)),
            ("Monthly Rate (incl. exclusivity premium)", f"${monthly_rate:,.2f}"),
            ("Term", f"{term_months} months"),
            ("Total Contract Value", f"${total_value:,.2f}"),
            ("Market(s)", ", ".join(markets)),
            ("Start Date", _format_date(start_date)),
            ("End Date", _format_date(end_date)),
            ("Auto-Renew", "Yes" if auto_renew else "No"),
            ("Exclusivity Scope", f"No competing {exclusive_category or 'category'} ads on your screens"),
            ("Content Creation", "Included at no additional cost"),
            ("Plays per Hour", "4 per screen"),
        ]

        self.docx_service.add_data_table(
            doc,
            headers=["Detail", "Value"],
            rows=details,
        )

    def _add_bundle_details(self, doc, business_name, client_name,
                             tier_name, screen_count, monthly_rate,
                             term_months, markets, start_date, end_date,
                             auto_renew, bundle_brands):
        """Add the multi-brand bundle partnership details table."""
        total_value = monthly_rate * term_months
        brand_count = len(bundle_brands) if bundle_brands else 1
        per_brand = monthly_rate / brand_count if brand_count > 0 else monthly_rate

        details = [
            ("Bundle Owner", business_name),
            ("Contact", client_name),
            ("Brands in Bundle", str(brand_count)),
        ]

        # List each brand
        if bundle_brands:
            for i, brand in enumerate(bundle_brands, 1):
                details.append((f"Brand {i}", brand))

        details += [
            ("Package", tier_name),
            ("Total Screen Count", str(screen_count)),
            ("Combined Monthly Rate", f"${monthly_rate:,.2f}"),
            ("Per-Brand Rate", f"${per_brand:,.2f}/mo"),
            ("Term", f"{term_months} months"),
            ("Total Contract Value", f"${total_value:,.2f}"),
            ("Market(s)", ", ".join(markets)),
            ("Start Date", _format_date(start_date)),
            ("End Date", _format_date(end_date)),
            ("Auto-Renew", "Yes" if auto_renew else "No"),
            ("Content Creation", f"Included for all {brand_count} brands"),
            ("Plays per Hour", "4 per screen per brand"),
        ]

        self.docx_service.add_data_table(
            doc,
            headers=["Detail", "Value"],
            rows=details,
        )

    # ------------------------------------------------------------------
    # Renewal details
    # ------------------------------------------------------------------

    def _add_renewal_details(self, doc, business_name, client_name,
                              tier_name, screen_count, monthly_rate,
                              term_months, markets, start_date, end_date,
                              auto_renew):
        """Add the renewal partnership details table."""
        total_value = monthly_rate * term_months

        details = [
            ("Renewing Advertiser", business_name),
            ("Contact", client_name),
            ("Package", tier_name),
            ("Screen Count", str(screen_count)),
            ("Monthly Rate", f"${monthly_rate:,.2f}"),
            ("Renewed Term", f"{term_months} months"),
            ("Total Contract Value", f"${total_value:,.2f}"),
            ("Market(s)", ", ".join(markets)),
            ("Renewal Start Date", _format_date(start_date)),
            ("Renewal End Date", _format_date(end_date)),
            ("Auto-Renew", "Yes" if auto_renew else "No"),
            ("Content Creation", "Included — existing creative carries over"),
            ("Plays per Hour", "4 per screen"),
        ]

        self.docx_service.add_data_table(
            doc,
            headers=["Detail", "Value"],
            rows=details,
        )

        self.docx_service.add_callout_box(
            doc,
            "THANK YOU FOR RENEWING\n\n"
            "We appreciate your continued partnership with MCTV Elite Advertising. "
            "Your existing ad creative and screen placements will continue "
            "seamlessly into your renewed term."
        )

    # ------------------------------------------------------------------
    # Multi-tier comparison table (Good / Better / Best)
    # ------------------------------------------------------------------

    def _add_tier_comparison_table(self, doc, tier_options, term_months,
                                    markets, start_date, end_date,
                                    auto_renew, business_name, client_name):
        """Add a side-by-side tier comparison table for multi-option contracts."""
        from docx.shared import Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self.docx_service.add_body_text(
            doc,
            f"We have prepared {len(tier_options)} package options for "
            f"{business_name}. Review the comparison below and select the "
            f"package that best fits your goals."
        )

        network = self.config.get("network", {})
        plays_per_hour = network.get("plays_per_hour", 4)
        hours_per_day = network.get("hours_per_day", 12)
        days_per_month = network.get("days_per_month", 30)

        # Recommended = middle tier (index 1 for 3 tiers, 0 for 2)
        rec_idx = 1 if len(tier_options) > 2 else 0

        # Build comparison rows
        row_labels = [
            "Package",
            "Screens",
            "Monthly Rate",
            "Daily Investment",
            "Monthly Ad Plays",
            "CPM",
        ]

        num_tiers = len(tier_options)
        table = doc.add_table(rows=len(row_labels) + 2, cols=num_tiers + 1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row with tier names
        header_row = table.rows[0]
        header_row.cells[0].text = ""
        for i, tier in enumerate(tier_options):
            cell = header_row.cells[i + 1]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label = tier.get("name", f"Option {i + 1}")
            if i == rec_idx:
                label += "\n★ RECOMMENDED"
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Gold background for recommended, navy for others
            tc_pr = cell._element.get_or_add_tcPr()
            bg = self.docx_service.c["accent_hex"] if i == rec_idx else self.docx_service.c["bg_hex"]
            shd = tc_pr.makeelement(qn("w:shd"), {
                qn("w:fill"): bg, qn("w:val"): "clear",
            })
            tc_pr.append(shd)

        # Data rows
        for r, label in enumerate(row_labels):
            row = table.rows[r + 1]
            # Label cell
            lp = row.cells[0].paragraphs[0]
            lr = lp.add_run(label)
            lr.font.size = Pt(9)
            lr.font.bold = True
            lr.font.name = "Arial"
            lr.font.color.rgb = self.docx_service.c["text"]

            for i, tier in enumerate(tier_options):
                cell = row.cells[i + 1]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                screens = tier.get("screens", 10)
                rate = float(tier.get("rate", 350))
                daily = rate / 30
                plays = screens * plays_per_hour * hours_per_day * days_per_month
                impressions = get_tier_impressions(self.config, screens)
                cpm = calculate_cpm(rate, impressions)

                values = {
                    "Package": tier.get("name", ""),
                    "Screens": str(screens),
                    "Monthly Rate": f"${rate:,.0f}",
                    "Daily Investment": f"${daily:.2f}",
                    "Monthly Ad Plays": f"{plays:,}",
                    "CPM": f"${cpm:.2f}" if cpm > 0 else "N/A",
                }
                val = values.get(label, "")
                vr = p.add_run(val)
                vr.font.size = Pt(9)
                vr.font.name = "Arial"
                vr.font.color.rgb = self.docx_service.c["text"]

                # Light gold background for recommended column
                if i == rec_idx:
                    tc_pr = cell._element.get_or_add_tcPr()
                    shd = tc_pr.makeelement(qn("w:shd"), {
                        qn("w:fill"): "FDF6E8", qn("w:val"): "clear",
                    })
                    tc_pr.append(shd)

        # "Select Your Package" row
        sel_row = table.rows[-1]
        lp = sel_row.cells[0].paragraphs[0]
        lr = lp.add_run("Your Selection")
        lr.font.size = Pt(9)
        lr.font.bold = True
        lr.font.name = "Arial"
        for i in range(num_tiers):
            cell = sel_row.cells[i + 1]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("☐")
            run.font.size = Pt(14)

        # Common details below the comparison
        self.docx_service.add_body_text(doc, "")
        common_details = [
            ("Advertiser", business_name),
            ("Contact", client_name),
            ("Term", f"{term_months} months"),
            ("Market(s)", ", ".join(markets)),
            ("Start Date", _format_date(start_date)),
            ("End Date", _format_date(end_date)),
            ("Auto-Renew", "Yes" if auto_renew else "No"),
            ("Content Creation", "Included at no additional cost"),
            ("Plays per Hour", "4 per screen"),
        ]
        self.docx_service.add_data_table(
            doc, headers=["Detail", "Value"], rows=common_details,
        )

    # ------------------------------------------------------------------
    # Signature section — two-column layout with e-sign notice above
    # ------------------------------------------------------------------

    def _add_signature_section(self, doc, business_name, client_name,
                                contract_type):
        """Add the signature and agreement block.

        Layout: e-sign notice callout → agreement text → two-column
        signature table (client left, MCTV right).
        """
        party_label = self.PARTY_LABELS.get(contract_type, "Advertiser")

        # E-signature notice — prominent, above signatures
        self.docx_service.add_callout_box(
            doc,
            "ELECTRONIC SIGNATURE NOTICE\n\n"
            "This contract may be signed electronically through the MCTV Client "
            "Portal. Electronic signatures made through the portal (typed name + "
            "'I Agree' confirmation) are legally equivalent to handwritten "
            "signatures under the Mississippi Uniform Electronic Transactions Act."
        )

        self.docx_service.add_body_text(doc, "")

        # Agreement text
        agreement_text = (
            f"By signing below, {client_name} on behalf of {business_name} "
            f"('{party_label}') agrees to the terms and conditions outlined in this "
            f"agreement. This agreement is legally binding upon signature by both parties."
        )
        self.docx_service.add_body_text(doc, agreement_text)
        self.docx_service.add_body_text(doc, "")

        # Two-column signature table
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.docx_service._clear_table_borders(table)

        # Left column — client signature
        left_cell = table.rows[0].cells[0]
        self._format_signature_cell(left_cell, party_label, client_name)

        # Right column — MCTV signature
        right_cell = table.rows[0].cells[1]
        self._format_signature_cell(right_cell, "MCTV Elite Advertising", "")

    def _format_signature_cell(self, cell, party_name, prefill_name):
        """Format a single signature cell with accent border and fields."""
        tc_pr = cell._element.get_or_add_tcPr()

        # Light background
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:fill"): self.docx_service.c["light_hex"],
            qn("w:val"): "clear",
        })
        tc_pr.append(shd)

        # Cell padding
        cell_margin = tc_pr.makeelement(qn("w:tcMar"), {})
        for side, val in (("top", "200"), ("bottom", "200"),
                          ("left", "240"), ("right", "200")):
            m = cell_margin.makeelement(qn(f"w:{side}"), {
                qn("w:w"): val, qn("w:type"): "dxa",
            })
            cell_margin.append(m)
        tc_pr.append(cell_margin)

        # Accent left border
        self.docx_service._set_cell_borders(
            cell,
            left_color=self.docx_service.c["accent_hex"],
            left_sz=36,
            other_color="D0D0D0",
            other_sz=2,
        )

        # Party name header
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(party_name)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.docx_service.c["primary"]
        run.font.name = "Arial"

        # Signature fields
        fields = [
            ("Signed By", prefill_name),
            ("Printed Name", prefill_name),
            ("Title", ""),
            ("Date", ""),
        ]
        for label, prefill in fields:
            fp = cell.add_paragraph()
            fp.paragraph_format.space_before = Pt(6)
            fp.paragraph_format.space_after = Pt(2)

            label_run = fp.add_run(f"{label}: ")
            label_run.font.size = Pt(9)
            label_run.font.bold = True
            label_run.font.color.rgb = self.docx_service.c["text"]
            label_run.font.name = "Arial"

            line_run = fp.add_run(prefill if prefill else "________________________________")
            line_run.font.size = Pt(9)
            line_run.font.color.rgb = self.docx_service.c["text"]
            line_run.font.name = "Arial"
