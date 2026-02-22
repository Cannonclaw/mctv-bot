"""Venue Partner Traction Report generator.

Generates branded Word documents showing a venue partner what value MCTV is
providing to their location. This is the venue-facing counterpart to the
advertiser traction report -- it shows the venue how many advertisers are
running on their screens, total ad plays, screen time, and which advertisers
are active at their location.

Used for hosts like restaurants, barbershops, medical offices, gyms, etc.
"""

import logging
from pathlib import Path
from datetime import datetime
from collections import defaultdict

from services.claude_service import ClaudeService
from services.docx_service import DocxService, NAVY, GOLD, GRAY, DARK_TEXT
from services.config_service import get_team_member
from models.report_data import TractionReportInput, VenueRecord

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class VenueReportGenerator:
    """Generates venue-partner-facing traction reports as Word documents.

    These reports show a venue host the value MCTV brings to their business --
    how many advertisers are running on their screens, total plays, screen
    time, and individual advertiser activity.

    Usage:
        generator = VenueReportGenerator(config, claude_service, docx_service)
        report_path = generator.generate(venue_data, progress_callback=update_ui)
    """

    def __init__(self, config: dict, claude: ClaudeService, docx: DocxService):
        """Initialize the venue report generator.

        Args:
            config: Master application config dict (from config.json).
            claude: ClaudeService instance for AI-generated insights.
            docx: DocxService instance for document creation and formatting.
        """
        self.config = config
        self.claude = claude
        self.docx = docx

    def generate(self, data: TractionReportInput, progress_callback=None) -> Path:
        """Generate a complete venue partner traction report.

        Builds a branded Word document with the venue's performance summary,
        advertiser breakdown table, optional AI insights about the venue's
        value to the network, and MCTV footer branding.

        Args:
            data: TractionReportInput populated with venue-centric data.
                  For venue reports, data.advertiser_name holds the venue name,
                  and data.venue_records contain per-advertiser play data
                  (each VenueRecord represents an advertiser's activity at
                  this venue, with host_name as the advertiser/content name).
            progress_callback: Optional callable(section_name, step, total)
                               for UI progress updates.

        Returns:
            Path to the saved .docx report file.
        """
        total_steps = self._count_steps(data)
        current_step = 0

        # --- Step 1: Create document ---
        doc = self.docx.create_document()
        current_step += 1
        if progress_callback:
            progress_callback("Creating document", current_step, total_steps)

        # --- Step 2: Title and subtitle ---
        self._add_title_block(doc, data)
        current_step += 1
        if progress_callback:
            progress_callback("Adding report header", current_step, total_steps)

        # --- Step 3: Summary metrics banner ---
        advertiser_records = self._get_advertiser_records(data)
        self._add_summary_metrics(doc, data, advertiser_records)
        current_step += 1
        if progress_callback:
            progress_callback("Adding summary metrics", current_step, total_steps)

        # --- Step 4: Advertiser breakdown table ---
        self._add_advertiser_table(doc, data, advertiser_records)
        current_step += 1
        if progress_callback:
            progress_callback("Building advertiser breakdown", current_step, total_steps)

        # --- Step 5: AI insights (if requested) ---
        if data.include_insights:
            current_step += 1
            if progress_callback:
                progress_callback("Generating AI insights", current_step, total_steps)
            self._add_insights_section(doc, data, advertiser_records)

        # --- Step 6: Footer ---
        self.docx.add_footer(doc)
        current_step += 1
        if progress_callback:
            progress_callback("Finalizing report", current_step, total_steps)

        # --- Save ---
        venue_name = data.advertiser_name  # For venue reports, this holds the venue name
        safe_name = self._make_safe_filename(venue_name)
        date_str = datetime.now().strftime("%Y-%m-%d")
        filename = f"MCTV_VenueReport_{safe_name}_{date_str}.docx"
        report_path = self.docx.save_report(doc, filename)

        logger.info(
            "Venue partner report saved: %s (%d advertisers, %s total plays)",
            report_path.name, len(advertiser_records),
            f"{data.total_plays:,}"
        )
        return report_path

    # ------------------------------------------------------------------
    # Document sections
    # ------------------------------------------------------------------

    def _add_title_block(self, doc, data: TractionReportInput):
        """Add the report title, venue name, and reporting period."""
        venue_name = data.advertiser_name

        # MCTV branding header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("MCTV ELITE ADVERTISING")
        run.font.size = Pt(12)
        run.font.color.rgb = GOLD
        run.font.bold = True
        run.font.name = "Calibri"

        # Spacer
        doc.add_paragraph()

        # Venue name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(venue_name)
        run.font.size = Pt(28)
        run.font.color.rgb = NAVY
        run.font.bold = True
        run.font.name = "Calibri"

        # Report type subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Venue Partner Traction Report")
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        run.font.bold = False
        run.font.name = "Calibri"

        # Reporting period
        doc.add_paragraph()
        if data.campaign_period:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Reporting Period: {data.campaign_period}")
            run.font.size = Pt(11)
            run.font.color.rgb = GRAY
            run.font.name = "Calibri"

        # Report generation date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        run.font.italic = True

        # Divider
        doc.add_paragraph()

    def _add_summary_metrics(self, doc, data: TractionReportInput,
                             advertiser_records: list):
        """Add the large summary metrics banner."""
        self.docx.add_section_header(doc, "Your Venue at a Glance")

        # Count unique advertisers from the venue records
        unique_advertisers = self._count_unique_advertisers(data)

        # Build metrics dict (key = display value, value = label)
        metrics = {}

        metrics[str(unique_advertisers)] = "Total Advertisers\nRunning"

        metrics[f"{data.total_plays:,}"] = "Total Ad Plays\nat Your Venue"

        if data.total_air_time:
            metrics[data.total_air_time] = "Total\nScreen Time"

        self.docx.add_metrics_banner(doc, metrics)

    def _add_advertiser_table(self, doc, data: TractionReportInput,
                              advertiser_records: list):
        """Add the per-advertiser breakdown table for this venue."""
        self.docx.add_section_header(doc, "Advertiser Activity at Your Venue")

        venue_name = data.advertiser_name

        if not advertiser_records:
            self.docx.add_body_text(
                doc,
                "No advertiser activity data is available for this reporting period."
            )
            return

        self.docx.add_body_text(
            doc,
            f"The following advertisers have been running campaigns on the screens "
            f"at {venue_name}. This activity demonstrates the value your venue "
            f"brings to the MCTV network and to local businesses in your community."
        )

        # Build table
        headers = ["Advertiser", "Plays at Your Venue", "Screen Time", "Active Since"]

        # Sort by plays descending
        sorted_records = sorted(
            advertiser_records, key=lambda r: r["total_plays"], reverse=True
        )

        rows = []
        for rec in sorted_records:
            rows.append([
                rec["advertiser_name"],
                f"{rec['total_plays']:,}",
                rec["total_air_time"],
                rec["active_since"] or "--",
            ])

        self.docx.add_data_table(doc, headers, rows)

    def _add_insights_section(self, doc, data: TractionReportInput,
                              advertiser_records: list):
        """Generate and add AI-powered insights about the venue's network value."""
        self.docx.add_section_header(doc, "Your Value to the MCTV Network")

        prompt = self._build_insights_prompt(data, advertiser_records)

        try:
            insights_text = self.claude.generate_insights(prompt)
            if insights_text:
                self.docx.add_body_text(doc, insights_text)
            else:
                logger.warning("Claude returned empty venue insights; using fallback.")
                self._add_insights_fallback(doc, data)
        except Exception as e:
            logger.error("Failed to generate venue AI insights: %s", e)
            self._add_insights_fallback(doc, data)

    def _add_insights_fallback(self, doc, data: TractionReportInput):
        """Add a static fallback when AI insights are unavailable."""
        self.docx.add_body_text(
            doc,
            f"As an MCTV venue partner, {data.advertiser_name} plays an important "
            f"role in connecting local advertisers with your customers. Your venue's "
            f"screen time and foot traffic contribute real value to our network. "
            f"Contact your MCTV representative for a detailed analysis of your "
            f"venue's performance and partnership benefits."
        )

    # ------------------------------------------------------------------
    # Helper methods
    # ------------------------------------------------------------------

    def _get_advertiser_records(self, data: TractionReportInput) -> list:
        """Extract per-advertiser records from the venue data.

        For venue reports, each VenueRecord in data.venue_records represents
        a content/advertiser that has played at the venue. We aggregate by
        the content_name or host_name field (depending on how data was built)
        to produce advertiser-level summaries.

        Returns a list of dicts with keys:
            advertiser_name, total_plays, total_air_time, active_since
        """
        advertiser_map = defaultdict(lambda: {
            "advertiser_name": "",
            "total_plays": 0,
            "total_seconds": 0,
            "active_since": "",
        })

        for venue in data.venue_records:
            # In venue reports, host_name is used as the advertiser/content name
            adv_name = venue.host_name
            entry = advertiser_map[adv_name]
            entry["advertiser_name"] = adv_name
            entry["total_plays"] += venue.total_plays
            entry["total_air_time"] = venue.total_air_time or ""

            # Track earliest active date
            if venue.first_aired:
                if not entry["active_since"] or venue.first_aired < entry["active_since"]:
                    entry["active_since"] = venue.first_aired

        # Convert air time -- use the stored string if single record,
        # otherwise note it from the venue record directly
        records = list(advertiser_map.values())

        # Ensure total_air_time is set for display
        for rec in records:
            if not rec.get("total_air_time"):
                rec["total_air_time"] = "--"

        return records

    def _count_unique_advertisers(self, data: TractionReportInput) -> int:
        """Count the number of unique content/advertiser names in venue records."""
        unique_names = set()
        for venue in data.venue_records:
            if venue.host_name and venue.host_name.strip():
                unique_names.add(venue.host_name.strip())
        return len(unique_names)

    def _build_insights_prompt(self, data: TractionReportInput,
                               advertiser_records: list) -> str:
        """Build a prompt for Claude to generate venue value insights."""
        venue_name = data.advertiser_name
        unique_advertisers = self._count_unique_advertisers(data)

        # Top advertisers at this venue
        top_advertisers = sorted(
            advertiser_records, key=lambda r: r["total_plays"], reverse=True
        )[:5]
        adv_summary = "\n".join(
            f"  - {a['advertiser_name']}: {a['total_plays']:,} plays, "
            f"screen time: {a['total_air_time']}"
            for a in top_advertisers
        )

        prompt = (
            f"You are writing a value insights section for a venue partner traction "
            f"report. The venue is {venue_name}, a host location in the MCTV indoor "
            f"digital billboard network.\n\n"
            f"Venue Performance Summary:\n"
            f"  - Total Advertisers Running: {unique_advertisers}\n"
            f"  - Total Ad Plays at This Venue: {data.total_plays:,}\n"
            f"  - Total Screen Time: {data.total_air_time}\n"
        )

        if data.campaign_period:
            prompt += f"  - Reporting Period: {data.campaign_period}\n"

        prompt += f"\nTop Advertisers at This Venue:\n{adv_summary}\n"

        prompt += (
            f"\nWrite 2-3 concise paragraphs explaining the value this venue provides "
            f"to the MCTV network and to the local advertisers running on their screens. "
            f"Emphasize the foot traffic and dwell time that makes {venue_name} an "
            f"effective advertising location. Mention that having {unique_advertisers} "
            f"active advertisers demonstrates real demand for their screen placement. "
            f"Keep the tone appreciative and professional -- this is a partner "
            f"relationship update. Do not use bullet points or headers -- just flowing "
            f"paragraphs."
        )

        return prompt

    def _count_steps(self, data: TractionReportInput) -> int:
        """Count total generation steps for progress reporting."""
        steps = 4  # create doc, title, metrics, advertiser table
        if data.include_insights:
            steps += 1
        steps += 1  # footer
        return steps

    @staticmethod
    def _make_safe_filename(name: str) -> str:
        """Convert a venue name to a filesystem-safe string."""
        safe = name.strip()
        for char in r' /\:*?"<>|':
            safe = safe.replace(char, "_")
        safe = safe.replace("'", "")
        while "__" in safe:
            safe = safe.replace("__", "_")
        return safe.strip("_")
