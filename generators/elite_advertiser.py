"""Elite Advertiser proposal generator (McGlawn / MS Urgent Care style)."""

from generators.base_proposal import BaseProposal
from services.config_service import get_team_member, get_pricing_tier, get_all_tiers


class EliteAdvertiserProposal(BaseProposal):
    """Generates the flagship 7-8 page advertiser proposal."""

    @property
    def proposal_type_key(self) -> str:
        return "elite_advertiser"

    def get_sections(self) -> list:
        return [
            ("opportunity", "The Opportunity"),
            ("whats_included", "What's Included"),
            ("market_coverage", "Your Market Coverage"),
            ("_pricing", "Partnership Pricing"),
            ("why_mctv", "Why MCTV Elite Advertising"),
            ("getting_started", "Let's Get Started"),
            ("_team", "Meet Your Team"),
        ]

    def get_prompt_variables(self, data) -> dict:
        rep = get_team_member(self.config, data.sales_rep)
        markets = ", ".join(data.selected_markets)

        # Build market details string
        market_details = []
        for market_name in data.selected_markets:
            market = self.config["markets"].get(market_name, {})
            screens = market.get("screens", 0)
            market_details.append(f"{market_name} ({screens} screens)")
        selected_markets = ", ".join(market_details)

        return {
            "business_name": data.business_name,
            "contact_name": data.contact_name,
            "contact_email": data.contact_email,
            "industry": data.industry,
            "city": data.city,
            "selected_markets": selected_markets,
            "sales_rep": rep["name"],
            "additional_notes": data.additional_notes or "No additional notes.",
            "host_free_screens": self.config["pricing"]["host_free_outside_screens"],
        }

    def build_section(self, doc, section_key, data, content):
        if section_key == "opportunity":
            self._build_opportunity(doc, data, content)
        elif section_key == "whats_included":
            self._build_whats_included(doc, content)
        elif section_key == "market_coverage":
            self._build_market_coverage(doc, data, content)
        elif section_key == "_pricing":
            self._build_pricing(doc, data)
        elif section_key == "why_mctv":
            self._build_why_mctv(doc, content)
        elif section_key == "getting_started":
            self._build_getting_started(doc, data, content)
        elif section_key == "_team":
            doc.add_page_break()
            self.docx.add_team_section(doc)

    def _build_opportunity(self, doc, data, content):
        self.docx.add_section_header(doc, "The Opportunity")
        self.docx.add_body_text(doc, content)

        # Network stats banner
        network = self.config["network"]
        self.docx.add_metrics_banner(doc, {
            network["total_screens"]: "Screens Across\nNorth Mississippi",
            network["monthly_impressions"]: "Monthly Impressions",
            f"{network['avg_dwell_time_minutes']}+ Min": "Avg. Dwell Time\nPer Visit",
            f"{network['plays_per_hour']}x/Hour": "Your Ad Plays\nEvery Day",
        })
        doc.add_page_break()

    def _build_whats_included(self, doc, content):
        self.docx.add_section_header(doc, "What's Included")
        self.docx.add_body_text(doc, content)

    def _build_market_coverage(self, doc, data, content):
        self.docx.add_section_header(doc, "Your Market Coverage")
        self.docx.add_body_text(doc, content)

        # Venue categories grid
        self.docx.add_sub_header(doc, "WHERE YOUR ADS PLAY")
        self.docx.add_venue_categories(doc)

        # Expand your reach callout
        expanding = [k for k, v in self.config["markets"].items() if v["status"] == "expanding"]
        if expanding:
            self.docx.add_sub_header(doc, "EXPAND YOUR REACH")
            self.docx.add_body_text(
                doc,
                f"As {data.business_name} grows, MCTV grows with you. "
                f"MCTV is currently expanding into {' and '.join(expanding)}, "
                f"adding to the 30 screens in Starkville and 25 in Tupelo. "
                f"Multi-market packages available."
            )
        doc.add_page_break()

    def _build_pricing(self, doc, data):
        self.docx.add_section_header(doc, "Partnership Pricing")

        self.docx.add_body_text(
            doc,
            "Choose the coverage level that fits your goals. As you scale up, "
            "your cost per screen drops and your monthly ad plays multiply \u2014 "
            f"giving {data.business_name} more visibility at a better value."
        )

        if data.custom_pricing:
            # Custom pricing display
            self.docx.add_sub_header(doc, f"YOUR PARTNERSHIP PACKAGE")
            self.docx.add_body_text(
                doc,
                f"${data.custom_monthly_rate:,.0f} / month  \u2014  "
                f"{data.custom_screen_count} Screens\n\n"
                f"Your ad plays across all {data.custom_screen_count} screens \u2014 "
                f"restaurants, gyms, offices, and retail. "
                f"Each ad slot gets {self.config['network']['plays_per_hour']} plays per hour "
                f"on a {self.config['network']['content_loop_minutes']}-minute loop, "
                f"all day, every day."
            )
        else:
            # Standard 4-tier pricing table
            tiers = get_all_tiers(self.config)
            self.docx.add_pricing_table(doc, tiers)

        # Contract terms
        self.docx.add_sub_header(doc, "PARTNERSHIP TERMS")
        self.docx.add_contract_terms(doc, self.config)
        doc.add_page_break()

    def _build_why_mctv(self, doc, content):
        self.docx.add_section_header(doc, "Why MCTV Elite Advertising")
        self.docx.add_body_text(doc, content)

    def _build_getting_started(self, doc, data, content):
        self.docx.add_section_header(doc, "Let's Get Started")
        self.docx.add_body_text(doc, content)

        # Contact card
        rep = get_team_member(self.config, data.sales_rep)
        self.docx.add_sub_header(doc, "YOUR PARTNERSHIP CONTACT")

        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from services.docx_service import NAVY, GOLD, GRAY

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(rep["name"])
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = NAVY

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{rep['email']}\n{rep['phone']}")
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("MCTV Elite Advertising  |  MCTVofMS.com")
        run.font.size = Pt(10)
        run.font.color.rgb = GOLD
