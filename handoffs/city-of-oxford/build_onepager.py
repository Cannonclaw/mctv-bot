# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Build the one-page city partnership leave-behind.

A single branded page to hand a mayor or city communications staffer. It states
the public-service offer, the reach behind it, and nothing else — this is a
leave-behind, not a proposal, so it must stay on one page.

    python handoffs/city-of-oxford/build_onepager.py
    python handoffs/city-of-oxford/build_onepager.py --city Starkville

Output lands in output/city/ (gitignored), as .docx plus .pdf when LibreOffice
is available.
"""

import argparse
import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from services.config_service import (  # noqa: E402
    get_team_member, get_tier_impressions,
)
from services.docx_service import DocxService  # noqa: E402
import json  # noqa: E402

OUTPUT_DIR = PROJECT_ROOT / "output" / "city"

def _rate_for_screens(config: dict, screens: int) -> float:
    """Monthly rate of the largest pricing tier this screen count qualifies for.

    Prices the donated slot at what a paying advertiser would owe. Deliberately
    conservative: a count between tiers takes the lower tier's rate, so the
    figure we hand a city is never inflated.
    """
    tiers = sorted(config["pricing"]["elite_tiers"], key=lambda t: t["screens"])
    rate = tiers[0]["monthly_rate"]
    for tier in tiers:
        if screens >= tier["screens"]:
            rate = tier["monthly_rate"]
    return rate


def _format_impressions(value: float) -> str:
    """1140000 -> '1.1M+'; 380000 -> '380K+'."""
    if value >= 1_000_000:
        return f"{value / 1_000_000:.1f}M+"
    return f"{value / 1_000:.0f}K+"


def _add_header(docx: DocxService, doc, city: str):
    """Logo wordmark, title, and one-line subtitle. Compact — this is page one of one."""
    logo = PROJECT_ROOT / "assets" / "branding" / "mctv_logo_on_light.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(logo), width=Cm(5.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run(f"A Community Partnership with the City of {city}")
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = docx.c["primary"]

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(2)
    sub.paragraph_format.space_after = Pt(6)
    run = sub.add_run(
        "A standing public-service slot on every screen we operate in "
        f"{city} — at no cost to the City."
    )
    run.font.size = Pt(10.5)
    run.font.name = "Arial"
    run.font.color.rgb = docx.c["text"]


def _add_contact(docx: DocxService, doc, rep: dict):
    """Closing rule plus a single contact line."""
    docx.add_section_divider(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{rep['name']}  ·  {rep['title']}")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = docx.c["primary"]

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(
        f"{rep['phone']}  ·  {rep['email']}  ·  "
        "MCTV Elite Advertising, Oxford, Mississippi"
    )
    run2.font.size = Pt(9)
    run2.font.name = "Arial"
    run2.font.color.rgb = docx.c["gray"]


def build(city: str = "Oxford", rep_name: str = "T. Creed Cannon",
          alert_turnaround: str = "same business day",
          color_scheme: str = "original") -> Path:
    """Build the one-pager for a city and return the saved .docx path."""
    config = json.loads(
        (PROJECT_ROOT / "config" / "config.json").read_text(encoding="utf-8")
    )

    market = config["markets"].get(city)
    if market is None:
        raise SystemExit(
            f"Unknown market {city!r}. Known: {', '.join(config['markets'])}"
        )

    screens = market["screens"]
    impressions = get_tier_impressions(config, screens)
    monthly_rate = _rate_for_screens(config, screens)
    annual_value = monthly_rate * 12
    network = config["network"]
    rep = get_team_member(config, rep_name)

    docx = DocxService(config, color_scheme=color_scheme)
    doc = docx.create_document()

    _add_header(docx, doc, city)

    docx.add_metrics_banner(doc, {
        str(screens): f"Screens Across\n{city}",
        _format_impressions(impressions): f"Monthly Impressions\nin {city}",
        f"{network['avg_dwell_time_minutes']}+ Min": "Average Time\nOn Location",
        "0%": "Skipped, Blocked,\nor Scrolled Past",
    })

    docx.add_sub_header(doc, "What the City Receives")

    docx.add_accent_card(
        doc,
        "A permanent public-service slot",
        f"One spot in the rotation on all {screens} {city} screens, running "
        f"{network['plays_per_hour']} times an hour, every hour we are on. The "
        "City decides what runs in it. We build the artwork at no charge and "
        "change it whenever you need it changed.",
    )

    docx.add_accent_card(
        doc,
        "Priority for urgent messages",
        "Boil water notices, severe weather, road and lane closures, missing "
        f"persons. Send it to us and we push it ahead of the normal rotation, "
        f"{alert_turnaround}.",
    )

    docx.add_accent_card(
        doc,
        "A report on what it did",
        "Every month, the same traction report our paying advertisers get: "
        "plays delivered, screens reached, estimated impressions. So the City "
        "knows what the space is actually worth to it.",
    )

    docx.add_sub_header(doc, "Where These Screens Are")

    docx.add_body_text(
        doc,
        f"MCTV operates {screens} indoor digital screens in {city} — in "
        "restaurants and bars, gyms, salons and barbershops, medical and dental "
        "waiting rooms, and retail counters. They reach people who are sitting "
        "still, not looking at a phone, and cannot scroll past. That is a room "
        "the City's website and social media do not reach.",
    )

    docx.add_callout_box(
        doc,
        f"What this is worth: an advertising package on {screens} {city} screens "
        f"sells for ${monthly_rate:,.0f} a month — ${annual_value:,.0f} a year. "
        f"We are donating it to the City of {city}. No cost, no contract, no "
        "purchase order, nothing for the Board to approve. We are a local "
        "company and these are our neighbors too.",
    )

    _add_contact(docx, doc, rep)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / f"MCTV_{city}_City_Partnership_OnePager.docx"
    doc.save(out_path)
    pdf = DocxService._convert_to_pdf(out_path)

    print(f"Saved: {out_path}")
    print(f"PDF:   {pdf if pdf else 'not generated (LibreOffice unavailable)'}")
    return out_path


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--city", default="Oxford",
                        help="Market name as spelled in config.json markets")
    parser.add_argument("--rep", default="T. Creed Cannon",
                        help="Team member whose contact info appears at the foot")
    parser.add_argument("--alert-turnaround", default="same business day",
                        help="Verified turnaround for pushing an urgent spot live. "
                             "Confirm the real number in n-compass before printing.")
    parser.add_argument("--color-scheme", default="original",
                        choices=["original", "light", "dark", "pastel"])
    args = parser.parse_args()

    build(city=args.city, rep_name=args.rep,
          alert_turnaround=args.alert_turnaround,
          color_scheme=args.color_scheme)


if __name__ == "__main__":
    main()
