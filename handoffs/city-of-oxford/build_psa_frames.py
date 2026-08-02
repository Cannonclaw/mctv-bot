# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Build the sample City public-service frames for the Oxford meeting.

Four 1920x1080 stills showing what the City could run on the network. These are
UNAPPROVED MOCKUPS for a partnership conversation — nothing here has been
issued by any authority and none of it should be scheduled until the City signs
off on the copy.

    python handoffs/city-of-oxford/build_psa_frames.py

Writes to output/city/psa/ (gitignored): a clean set for the printed handout, a
labeled set carrying a MOCKUP bar for anything that circulates digitally, and a
contact sheet of all four.

Design rules enforced in code rather than trusted to the eye — the build fails
if a frame breaks one. See LEGIBILITY below for where the numbers come from.
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass, field as dc_field
from pathlib import Path

from PIL import ImageDraw

sys.path.insert(0, str(Path(__file__).resolve().parent))
import psa_render as R  # noqa: E402

OUTPUT_DIR = Path(__file__).resolve().parent.parent.parent / "output" / "city" / "psa"

# ── Legibility budget ────────────────────────────────────────────────────────
# Sized for the worst common case: a 43" panel read at 25 feet by someone who is
# not trying to read it. Everything is a fraction of frame height so the spec
# survives a 32" salon panel and a 75" gym panel alike.
FRAME_H = 1080
SAFE_INSET = 0.05                  # 96px — bezel crop, overscan, off-axis
CAP_SIGNAL = 0.055                 # 59px  — categorical label
CAP_HEADLINE_MAX = 0.15            # 162px — the one thing a 2-second viewer reads
CAP_HEADLINE_MIN = 0.12            # 130px — hard floor, never ship below
CAP_ACTION = 0.070                 # 76px  — protective action detail
CAP_SCOPE = 0.058                  # 63px  — who and where (bold: stems must hold at 25ft)
CAP_TIME = 0.060                   # 65px  — when
CAP_TERTIARY = 0.040               # 43px  — source and confirmation (a 15ft element)
HEADLINE_MAX_WORDS = 5
HEADLINE_MAX_LINES = 2
FRAME_MAX_CHARS = 190
MIN_CONTRAST = 7.0                 # headline against its field; WCAG AA large is 3.0

# ── Severity tiers ───────────────────────────────────────────────────────────
# ANSI Z535 severity logic and its text pairings (orange and yellow are
# black-text colors), with fields darkened until they clear contrast on an
# emissive panel. Monotonic luminance ramp — dark for emergency up to bright for
# advisory — so severity survives colorblindness and a badly calibrated screen.
@dataclass(frozen=True)
class Tier:
    key: str
    field: tuple[int, int, int]
    ink: tuple[int, int, int]
    muted: tuple[int, int, int]


TIERS = {
    "emergency": Tier("emergency", (0xA4, 0x00, 0x0F), (0xFF, 0xFF, 0xFF), (0xF2, 0xC9, 0xCC)),
    "warning":   Tier("warning",   (0xF0, 0x8C, 0x00), (0x00, 0x00, 0x00), (0x4A, 0x2C, 0x00)),
    "advisory":  Tier("advisory",  (0xFF, 0xD1, 0x00), (0x00, 0x00, 0x00), (0x4A, 0x3D, 0x00)),
    "notice":    Tier("notice",    (0x0B, 0x3D, 0x91), (0xFF, 0xFF, 0xFF), (0xC2, 0xD2, 0xEF)),
}


@dataclass
class Frame:
    """One full-screen message.

    `signal` is the categorical label — the word people already know, and the
    ANSI redundancy layer backing the field color. `headline` is the imperative:
    hazard fused with protective action, because a viewer who reads exactly one
    thing must come away knowing what to do, not just what happened.
    """
    key: str
    tier: str
    signal: str
    headline: str
    action: str
    scope: str
    time_bound: str
    source: str
    confirm: str
    verify: list[str] = dc_field(default_factory=list)
    # Frames that impersonate a live alert too well keep the MOCKUP bar even in
    # the "clean" set. The tornado frame names a real federal issuing office
    # with real safety wording — photographed without the bar, it is
    # indistinguishable from a genuine NWS graphic.
    always_label: bool = False


def frames(city_domain: str) -> list[Frame]:
    return [
        Frame(
            key="boil_water",
            tier="warning",
            signal="BOIL WATER NOTICE",
            headline="Boil Your Water",
            action="Rolling boil, 1 full minute.",
            scope="[North Oxford service area]",
            time_bound="[Until further notice]",
            source="Oxford Utilities",
            confirm=city_domain,
            verify=[
                "Who issued it: a precautionary notice after a line break comes "
                "from Oxford Utilities; a notice after a bad sample comes from "
                "MSDH. The source line changes accordingly.",
                "The service area — system-wide, or a named area.",
                "Do not put an end time on this one. MSDH clears a notice only "
                "after two consecutive days of clean samples, so 'until further "
                "notice' is the honest line.",
                "'Rolling boil for one minute' is MSDH's own standard wording. "
                "Do not paraphrase it.",
            ],
        ),
        Frame(
            key="tornado_warning",
            tier="emergency",
            signal="TORNADO WARNING",
            headline="Take Shelter Now",
            action="Interior room. Lowest floor.",
            scope="[Oxford · W. Lafayette Co.]",
            time_bound="Until [9:45 PM CDT]",
            source="National Weather Service · Memphis",
            confirm="weather.gov",
            always_label=True,
            verify=[
                "Warnings for Lafayette County are issued by NWS Memphis, not by "
                "the City. Keeping NWS on the source line makes the screen match "
                "the phone alert people just got, which is what stops them "
                "second-guessing it.",
                "The expiration time has to be pulled from the live warning. A "
                "stale time is worse than showing nothing.",
                "This frame should take the whole screen with no ad rotation "
                "under it.",
                "Settle who at the City is authorized to trigger it, and how many "
                "seconds it takes to interrupt the loop network-wide.",
            ],
        ),
        Frame(
            key="road_closure",
            tier="advisory",
            signal="ROAD CLOSED",
            headline="[North Lamar] at [Molly Barr]",
            action="[Use posted detour]",
            scope="",
            time_bound="[Through Friday]",
            source="City of Oxford",
            confirm=city_domain,
            verify=[
                "Fill the street, the cross street, and the duration.",
                "If the closure is on Highway 7, University Avenue ramps, or any "
                "state route, the issuing authority is MDOT and not the City — "
                "change the source line or the mayor will catch it.",
                "The action line varies by closure: short event closures have no "
                "posted detour, so the alternates are 'Seek alternate route' or "
                "'Local traffic only'.",
            ],
        ),
        Frame(
            key="civic_election",
            tier="notice",
            signal="ELECTION DAY",
            headline="Vote Today",
            action="Polls open 7 a.m. to 7 p.m.",
            scope="[All wards] · [Date]",
            time_bound="",
            source="City of Oxford",
            confirm=f"Polling places: {city_domain}",
            verify=[
                "7 a.m. to 7 p.m. is Mississippi's statutory poll window and is "
                "safe to leave as-is; it does not change election to election.",
                "Fill the date and name the election (municipal general, primary, "
                "runoff, special).",
                "Jurisdiction check, verified in review: Oxford municipal "
                "elections are run by the Municipal Election Commission and "
                "organized by WARD, not precinct — and the Secretary of State's "
                "locator (sos.ms.gov) covers county-run elections only. This "
                "frame is the municipal version. A county/state election day "
                "version changes the source line to the Lafayette County "
                "Election Commission and the lookup to sos.ms.gov.",
                "This same frame becomes a public hearing notice or a city hiring "
                "notice by changing two lines — worth demonstrating in the room.",
            ],
        ),
    ]


# ── Contrast ─────────────────────────────────────────────────────────────────

def _luminance(rgb: tuple[int, int, int]) -> float:
    def channel(c: float) -> float:
        c = c / 255.0
        return c / 12.92 if c <= 0.04045 else ((c + 0.055) / 1.055) ** 2.4
    r, g, b = (channel(v) for v in rgb)
    return 0.2126 * r + 0.7152 * g + 0.0722 * b


def contrast(a: tuple[int, int, int], b: tuple[int, int, int]) -> float:
    la, lb = _luminance(a), _luminance(b)
    hi, lo = max(la, lb), min(la, lb)
    return (hi + 0.05) / (lo + 0.05)


# ── Type sizing ──────────────────────────────────────────────────────────────

_cap_ratio_cache: dict[str, float] = {}


def _cap_ratio(weight: str) -> float:
    """Cap height as a fraction of PIL's nominal font size, measured not assumed."""
    if weight not in _cap_ratio_cache:
        probe = R.font(weight, 400)
        box = probe.getbbox("H")
        _cap_ratio_cache[weight] = (box[3] - box[1]) / 400.0
    return _cap_ratio_cache[weight]


def size_for_cap(weight: str, cap_fraction: float) -> int:
    """PIL font size that yields a target cap height, given as a fraction of frame height."""
    return int(round(FRAME_H * cap_fraction / _cap_ratio(weight)))


def cap_fraction_of(weight: str, size: int) -> float:
    return (size * _cap_ratio(weight)) / FRAME_H


# ── Rendering ────────────────────────────────────────────────────────────────

def render(frame: Frame, labeled: bool = False) -> tuple:
    """Draw one frame. Returns (image, compliance report dict)."""
    tier = TIERS[frame.tier]
    img, draw = R.new_canvas(tier.field)

    inset_v = int(FRAME_H * SAFE_INSET)     # 54px — overscan and bezel crop
    left = int(1920 * SAFE_INSET)           # 96px — live area 1728x972
    live_w = 1920 - 2 * left

    report: dict = {"key": frame.key, "tier": frame.tier, "issues": []}

    # 1 — signal word: categorical, tracked, backs up the field color for anyone
    #     who cannot decode the hue.
    signal_size = size_for_cap("bold", CAP_SIGNAL)
    signal_font = R.font("bold", signal_size)
    y = inset_v
    R.draw_tracked(draw, (left, y), frame.signal, signal_font, tier.ink,
                   tracking=int(signal_size * 0.06))
    y += int(signal_size * 1.15)

    # Thin rule under the signal word, in the ink color — never a dark-on-dark
    # distinction, which is the first thing to dissolve off-axis.
    content_top = R.rule(draw, left, y + 8, live_w, 6, tier.ink) + int(FRAME_H * 0.035)

    # 6 and 7 — source and confirmation, on one baseline at the foot. Placed
    # before the body so the body gets a measured budget rather than a guess.
    # Attributed to the issuing AUTHORITY: the network's own logo is deliberately
    # absent, because a sign operator's mark on an emergency frame turns a
    # warning into an advertisement.
    tert_size = size_for_cap("regular", CAP_TERTIARY)
    tert_font = R.font("regular", tert_size)
    foot_y = 1080 - inset_v - tert_size
    R.rule(draw, left, foot_y - int(tert_size * 0.9), live_w, 3, tier.muted)
    draw.text((left, foot_y), frame.source, font=tert_font, fill=tier.ink)
    if frame.confirm:
        w = int(draw.textlength(frame.confirm, font=tert_font))
        draw.text((left + live_w - w, foot_y), frame.confirm, font=tert_font, fill=tier.ink)

    content_bottom = foot_y - int(tert_size * 1.2)
    available = content_bottom - content_top

    # Scope and time bound share one meta line. Two separate blocks put the frame
    # over the three-unit budget, and the spec is clear that the fix for
    # overflow is cutting copy, never shrinking type.
    meta = " · ".join(t for t in (frame.scope, frame.time_bound) if t)

    # 2 — headline: the imperative, and the only element allowed to dominate.
    # It gets the largest share of the budget and never yields it; if the copy
    # cannot fit at the 12% floor, that is a copy failure and gets reported.
    headline = frame.headline.upper() if len(frame.headline.split()) <= 3 else frame.headline
    head_block = R.fit_block(
        draw, headline, "bold", live_w, int(available * 0.62),
        max_size=size_for_cap("bold", CAP_HEADLINE_MAX),
        min_size=size_for_cap("bold", CAP_HEADLINE_MIN),
        line_spacing=1.02, max_lines=HEADLINE_MAX_LINES,
    )
    report["headline_cap_pct"] = round(cap_fraction_of("bold", head_block.fnt.size) * 100, 1)
    report["headline_lines"] = len(head_block.lines)
    report["headline_words"] = len(frame.headline.split())

    # Each subsequent block is fitted into what the ones above it left behind,
    # not into a fixed share of the frame. Fixed shares can sum past 100% and
    # silently overrun the footer.
    remaining = available - head_block.height

    # 3 — protective action detail, where the headline alone leaves a real gap.
    action_block = None
    if frame.action:
        action_block = R.fit_block(
            draw, frame.action, "bold", live_w, int(remaining * 0.55),
            max_size=size_for_cap("bold", CAP_ACTION),
            min_size=size_for_cap("bold", 0.060), line_spacing=1.12, max_lines=2,
        )
        remaining -= action_block.height

    # 4 and 5 — who, where, and until when.
    meta_block = None
    if meta:
        meta_block = R.fit_block(
            draw, meta, "bold", live_w, remaining,
            max_size=size_for_cap("bold", CAP_SCOPE),
            min_size=size_for_cap("bold", 0.034), line_spacing=1.10, max_lines=1,
        )

    # Stack them, spending whatever budget is left over as even breathing room.
    blocks = [b for b in (head_block, action_block, meta_block) if b]
    used = sum(b.height for b in blocks)
    slack = available - used
    # Breathing room is whatever is genuinely left over — a minimum gap that
    # pushes the stack past the footer is not breathing room, it is overflow.
    gap = max(0, min(int(FRAME_H * 0.045), slack // max(len(blocks) - 1, 1)))

    y = content_top + max(0, (available - used - gap * (len(blocks) - 1)) // 2)
    for i, blk in enumerate(blocks):
        y = R.draw_block(draw, blk, left, y, tier.ink)
        if i < len(blocks) - 1:
            y += gap

    report["overflow_px"] = max(0, y - content_bottom)
    report["meta_cap_pct"] = (round(cap_fraction_of("bold", meta_block.fnt.size) * 100, 1)
                              if meta_block else None)

    if labeled:
        _stamp_mockup(draw, tier)

    # ── compliance gate ──
    ratio = contrast(tier.ink, tier.field)
    report["contrast"] = round(ratio, 2)
    if ratio < MIN_CONTRAST:
        report["issues"].append(f"contrast {ratio:.2f}:1 below {MIN_CONTRAST}:1")
    if report["headline_cap_pct"] < CAP_HEADLINE_MIN * 100:
        report["issues"].append(
            f"headline cap {report['headline_cap_pct']}% below "
            f"{CAP_HEADLINE_MIN * 100:.0f}% floor"
        )
    if report["headline_words"] > HEADLINE_MAX_WORDS:
        report["issues"].append(f"headline {report['headline_words']} words > {HEADLINE_MAX_WORDS}")
    if report["headline_lines"] > HEADLINE_MAX_LINES:
        report["issues"].append(f"headline {report['headline_lines']} lines > {HEADLINE_MAX_LINES}")
    total_chars = sum(len(t) for t in (
        frame.signal, frame.headline, frame.action, frame.scope,
        frame.time_bound, frame.source, frame.confirm))
    report["total_chars"] = total_chars
    if total_chars > FRAME_MAX_CHARS:
        report["issues"].append(f"{total_chars} chars > {FRAME_MAX_CHARS}")
    if report["meta_cap_pct"] and report["meta_cap_pct"] < CAP_TERTIARY * 100:
        report["issues"].append(
            f"meta line at {report['meta_cap_pct']}% below the "
            f"{CAP_TERTIARY * 100:.1f}% tertiary floor"
        )
    if report["overflow_px"]:
        report["issues"].append(
            f"body overruns the footer by {report['overflow_px']}px"
        )

    return img, report


def _stamp_mockup(draw: ImageDraw.ImageDraw, tier: Tier):
    """Bar across the top for any copy that may travel outside the meeting.

    Degrades the design on purpose — it exists so a photograph of this frame can
    never be mistaken for a live municipal alert.
    """
    bar_h = 58
    draw.rectangle([0, 0, 1920, bar_h], fill=(0, 0, 0))
    fnt = R.font("bold", 30)
    text = "SAMPLE MOCKUP  ·  NOT AN ACTIVE ALERT  ·  NOT ISSUED BY ANY AUTHORITY"
    w = int(draw.textlength(text, font=fnt))
    draw.text(((1920 - w) // 2, (bar_h - 34) // 2), text, font=fnt, fill=(255, 255, 255))


# Short lines for the printed handout — what each frame is meant to prove. The
# frames themselves stay clean of any MCTV mark (a sign operator's logo on an
# emergency frame turns a warning into an advertisement), so the framing has to
# live here instead.
CAPTIONS = {
    "boil_water": "The workhorse. The most common real municipal alert — and the "
                  "one people miss on Facebook.",
    "tornado_warning": "Life safety. Full screen, no ad rotation underneath it.",
    "road_closure": "Routine. One message, one action.",
    "civic_election": "Non-emergency. The same frame becomes a public hearing "
                      "notice or a city job posting.",
}


def build_handout(out_dir: Path, city: str = "Oxford",
                  rep_name: str = "T. Creed Cannon") -> Path:
    """One printed page showing the four frames, to carry alongside the one-pager."""
    import json

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    root = Path(__file__).resolve().parent.parent.parent
    sys.path.insert(0, str(root))
    from services.config_service import get_team_member
    from services.docx_service import DocxService
    from build_onepager import _add_contact

    config = json.loads((root / "config" / "config.json").read_text(encoding="utf-8"))
    docx = DocxService(config)
    doc = docx.create_document()

    logo = root / "assets" / "branding" / "mctv_logo_on_light.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(logo), width=Cm(4.6))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run(f"What a City Message Looks Like on Our Screens")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = docx.c["primary"]

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(2)
    sub.paragraph_format.space_after = Pt(4)
    run = sub.add_run(
        f"Four samples, sized and colored to be read across a room in under two "
        f"seconds. Nothing here carries an MCTV logo — a City message should "
        f"look like it came from the City."
    )
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.font.color.rgb = docx.c["text"]

    # The printed page embeds the LABELED set. Paper leaves the room; the
    # SAMPLE MOCKUP bar is what keeps a photographed frame from passing for a
    # live alert.
    labeled = out_dir / "labeled"
    paths, captions = [], []
    for frame in frames(""):
        path = labeled / f"{frame.key}.png"
        if path.exists():
            paths.append(str(path))
            captions.append(CAPTIONS.get(frame.key, ""))
    docx.add_photos_grid(doc, paths, max_width=3.05, cols=2, captions=captions)

    docx.add_callout_box(
        doc,
        "These are unapproved samples, not live alerts. Copy, issuing authority, "
        f"and timing would all be the City of {city}'s call. What we are offering "
        "is the space and the speed.",
    )

    _add_contact(docx, doc, get_team_member(config, rep_name))

    out_path = out_dir / f"MCTV_{city}_Sample_City_Messages.docx"
    doc.save(out_path)
    DocxService._convert_to_pdf(out_path)
    return out_path


def build(city_domain: str = "oxfordms.net", out_dir: Path = OUTPUT_DIR) -> int:
    clean_dir, labeled_dir = out_dir / "clean", out_dir / "labeled"
    clean_dir.mkdir(parents=True, exist_ok=True)
    labeled_dir.mkdir(parents=True, exist_ok=True)

    reports, clean_paths = [], []
    for frame in frames(city_domain):
        img, report = render(frame, labeled=frame.always_label)
        path = clean_dir / f"{frame.key}.png"
        img.save(path)
        clean_paths.append(path)

        labeled_img, _ = render(frame, labeled=True)
        labeled_img.save(labeled_dir / f"{frame.key}.png")
        reports.append(report)

    R.contact_sheet(clean_paths, out_dir / "contact-sheet.png", cols=2)

    print(f"{'frame':<18}{'tier':<11}{'cap%':>6}{'lines':>7}{'words':>7}"
          f"{'chars':>7}{'contrast':>10}  issues")
    print("-" * 84)
    failures = 0
    for r in reports:
        issues = "; ".join(r["issues"]) if r["issues"] else "ok"
        failures += len(r["issues"])
        print(f"{r['key']:<18}{r['tier']:<11}{r['headline_cap_pct']:>6}"
              f"{r['headline_lines']:>7}{r['headline_words']:>7}"
              f"{r['total_chars']:>7}{r['contrast']:>9}:1  {issues}")

    print(f"\nclean:   {clean_dir}")
    print(f"labeled: {labeled_dir}")
    print(f"sheet:   {out_dir / 'contact-sheet.png'}")

    if not failures:
        # Only build the printed handout from frames that passed the gate.
        print(f"handout: {build_handout(out_dir)}")

    if failures:
        print(f"\nFAILED: {failures} legibility issue(s). Fix before printing.")
    return failures


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--city-domain", default="oxfordms.net",
                        help="City website shown on the confirmation line. "
                             "VERIFY this is the City's real domain before printing.")
    parser.add_argument("--out", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()
    raise SystemExit(1 if build(args.city_domain, args.out) else 0)


if __name__ == "__main__":
    main()
