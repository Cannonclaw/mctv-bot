# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Oxford location sheet for the Chamber packet.

Lists every active MCTV venue in Oxford and delineates venues that are
themselves members of the Oxford-Lafayette County Chamber of Commerce
(per data/oxford_chamber_members.json, cross-checked against the
chamber's online directory at business.oxfordms.com) from the rest of
the Oxford network.

Venue data comes from data/network_dashboard.json (NTV360 sync).

Usage:
    python scripts/generate_chamber_location_sheet.py
    python scripts/generate_chamber_location_sheet.py --scheme original
"""

import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.config_service import load_config
from services.docx_service import DocxService, PROJECT_ROOT

DASHBOARD_PATH = PROJECT_ROOT / "data" / "network_dashboard.json"
MEMBERS_PATH = PROJECT_ROOT / "data" / "oxford_chamber_members.json"

# NTV360 tags a few venues with the generic category "Building" — show
# something a customer would recognize instead.
CATEGORY_OVERRIDES = {
    "Cannon Chevrolet Buick Cadillac of Oxford": "Auto Dealer",
    "Cannon Collision Center": "Collision Center",
    "Rebel Yell Rooftop Bar": "Rooftop Bar",
}

# Short display names for the table (keeps rows on one line; the data
# file keeps full NTV360 names).
NAME_SHORTEN = {
    "Oxford-Lafayette County Chamber of Commerce": "Oxford-Lafayette County Chamber",
    "RedMed Urgent Clinic of Oxford - Jackson Ave": "RedMed Urgent Clinic",
    "Right Track Medical Group - Oxford": "Right Track Medical Group",
    "Uno Mas Tacos y Tequila Jackson Ave": "Uno Mas Tacos y Tequila",
    "The Velvet Ditch - Sports bar & Seafood": "The Velvet Ditch",
    "Internal Medicine Associates of Oxford": "Internal Medicine Associates",
    "4 Corner's Chevron (Chicken On A Stick)": "4 Corner's Chevron",
    "Little Bambini Children's Boutique": "Little Bambini",
    "Cannon Chevrolet Buick Cadillac of Oxford": "Cannon Chevrolet Buick Cadillac",
    "Mississippi Asthma & Allergy Clinic, P.A.": "MS Asthma & Allergy Clinic",
    "Booth's Barbeque and Yard": "Booth's Barbeque & Yard",
    "Rafters Music and Food": "Rafters Music & Food",
}


def street_address(full: str) -> str:
    """Trim ', Oxford MS 38655' and suite/unit fragments off a venue address."""
    addr = re.sub(r",?\s*Oxford\s+MS\s*\d*\s*$", "", full).strip().rstrip(",")
    return re.sub(r"\s*(,?\s*(Suite|Ste\.?|UNIT|Unit)\s*\S+|#\S+)\s*$", "", addr).strip()


def load_oxford_venues() -> list:
    data = json.loads(DASHBOARD_PATH.read_text(encoding="utf-8"))
    venues = [v for v in data["venues"].values()
              if "Oxford MS" in v.get("address", "")]
    venues.sort(key=lambda v: v["host_name"].lower())
    return venues


def load_membership() -> dict:
    if not MEMBERS_PATH.exists():
        return {"checked_at": "", "members": {}}
    return json.loads(MEMBERS_PATH.read_text(encoding="utf-8"))


def short_type(category: str) -> str:
    """Compress long NTV360 category labels so table rows stay single-line."""
    replacements = {
        "Barbecue Restaurant": "BBQ",
        "Mexican Restaurant": "Mexican",
        "Pizza Restaurant": "Pizza",
        "Seafood Restaurant": "Seafood",
        "Taco Restaurant": "Tacos",
        "Pho restaurant": "Pho",
        "Children's Clothing Store": "Kids' Clothing",
        "Men's Clothing Store": "Menswear",
        "Sports Nutrition Store": "Nutrition",
        "Live Music Bar": "Music Bar",
        "Urgent Care Center": "Urgent Care",
        "Mental Health Clinic": "Mental Health",
        "Chamber Of Commerce": "Chamber",
        "Recreation Center": "Recreation",
        "Health Food Store": "Health Food",
        "Flooring Store": "Flooring",
        "Party Store": "Party Goods",
    }
    return replacements.get(category, category)


def venue_rows(venues: list) -> list:
    rows = []
    for v in venues:
        imps = v.get("impressions") or 0
        category = CATEGORY_OVERRIDES.get(v["host_name"], v.get("category", ""))
        rows.append([
            NAME_SHORTEN.get(v["host_name"], v["host_name"]),
            short_type(category),
            street_address(v.get("address", "")),
            str(v.get("license_count", 1)),
            f"{imps:,.0f}",
        ])
    return rows


def set_column_widths(table, fractions: list, page_width_in: float = 6.3):
    """Rebalance data-table column widths (fractions must sum to ~1.0).

    Sets both the cell widths and the underlying tblGrid — LibreOffice
    lays the table out from the grid, so cell widths alone don't stick.
    """
    from docx.oxml.ns import qn
    from docx.shared import Inches

    table.autofit = False
    widths = [Inches(page_width_in * f) for f in fractions]
    grid = table._tbl.tblGrid
    for idx, grid_col in enumerate(grid.findall(qn("w:gridCol"))):
        grid_col.set(qn("w:w"), str(widths[idx].twips))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]


COLUMN_FRACTIONS = [0.35, 0.14, 0.23, 0.11, 0.17]


def generate_location_sheet(scheme: str = "original", also_pdf: bool = True) -> Path:
    config = load_config()
    docx = DocxService(config, color_scheme=scheme)

    venues = load_oxford_venues()
    membership = load_membership()
    member_status = membership.get("members", {})

    members = [v for v in venues
               if member_status.get(v["host_name"], {}).get("member") is True]
    others = [v for v in venues
              if member_status.get(v["host_name"], {}).get("member") is not True]
    unknowns = {v["host_name"] for v in venues
                if member_status.get(v["host_name"], {}).get("member") == "unknown"}

    total_screens = sum(v.get("license_count", 1) for v in venues)
    total_impressions = sum(v.get("impressions") or 0 for v in venues)
    member_screens = sum(v.get("license_count", 1) for v in members)

    doc = docx.create_document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

    # ── Header ──
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(0)
    brand.paragraph_format.space_after = Pt(2)
    brand_run = brand.add_run(config["company"]["name"].upper())
    brand_run.font.size = Pt(9)
    brand_run.font.bold = True
    brand_run.font.color.rgb = docx.c["accent"]

    docx.add_section_header(doc, "Oxford Screen Locations")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(4)
    sub.paragraph_format.space_after = Pt(4)
    run = sub.add_run("Chamber Packet Companion  •  Where Your Ad Plays in Oxford")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = docx.c["primary"]

    docx.add_metrics_banner(doc, {
        f"{len(venues)}": "Oxford\nVenues",
        f"{total_screens}": "Screens\nRunning Today",
        f"{total_impressions / 1000:,.0f}K": "Est. Monthly\nImpressions",
        f"{len(members)}": "Venues That Are\nChamber Members",
    })

    # ── Chamber member venues ──
    docx.add_sub_header(
        doc, f"CHAMBER MEMBER VENUES — YOUR AD RUNS IN FELLOW MEMBERS' ROOMS")
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(2)
    intro.paragraph_format.space_after = Pt(2)
    intro_run = intro.add_run(
        f"These {len(members)} venues ({member_screens} screens) are active "
        f"Oxford-Lafayette County Chamber members — including the Chamber's "
        f"own lobby. Your ad dollars circulate inside the member network."
    )
    intro_run.font.size = Pt(9.5)
    intro_run.font.color.rgb = docx.c["text"]

    headers = ["Venue", "Type", "Address", "Screens", "Imps/Mo"]
    table = docx.add_data_table(doc, headers, venue_rows(members))
    set_column_widths(table, COLUMN_FRACTIONS)

    # ── The rest of the Oxford network ──
    docx.add_sub_header(
        doc, "THE REST OF THE OXFORD NETWORK — INCLUDED IN YOUR ROTATION")
    intro2 = doc.add_paragraph()
    intro2.paragraph_format.space_before = Pt(2)
    intro2.paragraph_format.space_after = Pt(2)
    intro2_run = intro2.add_run(
        f"The Chamber deal doesn't stop at member venues — your ad also plays "
        f"across the other {len(others)} Oxford locations below."
    )
    intro2_run.font.size = Pt(9.5)
    intro2_run.font.color.rgb = docx.c["text"]

    other_rows = venue_rows(others)
    if unknowns:
        unknown_display = {NAME_SHORTEN.get(n, n) for n in unknowns}
        for row in other_rows:
            if row[0] in unknown_display:
                row[0] = f"{row[0]} †"
    table2 = docx.add_data_table(doc, headers, other_rows)
    set_column_widths(table2, COLUMN_FRACTIONS)

    # ── Source note ──
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(2)
    checked = membership.get("checked_at", "")
    note_text = (
        "Chamber membership cross-checked against the Oxford-Lafayette County "
        "Chamber of Commerce online directory (business.oxfordms.com)"
        + (f" on {checked}." if checked else ".")
        + " Impressions estimated from NTV360 venue traffic."
    )
    if unknowns:
        note_text += " † Membership could not be confirmed online."
    note_run = note.add_run(note_text)
    note_run.font.size = Pt(8)
    note_run.font.italic = True
    note_run.font.color.rgb = docx.c["gray"]

    # ── Contact line ──
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(4)
    contact.paragraph_format.space_after = Pt(0)
    contact_run = contact.add_run(
        f"{config['company']['name']}  •  {config['company']['tagline']}  •  "
        f"{config['company']['website']}"
    )
    contact_run.font.size = Pt(8)
    contact_run.font.color.rgb = docx.c["gray"]

    filename = "MCTV_Chamber_Oxford_Location_Sheet.docx"
    path = docx.save_proposal(doc, filename, also_pdf=also_pdf)
    return path


def main():
    parser = argparse.ArgumentParser(description="Generate the Oxford Chamber location sheet")
    parser.add_argument("--scheme", default="original",
                        choices=["original", "light", "dark", "pastel"],
                        help="Color scheme (default original)")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF conversion")
    args = parser.parse_args()

    path = generate_location_sheet(args.scheme, also_pdf=not args.no_pdf)
    print(f"Location sheet saved: {path}")


if __name__ == "__main__":
    main()
