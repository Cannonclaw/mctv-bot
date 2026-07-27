# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Chamber deal one-sheet: what $199/mo on 25 screens actually buys.

Single-page selling card for the Chamber of Commerce member deal.
Shows monthly plays, impressions, and CPM, compares MCTV against the
other six media channels, and breaks down the print and outdoor math
in dollars. All figures derive from config.json (network stats +
industry benchmarks) so the sheet stays current if those change.

Usage:
    python scripts/generate_chamber_one_sheet.py
    python scripts/generate_chamber_one_sheet.py --rate 199 --screens 25 --scheme original
"""

import argparse
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.config_service import load_config, get_tier_impressions, calculate_cpm
from services.docx_service import DocxService


def parse_cpm_range(est_cpm: str) -> tuple:
    """Parse a benchmark CPM string like '$20 - $40' into (20.0, 40.0)."""
    numbers = [float(n) for n in re.findall(r"[\d.]+", est_cpm)]
    if not numbers:
        return (0.0, 0.0)
    if len(numbers) == 1:
        return (numbers[0], numbers[0])
    return (numbers[0], numbers[1])


def impressions_bought(spend: float, cpm_low: float, cpm_high: float) -> str:
    """What a given spend buys at a CPM range, as a readable range string."""
    if cpm_low <= 0 or cpm_high <= 0:
        return "—"
    high = spend / cpm_low * 1000
    low = spend / cpm_high * 1000
    return f"~{low / 1000:,.0f}K–{high / 1000:,.0f}K"


def cost_of_impressions(impressions: float, cpm_low: float, cpm_high: float) -> str:
    """What it costs to buy a given impression count at a CPM range."""
    if cpm_low <= 0 or cpm_high <= 0:
        return "—"
    low = impressions / 1000 * cpm_low
    high = impressions / 1000 * cpm_high
    return f"${low:,.0f}–${high:,.0f}/mo"


def find_benchmark(config: dict, medium: str) -> dict:
    for b in config.get("industry_benchmarks", {}).get("media_comparison", []):
        if b["medium"] == medium:
            return b
    return {}


def generate_one_sheet(rate: float, screens: int, scheme: str = "original",
                       also_pdf: bool = True) -> Path:
    config = load_config()
    docx = DocxService(config, color_scheme=scheme)
    network = config["network"]

    # ── Deal math (all derived from config) ──
    plays_per_hour = network["plays_per_hour"]
    hours_per_day = network["hours_per_day"]
    days_per_month = network["days_per_month"]
    monthly_plays = plays_per_hour * hours_per_day * days_per_month * screens
    daily_plays = plays_per_hour * hours_per_day * screens
    monthly_impressions = get_tier_impressions(config, screens)
    cpm = calculate_cpm(rate, monthly_impressions)
    daily_cost = rate / days_per_month
    cost_per_screen = rate / screens
    impressions_per_dollar = monthly_impressions / rate
    dwell = network["avg_dwell_time_minutes"]

    # Published-rate anchor: what these screens cost at the closest standard tier
    tiers = config["pricing"]["elite_tiers"]
    anchor_tier = min(tiers, key=lambda t: abs(t["screens"] - screens))
    standard_cost = anchor_tier["cost_per_screen"] * screens
    savings_pct = (1 - rate / standard_cost) * 100 if standard_cost else 0

    print_bench = find_benchmark(config, "Print / Newspaper")
    outdoor_bench = find_benchmark(config, "Outdoor Billboard")
    print_lo, print_hi = parse_cpm_range(print_bench.get("est_cpm", ""))
    out_lo, out_hi = parse_cpm_range(outdoor_bench.get("est_cpm", ""))

    doc = docx.create_document()

    # One-sheet: tighter vertical margins than the standard proposal template
    # so everything lands on a single page.
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

    # ── Compact branded header (no full cover page — this is a one-sheet) ──
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(0)
    brand.paragraph_format.space_after = Pt(2)
    brand_run = brand.add_run(config["company"]["name"].upper())
    brand_run.font.size = Pt(9)
    brand_run.font.bold = True
    brand_run.font.color.rgb = docx.c["accent"]

    docx.add_section_header(doc, "Chamber Member Exclusive")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(4)
    sub.paragraph_format.space_after = Pt(4)
    run = sub.add_run(
        f"{screens} Screens Across North Mississippi  •  "
        f"${rate:,.0f}/Month  •  Cancel-Proof Visibility, Every Day"
    )
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = docx.c["primary"]

    # ── What the deal delivers every month ──
    docx.add_metrics_banner(doc, {
        f"{screens}": "Screens In\nLocal Venues",
        f"{monthly_plays:,}": "Ad Plays\nEvery Month",
        f"{monthly_impressions / 1000:,.0f}K": "Monthly\nImpressions",
        f"${cpm:.2f}": "Cost Per 1,000\nImpressions (CPM)",
    })

    # ── Full media comparison (MCTV row shows this deal's real numbers) ──
    docx.add_competitive_comparison(doc, monthly_rate=rate, screen_count=screens,
                                    monthly_impressions=monthly_impressions)

    # ── The print & outdoor math, dollar for dollar ──
    docx.add_sub_header(doc, f"THE PRINT & OUTDOOR MATH — SAME ${rate:,.0f}, DOLLAR FOR DOLLAR")
    docx.add_data_table(
        doc,
        headers=["Medium", f"What ${rate:,.0f} Buys",
                 f"Price of {monthly_impressions / 1000:,.0f}K Impressions", "Where You Show Up"],
        rows=[
            ["Print / Newspaper",
             f"{impressions_bought(rate, print_lo, print_hi)} impressions",
             cost_of_impressions(monthly_impressions, print_lo, print_hi),
             "One ad, one page, one run"],
            ["Outdoor Billboard",
             f"{impressions_bought(rate, out_lo, out_hi)} impressions",
             cost_of_impressions(monthly_impressions, out_lo, out_hi),
             "One board, 5-second drive-by"],
            [f"MCTV Chamber Deal",
             f"{monthly_impressions:,.0f} impressions",
             f"${rate:,.0f}/mo",
             f"{screens} venues, {dwell}-min avg visit"],
        ],
    )
    # ── What stands out ──
    docx.add_sub_header(doc, "WHY THIS DEAL IS HARD TO BEAT")
    docx.add_selling_point(
        doc, f"{savings_pct:.0f}% below our published rates.",
        f"Standard rate is ${anchor_tier['cost_per_screen']:.2f}/screen — Chamber members "
        f"pay ${cost_per_screen:.2f}/screen for a bigger network than our "
        f"${config['pricing']['elite_tiers'][1]['monthly_rate']:,.0f}/mo tier.",
    )
    docx.add_selling_point(
        doc, f"A {dwell}-minute audience, not a 5-second glance.",
        f"One outdoor board starts around "
        f"{outdoor_bench.get('monthly_cost', '$1,500 - $4,000').split(' - ')[0]}/mo for a "
        f"passing look at 60 mph. These screens play your ad {plays_per_hour}x an hour where "
        f"customers sit for {dwell} minutes — {daily_plays:,} plays a day.",
    )
    docx.add_selling_point(
        doc, "It cannot be skipped, tossed, or scrolled past.",
        "Newspapers get recycled and radio gets changed. These screens run 12 hours a day, "
        "with real play counts in the NTV360 dashboard.",
    )

    # ── Bottom line ──
    docx.add_callout_box(
        doc,
        f"THE BOTTOM LINE: ${daily_cost:.2f} a day = {screens} screens, {monthly_plays:,} plays, "
        f"{monthly_impressions:,.0f} local impressions a month. Every dollar buys "
        f"~{impressions_per_dollar:,.0f} impressions; print buys ~{1000 / print_hi:,.0f}–"
        f"{1000 / print_lo:,.0f}. Chamber members only — screens are limited.",
    )

    # ── Contact line (footer stand-in; add_footer skips single-section docs) ──
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(6)
    contact.paragraph_format.space_after = Pt(0)
    contact_run = contact.add_run(
        f"{config['company']['name']}  •  {config['company']['tagline']}  •  "
        f"{config['company']['website']}"
    )
    contact_run.font.size = Pt(8)
    contact_run.font.color.rgb = docx.c["gray"]

    filename = f"MCTV_Chamber_Deal_One_Sheet_{screens}_Screens.docx"
    path = docx.save_proposal(doc, filename, also_pdf=also_pdf)
    return path


def main():
    parser = argparse.ArgumentParser(description="Generate the Chamber deal one-sheet")
    parser.add_argument("--rate", type=float, default=199.0, help="Monthly rate (default 199)")
    parser.add_argument("--screens", type=int, default=25, help="Screen count (default 25)")
    parser.add_argument("--scheme", default="original",
                        choices=["original", "light", "dark", "pastel"],
                        help="Color scheme (default original)")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF conversion")
    args = parser.parse_args()

    path = generate_one_sheet(args.rate, args.screens, args.scheme,
                              also_pdf=not args.no_pdf)
    print(f"One-sheet saved: {path}")


if __name__ == "__main__":
    main()
