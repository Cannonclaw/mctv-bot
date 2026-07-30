# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
# Proprietary and confidential. Unauthorized copying, distribution,
# or modification of this file is strictly prohibited.
"""Letterhead template generator.

Produces three editable Word letterheads (plus PDF proofs) that share the
brand marks used everywhere else in the app:

  1. editorial  - cream stock, centered wordmark, hairline rules. Quietest option.
  2. masthead   - white stock, full-bleed navy band with the reversed-out mark.
  3. team       - cream stock, headshots + titles + direct phone and email.

Run:  python scripts/generate_letterhead.py
Out:  output/letterhead/MCTV_Letterhead_<variant>.docx  (+ .pdf when LibreOffice
      is available)
"""

from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from scripts.letterhead_assets import build_all  # noqa: E402

OUT_DIR = ROOT / "output" / "letterhead"

# --- Palette -----------------------------------------------------------------
# Navy and gold come from config.json; steel is sampled from the logo's own
# "ELITE ADVERTISING" line so type and mark agree.
NAVY = "1B1F3B"
STEEL = "486B8F"
GOLD = "C5A55A"
CREAM = "FBF8F1"
WHITE = "FFFFFF"
MUTED = "6E7480"

FONT = "Arial"

# Optional mailing address. Leave empty to omit the line entirely; drop the
# street address here once you want it printed on every sheet.
ADDRESS_LINE = ""

DOT = "  •  "
THIN_DOT = " · "  # tighter separator for lines that have to hold one line


# --- Low-level docx helpers --------------------------------------------------


def _el(tag: str, **attrs) -> OxmlElement:
    node = OxmlElement(tag)
    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), str(value))
    return node


# Elements that must precede <w:displayBackgroundShape/> in settings.xml.
# CT_Settings is an ordered sequence and Word refuses to open a file whose
# settings are out of order, so the flag has to be slotted in, not appended.
_SETTINGS_BEFORE_BACKGROUND = (
    "w:writeProtection",
    "w:view",
    "w:zoom",
    "w:removePersonalInformation",
    "w:removeDateAndTime",
    "w:doNotDisplayPageBoundaries",
)


def set_page_color(document: Document, hex_color: str) -> None:
    """Tint the page (cream stock) and tell Word to actually display it."""
    # <w:background> must be the first child of <w:document>.
    document.element.insert(0, _el("w:background", color=hex_color))

    settings = document.settings.element
    if settings.find(qn("w:displayBackgroundShape")) is not None:
        return

    index = 0
    for position, child in enumerate(settings):
        if child.tag in {qn(tag) for tag in _SETTINGS_BEFORE_BACKGROUND}:
            index = position + 1
    settings.insert(index, _el("w:displayBackgroundShape"))


def run(paragraph, text: str, *, size=9, color=NAVY, bold=False, italic=False,
        spacing: float = 0.0, caps=False, font=FONT):
    """Add a run. `spacing` is extra letter-spacing in points."""
    r = paragraph.add_run(text.upper() if caps else text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    # Arial for the East Asian/complex-script slots too, so Word never swaps in
    # a fallback face mid-line.
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _el("w:rFonts")
        rpr.insert(0, rfonts)
    for slot in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{slot}"), font)
    if spacing:
        rpr.append(_el("w:spacing", val=int(round(spacing * 20))))
    return r


def para(container, *, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0,
         line=None):
    p = container.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line is not None:
        pf.line_spacing = Pt(line)
    return p


def rule(paragraph, *, color=GOLD, weight=6, position="bottom", space=6) -> None:
    """Hairline rule drawn as a paragraph border. `weight` is in eighths of a point."""
    pbdr = _el("w:pBdr")
    pbdr.append(_el(f"w:{position}", val="single", sz=weight, space=space, color=color))
    paragraph._p.get_or_add_pPr().append(pbdr)


def shade_cell(cell, hex_color: str) -> None:
    cell._tc.get_or_add_tcPr().append(_el("w:shd", val="clear", fill=hex_color))


def shade_table(table, hex_color: str) -> None:
    """Shade the table itself as well as its cells.

    Without this, some renderers leave a hairline of unpainted page showing
    between two adjacent shaded cells - a visible seam down the navy band.
    """
    table._tbl.tblPr.append(_el("w:shd", val="clear", fill=hex_color))


def cell_margins(cell, *, top=0, bottom=0, left=0, right=0) -> None:
    """Cell padding in inches."""
    mar = _el("w:tcMar")
    for name, value in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        mar.append(_el(f"w:{name}", w=int(value * 1440), type="dxa"))
    cell._tc.get_or_add_tcPr().append(mar)


def border_cell(cell, hex_color: str) -> None:
    """Paint the cell's own edges in its fill color.

    Belt-and-braces with shade_table(): together they guarantee two adjacent
    navy cells read as one uninterrupted band with no seam at the joint.
    """
    borders = _el("w:tcBorders")
    for edge in ("top", "start", "bottom", "end"):
        borders.append(_el(f"w:{edge}", val="single", sz=8, space=0, color=hex_color))
    cell._tc.get_or_add_tcPr().append(borders)


def strip_table_borders(table) -> None:
    borders = _el("w:tblBorders")
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        borders.append(_el(f"w:{edge}", val="none", sz=0, space=0, color="auto"))
    table._tbl.tblPr.append(borders)


def bleed_table(table, *, left_margin_in: float) -> None:
    """Pull a table out past the text margins so shading runs edge to edge."""
    table._tbl.tblPr.append(_el("w:tblInd", w=int(-left_margin_in * 1440), type="dxa"))
    table._tbl.tblPr.append(_el("w:tblLayout", type="fixed"))


def fixed_row_height(row, inches: float) -> None:
    trpr = row._tr.get_or_add_trPr()
    trpr.append(_el("w:trHeight", val=int(inches * 1440), hRule="exact"))


def clear_para(paragraph) -> None:
    """Zero out a paragraph so an empty table cell adds no vertical slack."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(1)


def add_picture(paragraph, path: Path, width_in: float):
    r = paragraph.add_run()
    r.add_picture(str(path), width=Inches(width_in))
    return r


def set_margins(section, *, top, bottom, left, right, header, footer) -> None:
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.header_distance = Inches(header)
    section.footer_distance = Inches(footer)


def body_placeholder(document: Document, cfg: dict) -> None:
    """Sample letter copy so the proof shows real spacing. Delete when writing."""
    p = para(document, space_after=14)
    run(p, "Month Day, Year", size=10, color=MUTED)

    for line in ("Recipient Name", "Business Name", "Street Address", "City, MS ZIP"):
        p = para(document, space_after=0, line=13)
        run(p, line, size=10.5, color=NAVY)

    p = para(document, space_before=16, space_after=10)
    run(p, "Dear Recipient,", size=10.5, color=NAVY)

    p = para(document, space_after=10, line=15)
    run(
        p,
        "Body copy starts here. Replace this block with your letter. The margins, "
        "type size, and line spacing are already set, so anything you type picks up "
        "the letterhead formatting automatically.",
        size=10.5,
        color="222630",
    )

    p = para(document, space_before=14, space_after=0)
    run(p, "Sincerely,", size=10.5, color=NAVY)
    p = para(document, space_before=26, space_after=0, line=13)
    run(p, cfg["team"][0]["name"], size=10.5, bold=True, color=NAVY)
    p = para(document, space_after=0, line=13)
    run(p, cfg["team"][0]["title"], size=8.5, color=STEEL, spacing=0.6, caps=True)


# --- Variant 1: Editorial ----------------------------------------------------


def build_editorial(cfg: dict, art: dict) -> Path:
    """Cream stock. Centered wordmark, letterspaced tagline, hairline rules."""
    doc = Document()
    set_page_color(doc, CREAM)
    section = doc.sections[0]
    set_margins(section, top=2.35, bottom=1.5, left=1.2, right=1.2, header=0.85, footer=0.6)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_picture(p, art["wordmark_navy"], 2.45)

    p = para(header, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=7)
    run(p, cfg["company"]["tagline"], size=6.5, color=STEEL, spacing=2.4, caps=True)
    rule(p, color=GOLD, weight=4, space=10)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(0)
    rule(p, color=GOLD, weight=4, position="top", space=8)
    run(p, DOT.join(cfg["markets"].keys()), size=6.5, color=NAVY, spacing=1.9, caps=True)

    contact = [cfg["company"]["website"].replace("www.", "")]
    if ADDRESS_LINE:
        contact.insert(0, ADDRESS_LINE)
    contact += [cfg["team"][0]["phone"], cfg["team"][0]["email"]]
    p = para(footer, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, DOT.join(contact), size=7.5, color=STEEL)

    body_placeholder(doc, cfg)
    return _save(doc, "editorial")


# --- Variant 2: Masthead -----------------------------------------------------


def build_masthead(cfg: dict, art: dict) -> Path:
    """White stock. Full-bleed navy band, reversed-out mark, gold rule."""
    doc = Document()
    section = doc.sections[0]
    left = 1.05
    set_margins(section, top=2.1, bottom=1.15, left=left, right=left, header=0.0, footer=0.45)

    header = section.header
    header.is_linked_to_previous = False
    clear_para(header.paragraphs[0])

    page_width = 8.5

    # One navy cell painted edge to edge, with the two-column layout nested
    # inside it. Splitting the band itself into two cells leaves a hairline
    # tick where the cells meet; nesting keeps the paint unbroken.
    band = header.add_table(rows=1, cols=1, width=Inches(page_width))
    strip_table_borders(band)
    bleed_table(band, left_margin_in=left)
    shade_table(band, NAVY)
    fixed_row_height(band.rows[0], 1.3)

    band_cell = band.rows[0].cells[0]
    band_cell.width = Inches(page_width)
    shade_cell(band_cell, NAVY)
    border_cell(band_cell, NAVY)
    band_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(band_cell)
    clear_para(band_cell.paragraphs[0])

    inner = band_cell.add_table(rows=1, cols=2)
    inner.autofit = False
    inner._tbl.tblPr.append(_el("w:tblW", w=int(page_width * 1440), type="dxa"))
    strip_table_borders(inner)
    inner._tbl.tblPr.append(_el("w:tblLayout", type="fixed"))

    logo_cell, info_cell = inner.rows[0].cells
    logo_cell.width = Inches(3.5)
    info_cell.width = Inches(page_width - 3.5)
    for cell in (logo_cell, info_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_margins(logo_cell, left=left, right=0.1)
    p = logo_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_picture(p, art["wordmark_white"], 2.3)

    cell_margins(info_cell, left=0.1, right=left)
    p = info_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    run(p, cfg["company"]["website"].replace("www.", ""), size=9, color=WHITE,
        bold=True, spacing=1.0)

    p = para(info_cell, align=WD_ALIGN_PARAGRAPH.RIGHT, line=12)
    run(p, f"{cfg['team'][0]['phone']}{DOT}{cfg['team'][0]['email']}", size=8, color="C9CEDD")

    p = para(info_cell, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=5, line=11)
    run(p, THIN_DOT.join(cfg["markets"].keys()), size=5.8, color=GOLD, spacing=0.5, caps=True)

    # A cell must end with a paragraph; collapse the spare ones around the
    # nested table so they do not push the band's contents off center.
    for spare in band_cell.paragraphs:
        clear_para(spare)

    # Gold hairline immediately under the band, bleeding the same distance.
    accent = header.add_table(rows=1, cols=1, width=Inches(page_width))
    strip_table_borders(accent)
    bleed_table(accent, left_margin_in=left)
    fixed_row_height(accent.rows[0], 0.045)
    accent_cell = accent.rows[0].cells[0]
    accent_cell.width = Inches(page_width)
    shade_cell(accent_cell, GOLD)
    cell_margins(accent_cell)
    clear_para(accent_cell.paragraphs[0])

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    rule(p, color="D8DCE4", weight=4, position="top", space=8)
    tail = [cfg["company"]["legal_name"] + ", Inc.", cfg["company"]["tagline"]]
    if ADDRESS_LINE:
        tail.append(ADDRESS_LINE)
    run(p, DOT.join(tail), size=7, color=MUTED, spacing=0.5)

    body_placeholder(doc, cfg)
    return _save(doc, "masthead")


# --- Variant 3: Team ---------------------------------------------------------


def build_team(cfg: dict, art: dict) -> Path:
    """Cream stock. Header mark plus a footer roster: headshot, title, phone, email."""
    doc = Document()
    set_page_color(doc, CREAM)
    section = doc.sections[0]
    left = 1.1
    set_margins(section, top=1.85, bottom=2.55, left=left, right=left, header=0.7, footer=0.42)

    header = section.header
    header.is_linked_to_previous = False
    clear_para(header.paragraphs[0])

    text_width = 8.5 - (left * 2)
    top_row = header.add_table(rows=1, cols=2, width=Inches(text_width))
    strip_table_borders(top_row)
    top_row._tbl.tblPr.append(_el("w:tblLayout", type="fixed"))
    logo_cell, tag_cell = top_row.rows[0].cells
    logo_cell.width = Inches(2.3)
    tag_cell.width = Inches(text_width - 2.3)
    for cell in (logo_cell, tag_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        cell_margins(cell)

    p = logo_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_picture(p, art["wordmark_navy"], 2.15)

    p = tag_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    run(p, cfg["company"]["tagline"], size=6.0, color=STEEL, spacing=0.7, caps=True)
    p = para(tag_cell, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=3, line=11)
    run(p, cfg["company"]["website"].replace("www.", ""), size=8.5, color=NAVY,
        bold=True, spacing=0.8)

    p = para(header, space_before=7, space_after=0, line=1)
    rule(p, color=GOLD, weight=4, space=0)

    # Footer roster.
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = Pt(1)
    rule(p, color=GOLD, weight=4, position="top", space=2)

    team = cfg["team"][:3]
    roster = footer.add_table(rows=1, cols=len(team), width=Inches(text_width))
    strip_table_borders(roster)
    roster._tbl.tblPr.append(_el("w:tblLayout", type="fixed"))
    column = text_width / len(team)

    for cell, member in zip(roster.rows[0].cells, team):
        cell.width = Inches(column)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Symmetric padding so each block centers on its own column.
        cell_margins(cell, left=0.06, right=0.06)

        key = _headshot_key(member["name"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        add_picture(p, art[f"headshot_{key}"], 0.64)

        p = para(cell, align=WD_ALIGN_PARAGRAPH.CENTER, line=10)
        run(p, member["name"], size=8.5, bold=True, color=NAVY, spacing=0.3)

        p = para(cell, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=1, line=9)
        run(p, member["title"], size=6, color=STEEL, spacing=1.1, caps=True)

        p = para(cell, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, line=10)
        run(p, member["phone"], size=7.5, color="333844")

        p = para(cell, align=WD_ALIGN_PARAGRAPH.CENTER, line=10)
        run(p, member["email"], size=7.5, color=STEEL)

    p = para(footer, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=9)
    markets = DOT.join(cfg["markets"].keys())
    if ADDRESS_LINE:
        markets = f"{ADDRESS_LINE}{DOT}{markets}"
    run(p, markets, size=6.5, color=MUTED, spacing=1.6, caps=True)

    body_placeholder(doc, cfg)
    return _save(doc, "team")


def _headshot_key(name: str) -> str:
    lowered = name.lower()
    if "creed" in lowered:
        return "creed"
    if "mary" in lowered:
        return "mary_michael"
    return "swayze"


# --- Output ------------------------------------------------------------------


def _save(doc: Document, variant: str) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / f"MCTV_Letterhead_{variant}.docx"
    doc.save(path)
    return path


def to_pdf(path: Path) -> Path | None:
    """Render a PDF proof via LibreOffice when it is installed."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(path.parent), str(path)],
        check=True,
        capture_output=True,
        timeout=180,
    )
    pdf = path.with_suffix(".pdf")
    return pdf if pdf.exists() else None


def main() -> None:
    cfg = json.loads((ROOT / "config" / "config.json").read_text())
    art = build_all()

    builders = (
        ("editorial", build_editorial),
        ("masthead", build_masthead),
        ("team", build_team),
    )
    for name, builder in builders:
        docx_path = builder(cfg, art)
        pdf_path = to_pdf(docx_path)
        suffix = f"  ->  {pdf_path.name}" if pdf_path else "  (no PDF: LibreOffice missing)"
        print(f"{name:10} {docx_path.relative_to(ROOT)}{suffix}")


if __name__ == "__main__":
    main()
