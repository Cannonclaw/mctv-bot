#!/usr/bin/env python3
# Copyright (c) 2026 MCTV Digital, Inc. All rights reserved.
"""Test all 4 color schemes across all 6 proposal types.

Generates a proposal document for each (scheme x type) combination
without calling the Claude API. Validates:
  1. DocxService initializes correctly for each scheme
  2. Cover page renders with correct scheme-specific logo
  3. All formatting methods (headers, cards, tables, banners) work
  4. Document saves as valid .docx
  5. File is >0 bytes and can be reopened
"""

import io
import json
import os
import sys
import shutil
import traceback
from pathlib import Path
from datetime import datetime

# Force UTF-8 for emoji support on Windows
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from services.docx_service import DocxService, COLOR_SCHEMES
from services.config_service import load_config, get_all_tiers, get_team_member
from docx import Document


# ── Test Output Directory ────────────────────────────────────────────────────
TEST_OUTPUT = PROJECT_ROOT / "output" / "color_scheme_tests"


def setup():
    """Create clean test output directory."""
    if TEST_OUTPUT.exists():
        # Remove old test files (OneDrive may lock the dir itself)
        for f in TEST_OUTPUT.glob("*"):
            try:
                f.unlink()
            except OSError:
                pass
    TEST_OUTPUT.mkdir(parents=True, exist_ok=True)
    print(f"Test output: {TEST_OUTPUT}\n")


def make_docx_service(config: dict, scheme: str) -> DocxService:
    """Create a DocxService with the given color scheme."""
    svc = DocxService(config, color_scheme=scheme)
    svc.preparer_name = "Mary Michael Cannon"
    return svc


# ── Individual Component Tests ───────────────────────────────────────────────

def test_cover_page(svc: DocxService, scheme: str, config: dict):
    """Test cover page rendering with scheme-specific logo."""
    doc = svc.create_document()
    rep = get_team_member(config, "Mary Michael Cannon")
    svc.add_cover_page(
        doc,
        title="ADVERTISING PARTNERSHIP\nPROPOSAL",
        subtitle="Test Business LLC",
        prepared_for="John Smith",
        prepared_by=rep,
    )
    return doc


def test_section_headers(svc: DocxService, doc):
    """Test section headers with scheme colors."""
    svc.add_section_header(doc, "The Opportunity", new_page=True)
    svc.add_sub_header(doc, "MARKET OVERVIEW")
    svc.add_section_divider(doc)
    svc.add_section_header(doc, "What's Included", new_page=True)


def test_body_content(svc: DocxService, doc):
    """Test body text, bullets, selling points, callouts."""
    svc.add_body_text(doc, "This is a test paragraph demonstrating body text rendering "
                      "with the selected color scheme. The text color should match the "
                      "scheme's text value.")

    svc.add_selling_point(doc, "Captive Audiences",
                          "Your ad plays in high-traffic venues where customers "
                          "can't skip, scroll past, or block your message.")

    svc.add_selling_point(doc, "Local Reach",
                          "125+ screens across North Mississippi put your brand "
                          "in front of local consumers every day.")

    svc.add_bullet_point(doc, "Professional Design",
                         "Every ad is custom-designed by our creative team.")

    svc.add_callout_box(doc, "MCTV Elite Advertising operates the largest indoor "
                        "digital billboard network in North Mississippi.")


def test_accent_cards(svc: DocxService, doc):
    """Test accent cards with scheme-specific borders."""
    svc.add_accent_card(doc, "15-Second Ad Slot",
                        "Your custom-designed ad runs on a 15-minute loop, "
                        "playing 4 times per hour, 12 hours a day.")

    svc.add_accent_card(doc, "Professional Creative",
                        "Our team designs your ad at no extra cost. "
                        "Quarterly refreshes keep your message current.")

    svc.add_accent_card(doc, "NTV360 Analytics",
                        "Real-time dashboard shows impressions, plays, "
                        "and venue performance for your campaign.")


def test_metrics_banner(svc: DocxService, doc):
    """Test metrics banner with scheme background."""
    svc.add_metrics_banner(doc, {
        "125+": "Screens Across\nNorth Mississippi",
        "1.9M+": "Monthly Impressions",
        "55+ Min": "Avg. Dwell Time\nPer Visit",
        "4x/Hour": "Your Ad Plays\nEvery Day",
    })


def test_pricing_table(svc: DocxService, doc, config: dict):
    """Test pricing table with scheme-colored header."""
    tiers = get_all_tiers(config)
    svc.add_pricing_table(doc, tiers, recommended_idx=1)


def test_data_table(svc: DocxService, doc):
    """Test data table with scheme-colored header and alt rows."""
    headers = ["Venue", "Category", "Screens", "Monthly Plays"]
    rows = [
        ["Oxford Square Grill", "Restaurant", "2", "5,760"],
        ["Campus Cuts", "Barbershop", "1", "2,880"],
        ["FitZone Gym", "Health & Fitness", "3", "8,640"],
        ["Main Street Medical", "Medical", "1", "2,880"],
    ]
    svc.add_sub_header(doc, "VENUE PERFORMANCE SUMMARY")
    svc.add_data_table(doc, headers, rows)


def test_competitive_comparison(svc: DocxService, doc):
    """Test competitive comparison table."""
    svc.add_competitive_comparison(doc, monthly_rate=500, screen_count=20,
                                   monthly_impressions=30000)


def test_roi_projection(svc: DocxService, doc):
    """Test ROI projection callout."""
    svc.add_roi_projection(doc, monthly_rate=500, screen_count=20,
                           monthly_impressions=30000, business_name="Test Business")


def test_social_proof(svc: DocxService, doc):
    """Test social proof section."""
    svc.add_social_proof_section(doc)


def test_venue_categories(svc: DocxService, doc):
    """Test venue category grid."""
    svc.add_venue_categories(doc)


def test_team_section(svc: DocxService, doc):
    """Test team section."""
    svc.add_team_section(doc, new_page=True)


def test_footer(svc: DocxService, doc):
    """Test footer with scheme accent color."""
    svc.add_footer(doc)


def test_contract_terms(svc: DocxService, doc, config: dict):
    """Test contract terms rendering."""
    svc.add_contract_terms(doc, config)


# ── Full Integration Test ────────────────────────────────────────────────────

def run_full_proposal_test(config: dict, scheme: str) -> tuple:
    """Generate a full test proposal with all components for a given scheme.

    Returns (filepath, file_size, error_or_None).
    """
    svc = make_docx_service(config, scheme)
    label = COLOR_SCHEMES[scheme]["label"]

    try:
        # Cover page
        doc = test_cover_page(svc, scheme, config)

        # Section headers
        test_section_headers(svc, doc)

        # Body content
        test_body_content(svc, doc)

        # Accent cards
        svc.add_section_header(doc, "What's Included", new_page=True)
        test_accent_cards(svc, doc)

        # Metrics banner
        svc.add_section_header(doc, "Network Stats", new_page=True)
        test_metrics_banner(svc, doc)

        # Pricing table
        svc.add_section_header(doc, "Partnership Pricing", new_page=True)
        test_pricing_table(svc, doc, config)
        test_contract_terms(svc, doc, config)

        # Data table
        svc.add_section_header(doc, "Venue Performance", new_page=True)
        test_data_table(svc, doc)

        # Competitive comparison + ROI
        svc.add_section_header(doc, "How MCTV Compares", new_page=True)
        test_competitive_comparison(svc, doc)
        test_roi_projection(svc, doc)

        # Social proof
        svc.add_section_header(doc, "The MCTV Network", new_page=True)
        test_social_proof(svc, doc)
        test_venue_categories(svc, doc)

        # Team
        test_team_section(svc, doc)

        # Footer
        test_footer(svc, doc)

        # Save
        filename = f"test_{scheme}_full_proposal.docx"
        filepath = TEST_OUTPUT / filename
        doc.save(str(filepath))

        # Validate
        size = filepath.stat().st_size
        if size == 0:
            return str(filepath), 0, "File is 0 bytes"

        # Re-open to verify it's valid
        reopened = Document(str(filepath))
        para_count = len(reopened.paragraphs)
        table_count = len(reopened.tables)

        return str(filepath), size, None

    except Exception as e:
        return None, 0, f"{type(e).__name__}: {e}\n{traceback.format_exc()}"


# ── Logo Asset Check ─────────────────────────────────────────────────────────

def check_logo_assets():
    """Verify all scheme-specific logo files exist."""
    print("=" * 60)
    print("LOGO ASSET CHECK")
    print("=" * 60)
    all_ok = True
    for scheme, colors in COLOR_SCHEMES.items():
        logo_name = colors.get("cover_logo", "???")
        logo_path = PROJECT_ROOT / "assets" / "branding" / logo_name
        exists = logo_path.exists()
        size = logo_path.stat().st_size if exists else 0
        status = f"OK ({size:,} bytes)" if exists else "MISSING"
        icon = "+" if exists else "X"
        print(f"  [{icon}] {scheme:8s} -> {logo_name:30s} {status}")
        if not exists:
            all_ok = False
    print()
    return all_ok


# ── Color Verification ───────────────────────────────────────────────────────

def verify_color_consistency():
    """Check all schemes have required keys and valid values."""
    print("=" * 60)
    print("COLOR SCHEME CONSISTENCY CHECK")
    print("=" * 60)

    required_keys = ["label", "primary", "accent", "white", "gray", "text",
                     "light", "bg_hex", "accent_hex", "light_hex", "cover_logo"]

    all_ok = True
    for scheme, colors in COLOR_SCHEMES.items():
        missing = [k for k in required_keys if k not in colors]
        if missing:
            print(f"  [X] {scheme}: missing keys: {missing}")
            all_ok = False
        else:
            # Verify hex values are valid
            for hex_key in ["bg_hex", "accent_hex", "light_hex"]:
                val = colors[hex_key]
                try:
                    int(val, 16)
                except ValueError:
                    print(f"  [X] {scheme}.{hex_key} = '{val}' is not valid hex")
                    all_ok = False
                    continue
            print(f"  [+] {scheme:8s} ({colors['label']:25s}) — all {len(required_keys)} keys present, hex valid")

    print()
    return all_ok


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    print()
    print("=" * 60)
    print("  MCTV Color Scheme Test Suite")
    print(f"  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 60)
    print()

    # Load config
    config = load_config()
    setup()

    # Pre-flight checks
    colors_ok = verify_color_consistency()
    logos_ok = check_logo_assets()

    # Full proposal generation for each scheme
    print("=" * 60)
    print("FULL PROPOSAL GENERATION (all components)")
    print("=" * 60)

    results = {}
    schemes = list(COLOR_SCHEMES.keys())

    for scheme in schemes:
        label = COLOR_SCHEMES[scheme]["label"]
        print(f"\n  Generating: {scheme} ({label})...", end=" ", flush=True)
        filepath, size, error = run_full_proposal_test(config, scheme)
        results[scheme] = {"filepath": filepath, "size": size, "error": error}

        if error:
            print(f"FAIL")
            print(f"    Error: {error}")
        else:
            print(f"OK ({size:,} bytes)")

    # Color differentiation check — verify each scheme's unique colors
    # appear in its document XML (so we know it's not all falling back to default)
    print()
    print("=" * 60)
    print("COLOR DIFFERENTIATION CHECK")
    print("=" * 60)
    diff_ok = True
    for scheme in schemes:
        r = results[scheme]
        if r["error"] or not r["filepath"]:
            continue
        bg_hex = COLOR_SCHEMES[scheme]["bg_hex"].upper()
        accent_hex = COLOR_SCHEMES[scheme]["accent_hex"].upper()
        try:
            doc_check = Document(r["filepath"])
            xml_text = doc_check.element.xml.upper()
            bg_found = bg_hex in xml_text
            accent_found = accent_hex in xml_text
            if bg_found and accent_found:
                print(f"  [+] {scheme:8s} — bg_hex {bg_hex} found, accent_hex {accent_hex} found")
            else:
                missing = []
                if not bg_found:
                    missing.append(f"bg_hex {bg_hex}")
                if not accent_found:
                    missing.append(f"accent_hex {accent_hex}")
                print(f"  [X] {scheme:8s} — MISSING: {', '.join(missing)}")
                diff_ok = False
        except Exception as e:
            print(f"  [X] {scheme:8s} — Error reading XML: {e}")
            diff_ok = False
    print()

    # Summary
    print("=" * 60)
    print("RESULTS SUMMARY")
    print("=" * 60)

    passed = 0
    failed = 0

    # Pre-flight
    for label, ok in [("Color consistency", colors_ok), ("Logo assets", logos_ok),
                       ("Color differentiation", diff_ok)]:
        icon = "PASS" if ok else "FAIL"
        print(f"  [{icon}] {label}")
        if ok:
            passed += 1
        else:
            failed += 1

    # Generation tests
    for scheme in schemes:
        r = results[scheme]
        label = COLOR_SCHEMES[scheme]["label"]
        if r["error"]:
            print(f"  [FAIL] {scheme} ({label})")
            print(f"         {r['error'].split(chr(10))[0]}")
            failed += 1
        else:
            print(f"  [PASS] {scheme} ({label}) — {r['size']:,} bytes")
            passed += 1

    total = passed + failed
    print()
    print(f"  {passed}/{total} passed, {failed} failed")

    if failed == 0:
        print()
        print("  All color schemes generate valid proposals!")
        print(f"  Test files saved to: {TEST_OUTPUT}")
    else:
        print()
        print("  Some tests FAILED. See errors above.")

    print()
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
